"""
DSCR Advisor-Grade Decision Engine: Cross-Document Synthesis v2.0
================================================================
Covers ALL material in workspace:
  - 17 user-supplied docs (8 MDs + 9 PDFs from initial round)
  - 5 additional PDFs (Round 22)
  - 6 Sprint docs (Sprint 0-6 research execution)
  - 3 Sovereign Master docs (definitive blueprint, complete sovereign, definitive product spec)
  - v14/v15/v16 master consolidated specs
  - DSCR Forumals, Dual Truth Engine research, Master Knowledge Document, 2026 Master Knowledge Paper
  - Definitive Blueprint correction log, architectural-debt analysis
  - Slice 1 source code (payment/dscr/leverage/ltv/compliance)

Total: ~60+ source documents, ~3,000+ pages analyzed.
Output: DSCR_Advisor_Engine_Cross_Doc_Synthesis_v2_20260619.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import os

# ============== CONFIG ==============
OUT_DIR = Path(r"C:\Users\serge\OneDrive\Documents\DSCR_LOAN OFFICE\output\doc")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_FILE = OUT_DIR / "DSCR_Advisor_Engine_Cross_Doc_Synthesis_v2_20260619.docx"

# ============== STYLE HELPERS ==============
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
ACCENT = RGBColor(0x0B, 0x5E, 0x4B)
WARN = RGBColor(0xB7, 0x47, 0x1A)
CRIT = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x7E, 0x34)

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(20)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(15)
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(12.5)
    return p

def add_h4(doc, text):
    p = doc.add_heading(text, level=4)
    for run in p.runs:
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(11)
    return p

def add_para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0, color=None, bold=False):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)
        if color:
            run.font.color.rgb = color
        if bold:
            run.font.bold = True
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, widths=None, header_bg="1F3A5F"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, 1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F0F4F8")
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_callout(doc, text, color=WARN, label="WARNING"):
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}] {text}")
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF4E5' if color == WARN else 'FFE5E5' if color == CRIT else 'E5F5E5')
    pPr.append(shd)
    return p

# ============== BUILD DOC ==============
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ============== TITLE PAGE ==============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("DSCR Advisor-Grade Decision Engine")
trun.font.size = Pt(28); trun.font.bold = True; trun.font.color.rgb = NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = subtitle.add_run("Cross-Document Synthesis v2.0 — Full Workspace Audit")
srun.font.size = Pt(18); srun.font.italic = True; srun.font.color.rgb = DARK_GRAY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mrun = meta.add_run("60+ source documents analyzed in full\n"
                    "8 master MDs + 14 architectural PDFs + 6 Sprint research execution reports\n"
                    "v14/v15/v16 master consolidated specs + Definitive Blueprint v3 + Definitive Product Spec v12\n"
                    "Slice 1 source code + audit + golden vectors + 132 tests")
mrun.font.size = Pt(11); mrun.font.color.rgb = DARK_GRAY

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
drun = date_p.add_run("Prepared: 2026-06-19  -  Workspace: DSCR_LOAN OFFICE\n"
                     "Status: Synthesis v2 — corrects v1 errors, adds Sprint 0-6, FCRA, Definitive Blueprint,\n"
                     "Definitive Product Spec, v16 fixes, Master Knowledge, 2026 Master Knowledge Paper")
drun.font.size = Pt(10); drun.font.color.rgb = DARK_GRAY

doc.add_paragraph()
warn_p = doc.add_paragraph()
warn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
wr = warn_p.add_run("ADVISOR-GRADE  -  RESEARCH-BACKED  -  AUDITABLE  -  SR 26-02 COMPLIANT\n"
                   "Every formula source-cited  -  Every assumption labeled [ASSUMPTION]\n"
                   "Deterministic math first, AI explanation second\n"
                   "v2.0 corrects 7 errors in v1.0 synthesis — see Section 1.4")
wr.font.size = Pt(11); wr.font.bold = True; wr.font.color.rgb = ACCENT

add_page_break(doc)

# ============== TOC ==============
add_h1(doc, "Table of Contents")
toc_items = [
    "PART I — STRATEGIC FOUNDATION",
    "  1. Executive Summary & v2 Corrections",
    "  2. Six-Function Doctrine & Three-Audience Framework",
    "  3. Three-Plane Architecture (Graph-Native OS)",
    "PART II — DETERMINISTIC FINANCIAL CORE",
    "  4. Dual-Track DSCR + Stabilized + All-In (4 tracks)",
    "  5. BUG/FLAW Catalog (v16 reconciled)",
    "  6. After-Tax Engine (OBBBA + §1250 + NIIT + PAL + REP)",
    "  7. ARM Reset Engine + IO Reversion Cliff",
    "PART III — STRESS & RISK",
    "  8. t-Copula Monte Carlo (Sprint 6 implementation)",
    "  9. Macro Archetypes + Sequential Drawdown + MCID",
    " 10. Cap-Rate Linked Refi Solver (Sprint 3)",
    " 11. STR Risk Scoring (Sprint 2/3)",
    "PART IV — LENDER & COMPLIANCE",
    " 12. Lender Footprint Matrix (Sprint 3 + Definitive Blueprint)",
    " 13. PPP State Matrix (Sprint 2 — all 50 states)",
    " 14. ECOA / FCRA / SR 26-02 Compliance Layer",
    " 15. Adverse-Action Notice Payload (FCRA PDF spec)",
    "PART V — LIVE DATA & ARCHITECTURE",
    " 16. Live Rate Triplet (FRED + CME SOFR, June 17-18 2026)",
    " 17. RentCast + AirDNA + Optimal Blue Integration",
    " 18. Three-Metric Credit Standard + Dual-Audience Architecture",
    "PART VI — IMPLEMENTATION ROADMAP",
    " 19. v1 → v2 Corrections Summary",
    " 20. Slice-by-Slice Build Plan (consolidated)",
    " 21. Specific Code-Level Action Items",
    " 22. SR 26-02 Compliance Status",
    "PART VII — APPENDICES",
    "  A. Document Inventory (60+ sources)",
    "  B. Pseudocode Library (canonical, expanded)",
    "  C. Live Rate Triplet + Market Data (June 17-18 2026)",
    "  D. Lender Footprint Matrix",
    "  E. References & Source Anchors",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)

add_page_break(doc)

# ============== SECTION 1: EXEC SUMMARY + v2 CORRECTIONS ==============
add_h1(doc, "Part I — Strategic Foundation")

add_h2(doc, "1. Executive Summary & v2 Corrections")

add_para(doc, 
    "Sixty-plus source documents were analyzed end-to-end to inform the upgrade of DSCR Sovereign OS into an "
    "Advisor-Grade Decision Engine. The corpus spans master specifications (Markdown, ~1.5 MB), architectural "
    "blueprint PDFs (~2.5 MB / ~750 pages), Sprint 0-6 research execution reports, the v14/v15/v16 master "
    "consolidated specs, the Definitive Master Blueprint v3, the Definitive Product Specification v12, and the "
    "Slice 1 source code itself (payment/dscr/leverage/ltv/compliance + audit + golden vectors).",
    bold=False)

add_h2(doc, "1.1 The Verdict (v2)")
add_para(doc, 
    "The corpus converges on a remarkably consistent architectural vision but diverges significantly on "
    "specific numerical parameters and feature names. The flagship value is the Qualifies-but-Dangerous "
    "(QbD) detector. The Six-Function Doctrine + Three-Audience Framework + Three-Plane Graph-Native "
    "Architecture are the canonical strategic foundation. The Dual-Track DSCR + After-Tax Engine + "
    "Monte Carlo + Live Data + Audit Trail are the canonical technical foundation.",
    bold=True)

add_h2(doc, "1.2 Headline Architecture (SR 26-02 Critical Insight)")
add_callout(doc, 
    "SR 26-02 (OCC Bulletin 2026-13, effective April 17, 2026) EXPLICITLY EXCLUDES simple arithmetic calculations "
    "and deterministic rule-based processes from model governance. The DSCR calculator + Legal Rules Engine + "
    "After-Tax engine are NOT models. Only the Monte Carlo engine, ML forecasters, and approval predictor "
    "fall under model governance. This eliminates validation overhead on the most-used layer — a deliberate moat.",
    color=GREEN, label="SR 26-02 INSIGHT")

add_table(doc,
    ["Component", "SR 26-02 Classification", "Governance"],
    [
        ["DSCR calculator (QuantLib/pyxirr)", "Not a model", "Unit tests + CI/CD"],
        ["Legal Rules Engine", "Not a model", "Quarterly counsel review"],
        ["After-Tax Engine", "Not a model", "Unit tests + IRS source verification"],
        ["Monte Carlo Risk Engine", "High-materiality model", "Full model card + challenger"],
        ["TimesFM/TFT Forecasters", "Medium-high model", "Model card + backtesting"],
        ["Approval Predictor (XGBoost)", "High-materiality model", "Full card + outcomes analysis"],
    ],
    widths=[2.5, 1.8, 2.0])

add_h2(doc, "1.3 The Headline Numbers (v2)")
add_table(doc,
    ["Metric", "Baseline", "Target After v2 Improvements"],
    [
        ["DSCR measurement", "Single static ratio (Track 1 only)", "4 tracks: Track 1 Lender + Track 2 Investor + Track 3 Stabilized + All-In"],
        ["Time horizon", "Year-1 snapshot", "10-year trajectory + 36mo sequential drawdown + IO+ARM reset"],
        ["Stress scenarios", "3 independent shocks", "4+ correlated macro archetypes (Stagflation/Recession/Climate/Local)"],
        ["Monte Carlo", "Absent", "t-copula, 10K trials, KBRA-calibrated 5-factor correlation"],
        ["Breakpoint solver", "Single-variable algebra", "Multi-variable constrained opt + Matrix grid + Counterfactual"],
        ["ECOA compliance", "Abstract scores", "40 reason codes (Form C-1 verbatim) + FCRA disclosure + state notices"],
        ["Refi realism", "Break-even rate only", "Cap-rate-linked dual gate (DSCR + LTV) with Cap_Rate_Beta"],
        ["Audit trail", "Log file", "SHA-256 + Merkle + SR 26-02 model cards + version-pinned JSON"],
        ["After-tax", "Pre-tax only", "OBBBA 100% bonus + §1250 recapture + NIIT + PAL + 1031"],
        ["SR 26-02 status", "Not classified", "Deterministic core excluded from model governance"],
    ],
    widths=[1.7, 2.0, 3.6])

add_h2(doc, "1.4 v1 → v2 Corrections — Important")
add_callout(doc, 
    "v1.0 synthesis contained 7 errors corrected in v2.0. Do NOT act on v1.0 numbers — they are stale or wrong.",
    color=CRIT, label="CRITICAL CORRECTIONS")

add_table(doc,
    ["#", "v1.0 Statement", "v2.0 Correction", "Source"],
    [
        ["1", "ECOA codes 19/21/26/27/28 should be re-numbered to match Form C-1 verbatim", "DO NOT re-number — codes are CORRECT. v1.0 was wrong.", "FCRA PDF p.6"],
        ["2", "PPP PA threshold = $319,777", "$319,777 was 2025; 2026 figure = **$329,411** (LIPL-indexed)", "Sprint 2 + Definitive Blueprint v3"],
        ["3", "LLC-vested non-bank financing triggers FinCEN BOI reporting", "**WRONG** under current law — non-bank lenders exempt per FinCEN Mar 2025 interim final rule", "Definitive Blueprint v3 §FinCEN"],
        ["4", "Architecture: deterministic core is a model under SR 11-7", "**NOT a model** under SR 26-02 (April 2026); simpler compliance", "Definitive Blueprint v3 §SR 26-02"],
        ["5", "TimesFM 2.0 with 500M params", "**TimesFM 2.5**: 200M params (faster), 15,360 context (7.5x), native quantile head, XReg covariates", "Google BigQuery docs Jun 12 2026"],
        ["6", "Three DSCR tracks (Track 1/2 only)", "**Four tracks**: Track 1 Lender + Track 2 Investor + Track 3 Stabilized + All-In", "v16 Master Consolidated"],
        ["7", "v1.0 said 'compress 36mo simulation to scalar LSC'", "v2.0 keeps LSC as quick diagnostic but uses **Sequential Drawdown Array** for gating", "Sprint 6 + multiple docs"],
    ],
    widths=[0.3, 3.0, 3.2, 1.5])

add_h2(doc, "1.5 The Six Big Findings (v2)")
add_para(doc, "Finding 1: SR 26-02 architectural split is the biggest moat.", bold=True)
add_para(doc, 
    "Deterministic math (DSCR calc, legal rules engine, after-tax engine) is explicitly excluded from model "
    "governance. Only Monte Carlo + ML need full model cards. This means the engine can ship the deterministic "
    "layer 5x faster than competitors still operating under SR 11-7's blanket definition.")

add_para(doc, "Finding 2: Dual-audience architecture (Borrower + Loan Officer) is the product frame.", bold=True)
add_para(doc, 
    "Borrower wants truth (DSCR survives? deal makes money?). Loan Officer wants speed (10-min verdict, "
    "zero-defect file). The same engine must serve both simultaneously via the Three-Plane architecture "
    "(Projection/Graph/Ledger). The 10-Minute Committee-Grade Verdict + 150-word Investment Thesis Block + "
    "Adverse-Case Recourse Table are the LO-facing outputs.")

add_para(doc, "Finding 3: The Three-Metric Credit Standard replaces single DSCR for credit decisions.", bold=True)
add_para(doc, "DSCR (Cash Control) + Debt Yield ≥ 9% (Workout Metric) + LTV (Loss-Given-Default). A deal "
    "passing DSCR but failing Debt Yield is fragile; one with low LTV but bad DSCR is liquidity-stressed. "
    "All three are required outputs.")

add_para(doc, "Finding 4: After-Tax IRR is a deal-flopper. OBBBA changes everything.", bold=True)
add_para(doc, 
    "OBBBA 100% bonus depreciation (permanent, acquired after Jan 19, 2025) + §1250 recapture 25% on residential "
    "(Bucket 1 = 0; Bucket 2 = 25%; Bucket 3 = LTCG) + NIIT 3.8% on rental income for MFJ MAGI > $250K + PAL "
    "phase-out $150K complete + 1031 exchange deferral. A negative-carry deal can be a winner after depreciation "
    "shelters income. The engine MUST compute after-tax IRR alongside pre-tax.")

add_para(doc, "Finding 5: S&P institutional stress = DSCR × 1.5-2.5x.", bold=True)
add_para(doc, 
    "S&P Global Ratings applies a DSCR adjustment factor of 1.50x-2.50x when rating non-QM/DSCR pools (NRMLT "
    "2026-NQM1 methodology). A borrower's 1.20 DSCR looks like 0.48-0.80 under S&P's stressed view. The engine's "
    "Monte Carlo must reproduce this range to validate institutional alignment.")

add_para(doc, "Finding 6: t-Copula (NOT Gaussian) for correlated RE risk.", bold=True)
add_para(doc, 
    "Gaussian copula = 2008 CDO failure mechanism. t-Copula with ν=5-7 df captures tail dependence (under stress, "
    "rent declines, vacancy spikes, cap rate expansion co-move). KBRA-calibrated 5-factor correlation matrix: "
    "rent↔vac -0.55, rent↔exp 0.25, rent↔cap 0.35, vac↔cap -0.30, cap↔rate 0.20. **Gaussian copula is BANNED from "
    "production use**.")

add_page_break(doc)

# ============== SECTION 2: SIX-FUNCTION DOCTRINE ==============
add_h2(doc, "2. Six-Function Doctrine & Three-Audience Framework")

add_h3(doc, "2.1 The Six Functions (Godmode v7 — from THE COMPLETE SOVEREIGN MASTER DOCUMENT)")
add_para(doc, "Every feature, every code module, every operational decision must trace back to exactly one of these six functions. The Iron Rule prevents the platform from becoming a feature graveyard.")

add_table(doc,
    ["#", "Function", "Elite Standard", "Module"],
    [
        ["01", "Scenario Accuracy", "GO/NO-GO verdict with confidence score in under 10 minutes", "engine.ts + preflightGate.ts + rentCompAggregator.ts"],
        ["02", "Guideline Intelligence", "25+ verified lenders with auto-fit scoring and two-quote rule", "lenders.ts + lenderGuidelines.ts + fitScorer.ts"],
        ["03", "Borrower Trust", "Every quote regulator-ready, backed by full constraint disclosure", "quoteExplainer.ts + pdfQuotePack.ts"],
        ["04", "Capital Partner Trust", "Zero-defect file standard, first-pass clean rate above 90%", "fileCompletenessEngine.ts + defectScorer.ts"],
        ["05", "Distribution", "60%+ of revenue from repeat referral channels", "referralPortal.ts + channelAttribution.ts"],
        ["06", "Risk Discipline", "Hard decline gates + adverse-action compliance, false-decline below 5%", "declineGate.ts + adverseActionEngine.ts"],
    ],
    widths=[0.4, 1.5, 2.5, 2.5])

add_h3(doc, "2.2 Three Audiences of Every Quote")
add_para(doc, "Every DSCR quote is read by three audiences simultaneously:")
add_bullet(doc, "BORROWER — cares whether deal closes and at what cost; judges by rate, fees, fairness")
add_bullet(doc, "CAPITAL PARTNER (lender UW, investor asset manager, credit committee) — cares whether file is clean and defensible; judges by defect rate, audit trail")
add_bullet(doc, "OPERATOR (loan officer) — cares whether 10 minutes produced a verdict that holds through closing")
add_para(doc, "A quote that satisfies only one audience is a failure.", italic=True)

add_page_break(doc)

# ============== SECTION 3: THREE-PLANE ARCHITECTURE ==============
add_h2(doc, "3. Three-Plane Architecture (Graph-Native OS)")

add_para(doc, "From DEFINITIVE PRODUCT SPEC v12 + THE COMPLETE SOVEREIGN MASTER DOCUMENT. The OS is not a flat database app — it's a Graph-Native Financial OS.")

add_table(doc,
    ["Plane", "Definition", "Implementation"],
    [
        ["Projection Plane", "Human-facing views", "Scenario Builder, Lender Matchmaker, After-Tax IRR Studio, IC Memo Command"],
        ["Graph Plane", "Causal central nervous system", "Nodes (Borrower, Property, Lender, Law, Rate) with Typed Edges (Qualifies, Conflicts, Supersedes, Shocks)"],
        ["Ledger Plane", "Immutable append-only event log", "Every mutation, approval, export captured with full provenance"],
    ],
    widths=[1.5, 2.0, 3.5])

add_h3(doc, "3.1 Semantic Diff Engine")
add_para(doc, "Classifies changes by facet (Location, Timing, Budget, Legal). A structural change (e.g., LLC to Individual vesting) triggers causal propagation through the PPP Legal Branching Gate without destroying unrelated underwriting work. A cosmetic change (typo fix) produces no propagation.")

add_h3(doc, "3.2 Evidence Vault (JSONB in PostgreSQL)")
add_table(doc,
    ["Field", "Type", "Description"],
    [
        ["source_id", "string", "Vendor identifier (rentcast, airdna, fred, ocrolus, etc.)"],
        ["as_of_timestamp", "datetime (UTC)", "When the data was retrieved from source"],
        ["effective_date", "datetime", "What date the data describes"],
        ["confidence_score", "float [0–1]", "Source-specific reliability rating"],
        ["hash", "string", "SHA-256 of raw response payload"],
        ["ttl_hours", "int", "Time-to-live before staleness flag"],
        ["provenance_tier", "enum", "primary_source / vendor_model / derived / user_input"],
        ["decay_rate", "float", "Confidence reduction per hour after TTL"],
    ],
    widths=[1.4, 1.4, 4.0])

add_para(doc, "Staleness is active: evidence older than its TTL decays in confidence automatically. Stale data self-flags. No manual discipline required.")

add_page_break(doc)

# ============== SECTION 4: DUAL-TRACK + TRACK 3 + ALL-IN ==============
add_h1(doc, "Part II — Deterministic Financial Core")

add_h2(doc, "4. Four DSCR Tracks (v16 Master Consolidated)")

add_callout(doc, 
    "v1.0 said dual-track. v2.0 (v16) establishes FOUR tracks. The engine must implement all four.",
    color=WARN, label="EXPANDED FROM v1.0")

add_table(doc,
    ["Track", "Numerator", "Denominator", "Frequency", "Use Case", "Rule"],
    [
        ["Track 1 — Lender Qualifying DSCR", "Qualifying rent: lower of lease or market rent (Form 1007)", "PITIA", "Monthly", "DSCR loan qualification", "Qualifying_Rent_Monthly / PITIA_Monthly"],
        ["Track 2 — Investor Survival DSCR", "NOI = EGI − OpEx (excl. debt service)", "ADS = annual debt service", "Annual", "True property-level debt coverage", "NOI_Annual / ADS_Annual"],
        ["Track 3 — Stabilized DSCR", "Stabilized Year-N NOI (usually Year 3 after rehab/lease-up)", "ADS after recast or perm debt", "Annual", "Value-add, bridge-to-perm, repositioning", "Stabilized_NOI / Stabilized_ADS"],
        ["All-In DSCR", "NOI", "PI + T + I + HOA annualized", "Annual", "Conservative lender/investor variant", "NOI / All_In_Housing_Cost"],
    ],
    widths=[1.3, 1.4, 1.4, 0.7, 1.2, 1.0])

add_para(doc, "Engine Enforcement: no cross-track comparison is allowed without explicit labeling. A deal can pass Track 1 lender qualification while failing Track 2 investor survival.")

add_h2(doc, "4.1 Golden Vector (PIN as Unit Tests — matches Slice 1)")
add_code_block(doc,
    "Payment Factor Formula: factor(r) = r(1+r)^360 / ((1+r)^360 - 1), r = annual_rate / 12\n"
    "Verified factors: 6.125% -> 0.0060761 | 7.00% -> 0.0066530 | 8.25% -> 0.0075127\n"
    "IO: Monthly_IO = Loan * rate / 12\n"
    "\n"
    "Reference Deal ($425K / 75% LTV / 7.00% / lease $3,000 = 1007 / tax $5K / ins $2K / HOA $150):\n"
    "  P&I = $318,750 * 0.0066530 = $2,121\n"
    "  PITIA = $2,121 + $416.67 + $166.67 + $12.50 = $2,855 monthly\n"
    "  Track 1 DSCR @ 7.00% = $3,000 / $2,855 = 1.05\n"
    "  Track 1 DSCR @ 8.25% = $3,000 / $3,192 = 0.96\n"
    "  Track 2 DSCR (8% vac, 8% mgmt) = 0.88 -> negative $335/mo\n"
    "  Rent break-even (T1=1.0) = $2,855 (-4.83%)\n"
    "  Deal-break rate ~= 7.67%\n"
    "  Max price at T1=1.0 ~= $454,100")

add_h2(doc, "4.2 Returns Engine (Pre-Tax)")
add_code_block(doc,
    "Accounting Split (define once, never mix):\n"
    "  EGI = GPR * (1 - Vacancy)\n"
    "  OpEx = Mgmt + Maint + Tax + Ins + HOA + Util + Turnover  [NO debt, NO capex]\n"
    "  NOI = EGI - OpEx\n"
    "  ADS = P&I * 12\n"
    "  CapEx reserve: modeled separately at 5-8% EGI\n"
    "  PITIA is the LENDER denominator; NOI is the INVESTOR result.")

add_table(doc,
    ["Metric", "Formula"],
    [
        ["Cap Rate", "NOI / Price"],
        ["Yield-on-Cost", "Stabilized_NOI / Total_Cost"],
        ["CoC", "(NOI - ADS) / Cash_Invested [Year 1, 3, 5]"],
        ["Debt Yield", "NOI / Loan [target >= 9% institutional]"],
        ["Break-even Occupancy", "(OpEx + ADS) / GPR"],
        ["Equity Multiple", "Total_Distributions / Total_Equity_Invested"],
        ["DSCR Cushion", "Track1 - Lender_Floor"],
    ],
    widths=[1.8, 4.0])

add_h2(doc, "4.3 Levered IRR + Exit Model (Sensitivity Grid)")
add_code_block(doc,
    "m0: -Cash_Invested\n"
    "m1..n: (NOI/12 - P&I)\n"
    "mn: + [Exit_NOI/Exit_Cap - Selling_Costs - Remaining_Balance - Prepay(exit_year)]\n"
    "Sensitivity grid: 4 hold periods (3/5/7/10yr) x 3 exit cap scenarios (bear/base/bull) x 4 rent growth rates (0/1/2/3%) = 48-cell matrix\n"
    "Flag: IRR's sensitivity to EXIT CAP is the most fragile input in every model.")

add_page_break(doc)

# ============== SECTION 5: BUG/FLAW CATALOG ==============
add_h2(doc, "5. BUG / FLAW Catalog (v16 Master Consolidated — fully reconciled)")

add_callout(doc, 
    "Slice 1 already fixed BUG-01/05/06. v16 Master Consolidated adds BUG-02/03 + FLAW-01/02 + IMP-08. These must be added in Slice 2.",
    color=WARN, label="ACTION REQUIRED")

add_h3(doc, "5.1 Critical Bugs")
add_table(doc,
    ["ID", "Bug", "Fix", "Slice 1 Status"],
    [
        ["BUG-01", "LTV denominator must use min(purchase, appraisal) for purchases", "value_for_ltv() per transaction type", "FIXED in ltv.py"],
        ["BUG-02", "NOI growth exponent off-by-one (Year3 = Year1 * (1+g)^2 not ^3)", "year1_noi * (1+g)^(year-1)", "NEW — Slice 2"],
        ["BUG-03", "Vacancy tornado labels swapped (low vac = best case, high vac = worst)", "vac_best = max(0, vac - 0.05); vac_worst = min(1, vac + 0.05)", "NEW — Slice 2"],
        ["BUG-05", "Breakeven Occupancy must include OpEx", "(ADS + OpEx) / Potential_Gross_Income", "FIXED in ltv.py"],
        ["BUG-06", "IO Max Loan must use decimal rate (not percent)", "Max_Loan_IO = (Max_PI * 12) / annual_rate_decimal", "FIXED in ltv.py"],
    ],
    widths=[0.6, 2.6, 2.5, 1.2])

add_h3(doc, "5.2 Logic Flaws")
add_table(doc,
    ["ID", "Flaw", "Fix"],
    [
        ["FLAW-01", "Required DSCR risk factors don't stack", "Additive model: product_minimum + sum of risk_adjustments (capped at 1.30)"],
        ["FLAW-02", "Waterfall promote is cliff model not tranched", "Sequential tranche waterfall with LP IRR hurdle"],
        ["IMP-08", "Stressed DS re-amortization not handled", "Recalc PMT over shortened remaining term after IO expiry + rate reset"],
    ],
    widths=[0.6, 2.6, 3.5])

add_h3(doc, "5.3 Dimensional Types (Required)")
add_code_block(doc,
    "class Rate:\n"
    "    annual: number  # decimal, e.g. 0.07\n"
    "    @property\n"
    "    def monthly(self): return self.annual / 12\n"
    "    @property\n"
    "    def display(self): return f'{(self.annual*100).toFixed(3)}%'\n"
    "FORBIDDEN: raw rate%, ambiguous 7 vs 0.07, mixing monthly numerator with annual denominator")

add_page_break(doc)

# ============== SECTION 6: AFTER-TAX ENGINE ==============
add_h2(doc, "6. After-Tax Engine (OBBBA + §1250 + NIIT + PAL + REP)")

add_h3(doc, "6.1 OBBBA 100% Bonus Depreciation (Sprint 4 verified)")
add_table(doc,
    ["Test", "Rule"],
    [
        ["TEST 1 — Acquisition Date", "Property must be acquired after January 19, 2025"],
        ["TEST 2 — Binding Contract", "If binding written contract executed before Jan 20, 2025: NOT ELIGIBLE for 100% bonus"],
        ["TEST 3 — 10% Safe Harbor", "If more than 10% of total hard costs incurred before Jan 20, 2025: NOT ELIGIBLE"],
    ],
    widths=[1.8, 4.5])

add_para(doc, "Source: Treasury/IRS Notice 2026-11 (issued Jan 13, 2026).")
add_para(doc, "ELIGIBLE: 5-yr property (appliances, carpet, furniture), 7-yr (office furniture, fixtures), 15-yr (land improvements, fencing, paving), used property if not previously owned by this taxpayer, cost-seg components.")
add_para(doc, "NOT ELIGIBLE: building structure (always 27.5-yr straight-line), land, residential rental building.")

add_h3(doc, "6.2 Cost Segregation Reclassification (Sprint 4)")
add_code_block(doc,
    "COST_SEGREGATION_RECLASSIFICATION = {\n"
    "    'total_building_basis': 1.0,           # 100% basis\n"
    "    'structure_27_5_yr': 0.60,             # 60%: non-reclassifiable (straight-line)\n"
    "    'land_improvements_15_yr': 0.15,       # 15%: eligible for 100% bonus\n"
    "    'personal_property_5_yr': 0.15,        # 15%: eligible for 100% bonus\n"
    "    'personal_property_7_yr': 0.10,        # 10%: eligible for 100% bonus\n"
    "    'reclassifiable_total': 0.40,\n"
    "    'study_cost_range': (5000, 15000),\n"
    "    'minimum_basis_for_study': 500000,\n"
    "}\n"
    "# $750K acquisition example: $253,091 Year-1 depreciation vs $21,818 standard = $231,273 advantage\n"
    "# At 37% marginal rate = $85,571 deferred tax; net Year-1 benefit ~$79,571")

add_h3(doc, "6.3 §1250 Recapture — Three Buckets")
add_table(doc,
    ["Bucket", "Definition", "Tax Rate", "Applies to"],
    [
        ["Bucket 1 — Recaptured §1250", "Depreciation in EXCESS of straight-line", "ORDINARY (up to 37%)", "Commercial (39-yr) where accelerated taken; ZERO for residential 27.5-yr"],
        ["Bucket 2 — Unrecaptured §1250", "Total straight-line depreciation taken", "Max 25% federal", "Every residential rental seller (Form 4797 Part III)"],
        ["Bucket 3 — Remaining Gain", "Sale price - adjusted basis - Bucket 2", "LTCG 0/15/20%", "True appreciation"],
    ],
    widths=[1.5, 1.7, 1.5, 1.6])

add_h3(doc, "6.4 NIIT — Fixed (NOT CPI-Adjusted) Since 2013")
add_table(doc,
    ["Filing Status", "MAGI Threshold", "NIIT Rate"],
    [
        ["Married Filing Jointly", "$250,000", "3.8%"],
        ["Single / Head of Household", "$200,000", "3.8%"],
        ["Married Filing Separately", "$125,000", "3.8%"],
        ["Qualifying Widow(er)", "$250,000", "3.8%"],
    ],
    widths=[2.0, 1.8, 1.0])

add_callout(doc, 
    "NIIT thresholds are FIXED by statute (IRC §1411). More investors hit threshold every year from "
    "inflation bracket creep. Engine must surface as compounding risk for near-threshold investors.",
    color=WARN, label="CRITICAL TAX INSIGHT")

add_h3(doc, "6.5 PAL — Phase-Out at $150K (Canonical Resolution)")
add_callout(doc, 
    "v1.0 had Uncle Kam CPA's $200K MFJ figure. Sprint 4 statutory analysis confirms $150K for ALL "
    "individual filers (Single/HH/MFJ). The $100K-$150K applies to Single; Uncle Kam's $200K reflects "
    "NIIT threshold not PAL.",
    color=WARN, label="v1.0 CORRECTION")

add_code_block(doc,
    "PAL_ENGINE = {\n"
    "    'standard_allowance': 25000,\n"
    "    'phase_out_threshold': 100000,\n"
    "    'phase_out_rate': 0.50,\n"
    "    'phase_out_complete_magi': 150000,  # NOT 200K\n"
    "    'rep_hours_minimum': 750,\n"
    "    'rep_time_pct_minimum': 0.50,\n"
    "}")

add_h3(doc, "6.6 REP Status — Highest-Leverage Tax Planning")
add_para(doc, "Real Estate Professional under IRC §469(c)(7):")
add_bullet(doc, "More than 750 hours in real property trades/businesses AND")
add_bullet(doc, "More than 50% of personal services in real property")
add_bullet(doc, "Result: ALL rental losses deductible against W-2 + active business income", bold=True)
add_bullet(doc, "NIIT also eliminated on rental income", bold=True)
add_para(doc, "REP status is the highest-leverage tax status the engine can advise.", italic=True)

add_h3(doc, "6.7 After-Tax IRR Computation (Full Pipeline)")
add_code_block(doc,
    "def compute_after_tax_irr(deal, tax_profile, hold_years):\n"
    "    # Phase 1: Cash flows\n"
    "    annual_pre_tax_cf = deal.monthly_noi * 12 - deal.monthly_pitia * 12\n"
    "    # Phase 2: Depreciation (OBBBA + cost-seg)\n"
    "    building_basis = deal.purchase_price - deal.land_value\n"
    "    if deal.cost_seg_elected and deal.purchase_price >= 500000:\n"
    "        bonus_eligible = building_basis * 0.40\n"
    "        structure = building_basis * 0.60\n"
    "        yr1_dep = bonus_eligible + (structure / 27.5)\n"
    "        ongoing_dep = structure / 27.5\n"
    "    # Phase 3: Year-by-year tax computation\n"
    "    # Phase 4: NIIT, PAL, REP status, §1250 recapture at exit\n"
    "    # Phase 5: XIRR via pyxirr")

add_page_break(doc)

# ============== SECTION 7: ARM RESET ENGINE ==============
add_h2(doc, "7. ARM Reset Engine + IO Reversion Cliff")

add_h3(doc, "7.1 SOFR Anchors (June 17, 2026 — Sprint 5 verified)")
add_table(doc,
    ["Series", "FRED ID", "Value", "Source"],
    [
        ["10-Year Treasury (DGS10)", "DGS10", "4.43%", "FRED"],
        ["SOFR Overnight", "SOFR", "3.63%", "FRED"],
        ["30-Day Average SOFR", "NY Fed", "3.609%", "NY Fed"],
        ["90-Day Average SOFR", "NY Fed", "3.636%", "NY Fed"],
        ["1-Month Term SOFR", "CME", "3.637%", "CME Term SOFR"],
        ["3-Month Term SOFR", "CME", "3.668%", "CME Term SOFR"],
        ["6-Month Term SOFR", "CME", "3.731%", "CME Term SOFR"],
        ["12-Month Term SOFR", "CME", "3.869%", "CME Term SOFR"],
    ],
    widths=[2.0, 0.8, 0.8, 1.0])

add_callout(doc, 
    "CME Term SOFR requires Category One Use License for commercial use ($ enterprise pricing). "
    "FREE alternative: NY Fed SOFR Averages + FRED historical. Most DSCR ARMs use 30-day SOFR avg, not CME Term.",
    color=WARN, label="LICENSING CONSTRAINT")

add_h3(doc, "7.2 ARM Reset Formula (Canonical)")
add_code_block(doc,
    "Fully_Indexed_Rate = Index_t + Margin\n"
    "New_Rate = min(max(SOFR_t + Margin, Floor),\n"
    "                min(Current_Rate + Periodic_Cap, Initial_Rate + Lifetime_Cap))\n"
    "Reset_Payment = Remaining_Balance * New_Rate/12 / (1 - (1 + New_Rate/12)^(-n_remaining))")

add_h3(doc, "7.3 IO + ARM Double-Shock — Kill-Switch Year")
add_para(doc, "When IO expires, loan recasts to amortizing simultaneously with potential rate reset. The double-shock year must be flagged as a kill-criterion checkpoint.")
add_para(doc, "Engine output: \"Kill-Switch Year: Year N — IO expires and rate resets simultaneously.\"", italic=True)

add_h3(doc, "7.4 ARM Worked Example (Sprint 5)")
add_para(doc, "5/6 ARM resetting at month 60, start rate 7.25%:")
add_bullet(doc, "Index = 6-Month Term SOFR at reset (currently 3.731%)")
add_bullet(doc, "Margin = 2.50% (typical DSCR)")
add_bullet(doc, "Computed new rate = 3.731% + 2.50% = 6.231%", bold=True)
add_bullet(doc, "Periodic cap +2%, Lifetime cap +5%")
add_callout(doc, 
    "With current inverted-to-flat curve, ARM reset rate (6.231%) is LOWER than start rate (7.25%). "
    "Engine must compute explicitly rather than assuming rates move adversely.",
    color=GREEN, label="COUNTERINTUITIVE FINDING")

add_page_break(doc)

# ============== SECTION 8: T-COPULA MONTE CARLO ==============
add_h1(doc, "Part III — Stress & Risk")

add_h2(doc, "8. t-Copula Monte Carlo (Sprint 6 — full implementation)")

add_h3(doc, "8.1 Why t-Copula, NOT Gaussian")
add_callout(doc, 
    "Gaussian copula = 2008 CDO failure mechanism. t-Copula with ν=5-7 df captures tail dependence "
    "that Gaussian misses. Under stress, rent declines, vacancy spikes, cap rate expansion co-move. "
    "Gaussian copula is BANNED from production use.",
    color=CRIT, label="MANDATORY")

add_h3(doc, "8.2 KBRA-Calibrated Correlation Matrix (5 factors)")
add_table(doc,
    ["Factor Pair", "Correlation", "Rationale"],
    [
        ["Rent ↔ Vacancy", "-0.55", "Inverse: high rent → low vacancy"],
        ["Rent ↔ Expense", "+0.25", "Both rise with market"],
        ["Rent ↔ Cap Rate", "+0.35", "Cap rate expands when rents pressured"],
        ["Vacancy ↔ Cap", "-0.30", "Vacancy up, values down"],
        ["Cap ↔ Rate", "+0.20", "Rate hikes → cap expansion"],
        ["Expense ↔ Cap", "+0.10", "Modest co-movement"],
    ],
    widths=[2.0, 1.0, 3.0])

add_h3(doc, "8.3 Marginal Distributions (Sprint 6 calibrated)")
add_table(doc,
    ["Factor", "Distribution", "Parameters", "Source"],
    [
        ["LTR Rental Growth", "Normal", "μ=2%, σ=5%", "KBRA / MMCG"],
        ["STR Gross Revenue", "Lognormal", "μ=0%, σ=18-25%", "AirDNA"],
        ["LTR Vacancy", "Beta(α=2, β=36)", "mean ≈ 5.3%", "CoStar/Trepp"],
        ["STR Vacancy", "Beta(α=3, β=7)", "mean 20-40%", "AirDNA"],
        ["Insurance Escalation", "Lognormal", "μ=7%, σ=5%; coastal μ=12%", "Post-2024 crisis"],
        ["Property Tax Growth", "Truncated Normal", "μ=3%, σ=1%; CA cap 2%", "Prop 13; TX/FL uncapped"],
        ["10Y Treasury Path", "CIR / Hull-White", "Calibrated to live SOFR", "FRED + QuantLib"],
        ["Rate Drift", "Normal", "μ=0%, σ=0.5%", "Current flat curve"],
    ],
    widths=[1.5, 1.5, 1.8, 1.2])

add_h3(doc, "8.4 Verdict Bands (Sprint 6)")
add_table(doc,
    ["Verdict", "P5 DSCR", "P(DSCR<1.0)"],
    [
        ["RESILIENT", ">= 1.0", "< 5%"],
        ["MODERATE_RISK", ">= 0.90", "< 15%"],
        ["ELEVATED_RISK", ">= 0.75", "< 35%"],
        ["STRESSED — REVIEW", "< 0.75", ">= 35%"],
    ],
    widths=[2.0, 1.0, 1.5])

add_callout(doc, 
    "S&P Global Ratings applies DSCR adjustment factor 1.50x-2.50x when rating non-QM/DSCR pools. "
    "A borrower's 1.20 DSCR looks like 0.48-0.80 under S&P's stressed view. Engine Monte Carlo must "
    "reproduce this range to validate institutional alignment.",
    color=ACCENT, label="INSTITUTIONAL VALIDATION")

add_h3(doc, "8.5 Verdict Triage + Hard NO-GO Rule")
add_para(doc, "P(DSCR<1.00) > 15% → HARD NO-GO regardless of lender Track 1 qualification.", italic=True, bold=True)

add_page_break(doc)

# ============== SECTION 9: STRESS ==============
add_h2(doc, "9. Macro Archetypes + Sequential Drawdown + MCID")

add_h3(doc, "9.1 Deterministic Macro Archetypes (4 — from Sprint 6)")
add_table(doc,
    ["Archetype", "Rent", "Vacancy", "OpEx", "Rate", "Cap Rate"],
    [
        ["Stagflation", "flat", "flat", "+10%", "+200 bps", "+50 bps"],
        ["Recession", "-15%", "+10%", "-5%", "-100 bps", "+75 bps"],
        ["Climate/Regional", "-5%", "-3%", "+50% insurance", "flat", "-10% value"],
        ["Local Distress", "-10%", "+5%", "flat", "flat", "+100 bps"],
    ],
    widths=[1.4, 0.9, 0.9, 1.3, 1.0, 1.0])

add_h3(doc, "9.2 Sequential Drawdown Array (Sprint 6)")
add_code_block(doc,
    "def sequential_drawdown(\n"
    "    monthly_rent: list[float],     # 36 months\n"
    "    monthly_opex: list[float],\n"
    "    monthly_debt_service: list[float],\n"
    "    capex_events: dict[int, float],\n"
    "    starting_reserves: float,\n"
    ") -> dict:\n"
    "    balance = starting_reserves\n"
    "    ruin_month = None\n"
    "    for t in range(36):\n"
    "        ncf = monthly_rent[t] - monthly_opex[t] - monthly_debt_service[t]\n"
    "        if t in capex_events: ncf -= capex_events[t]\n"
    "        balance += ncf\n"
    "        if balance < 0 and ruin_month is None: ruin_month = t + 1\n"
    "    return {'liquidity_failure': ruin_month is not None,\n"
    "            'ruin_month': ruin_month, 'min_balance': min(balance, 0)}")

add_h3(doc, "9.3 STR Seasonality — Maximum Cumulative Intra-Year Deficit (MCID)")
add_code_block(doc,
    "def compute_mcid(monthly_net_cashflow):\n"
    "    cumulative = 0; min_cumulative = 0\n"
    "    for cash in monthly_net_cashflow:\n"
    "        cumulative += cash\n"
    "        if cumulative < min_cumulative: min_cumulative = cumulative\n"
    "    return abs(min_cumulative)\n"
    "# Required reserves >= MCID, else fail seasonality trough test")

add_page_break(doc)

# ============== SECTION 10: CAP-RATE REFI ==============
add_h2(doc, "10. Cap-Rate Linked Refi Solver")

add_code_block(doc,
    "def break_even_refi_cap_linked(noi, loan_balance, current_rate,\n"
    "                                target_dscr, current_cap_rate,\n"
    "                                cap_rate_beta,  # 0.5-1.5 multifamily\n"
    "                                max_matrix_ltv, remaining_term_months):\n"
    "    def check(rate):\n"
    "        pmt = (loan_balance * (rate/12) * (1 + rate/12)**remaining_term_months) / \\\n"
    "              ((1 + rate/12)**remaining_term_months - 1)\n"
    "        dscr = noi / (pmt * 12)\n"
    "        projected_cap = current_cap_rate + cap_rate_beta * (rate - current_rate)\n"
    "        projected_value = noi / projected_cap\n"
    "        ltv = loan_balance / projected_value\n"
    "        return (dscr >= target_dscr) and (ltv <= max_matrix_ltv)\n"
    "    # Binary search...\n"
    "    return break_even_rate")

add_para(doc, "Cap_Rate_Beta default by asset class:")
add_bullet(doc, "Multifamily: 0.5-1.5 (UNC cap-rate study)")
add_bullet(doc, "Office: 0.8-2.0 (higher due to secular headwinds)")
add_bullet(doc, "Retail: 0.6-1.5")
add_bullet(doc, "Industrial: 0.3-0.8 (lowest due to secular tailwinds)")

add_page_break(doc)

# ============== SECTION 11: STR RISK ==============
add_h2(doc, "11. STR Risk Scoring (Sprint 2/3)")

add_h3(doc, "11.1 STR Legality (Gating)")
add_table(doc,
    ["Status", "Action"],
    [
        ["PROHIBITED (NYC, SF, LA, Santa Monica, Manhattan Beach, NO French Quarter)", "KILL — STR income disabled"],
        ["RESTRICTED (Nashville, Austin, Denver, Jersey City)", "Verify permit before underwriting"],
        ["PERMITTED with conditions (Miami, San Diego)", "Check zoning, permit availability"],
        ["PERMITTED (Houston, Dallas, Raleigh)", "No restrictions"],
    ],
    widths=[3.5, 3.0])

add_callout(doc, 
    "California SB 346 (effective Jan 1, 2026) — Airbnb/VRBO must share host data with local "
    "governments. Enforcement risk dramatically increased. Engine must flag CA STR deals.",
    color=WARN, label="CA SB 346 ALERT")

add_h3(doc, "11.2 STR Top 10 Markets (AirDNA 2026 — small/mid-tier focus)")
add_table(doc,
    ["Rank", "Market", "Cap Rate", "Median Price", "Annual Revenue"],
    [
        ["1", "Jackson, MS", "15.95%", "$84,672", "$24,550"],
        ["2", "Abilene, TX", "14.01%", "$201,493", "$51,330"],
        ["3", "Akron, OH", "11.66%", "$139,633", "$29,612"],
        ["4", "Montgomery, AL", "11.64%", "$143,500", "$30,364"],
        ["5", "Port Arthur, TX", "10.38%", "$124,353", "$23,477"],
        ["6", "Springfield, IL", "10.09%", "$159,667", "$29,283"],
        ["7", "Charleston, WV", "9.80%", "$158,399", "$28,211"],
        ["8", "Lebanon, PA", "8.68%", "$281,650", "$44,457"],
        ["9", "Lake Charles, LA", "8.41%", "$212,333", "$32,453"],
        ["10", "St. Paul, MN", "6.84%", "$289,137", "$35,968"],
    ],
    widths=[0.5, 1.5, 0.8, 1.0, 1.2])

add_para(doc, "Avg purchase $296K, avg yield 13.7%. Trophy markets (NYC/SF/LA) effectively dead for investors — engine must weight STR income toward these market types.")

add_h3(doc, "11.3 STR Income Qualification")
add_bullet(doc, "If multiple sources: LOWEST monthly income figure used")
add_bullet(doc, "Mandatory 20% haircut on gross AirDNA projection")
add_bullet(doc, "STR income capped at LTR market rent per Form 1007 (LTR floor)")
add_bullet(doc, "AirDNA: 12mo forecast, 3 comps, market score ≥ 60, dated within 90 days, ≤ 2 individuals/bedroom")

add_page_break(doc)

# ============== PART IV — LENDER & COMPLIANCE ==============
add_h1(doc, "Part IV — Lender & Compliance")

add_h2(doc, "12. Lender Footprint Matrix (Sprint 3 + Definitive Blueprint)")

add_table(doc,
    ["Lender", "NMLS", "States", "Key Specs"],
    [
        ["Visio Lending", "1935590", "41+DC", "DSCR floor 1.00; FICO 680; LTV 80%/75% CO; $100K-$5M; 5/4/3/2/1 or 3/2/1 PPP; #1 DSCR lender ($854.6M 2024)"],
        ["Kiavi", "—", "49+DC", "Tech-forward, AVM-heavy, rapid closings"],
        ["Angel Oak", "1160240", "47+DC", "Rental AVM (Clear Capital) Nov 2025 — locks at pre-qual; non-warrantable condo specialist"],
        ["Griffin Funding", "—", "46+DC", "Sub-1.0/low-DSCR/no-ratio/jumbo DSCR"],
        ["LendingOne", "—", "All+exempt", "Licensed or exempt in all other states"],
        ["Lima One Capital", "—", "National", "Premier business-purpose; $10B+ lifetime; DSCR/bridge/construction"],
        ["Deephaven", "—", "National", "DSCR First + DSCR Second Mortgage ($75K-$500K, no reserves, no income docs, AVM option)"],
        ["MortgageDepot", "—", "Verified", "40-yr amortized + 40-yr IO (up to $3M)"],
        ["Sistar Mortgage", "—", "Verified", "40-yr IO confirmed 2026"],
    ],
    widths=[1.2, 0.8, 0.9, 4.5])

add_h3(doc, "12.1 Deephaven DSCR Second — Use Cases")
add_bullet(doc, "Preserve Legacy Rate: keep 3.75% first mortgage, second lien at 8-10% beats cash-out refi to 7%+")
add_bullet(doc, "Bridge Down Payment: pull equity from Property A as down payment for Property B")
add_bullet(doc, "Renovation Capital: $150K without refi; no reserves; AVM at < $400K")

add_h3(doc, "12.2 Pool Benchmark Calibration (AOMT 2025-6 / NRMLT 2026-NQM1)")
add_table(doc,
    ["Pool Statistic", "AOMT 2025-6 (Angel Oak)", "NRMLT 2026-NQM1 (Rithm)"],
    [
        ["WA FICO", "746", "758"],
        ["WA CLTV", "71.95%", "Per filing"],
        ["WA DSCR (DSCR loans)", "1.19", "N/A"],
        ["Sub-1.0 DSCR concentration", "4.20%", "—"],
        ["DSCR loan % of pool", "42.43%", "—"],
        ["IO feature", "11.91%", "—"],
        ["Fixed vs ARM", "99.01% / 0.99%", "—"],
        ["Pool balance", "$349.65M", "$502.1M"],
    ],
    widths=[1.8, 2.3, 2.3])

add_callout(doc, 
    "When deal inputs EXCEED pool center (higher FICO, lower LTV, higher DSCR), engine gives "
    "confidence premium. When BELOW pool center, widen Monte Carlo confidence interval.",
    color=ACCENT, label="POOL-CENTER CALIBRATION")

add_h2(doc, "12.3 Competitive Threat Map")
add_table(doc,
    ["Competitor", "Threat", "Counter-Moat"],
    [
        ["YieldStack AI (Apr 2026)", "180+ lender programs, program-level matching, zero cost", "YieldStack matches; we ANALYZE (dual-track, MC, AEY, after-tax) + compliance gates + IC memo"],
        ["LenderSA 3.2 (Jan 2026)", "Hard money focus, AI negotiation, hundreds of lenders", "Different segment (fix-and-flip not DSCR); moat in analytical depth"],
        ["Angel Oak Rental AVM", "Industry-first AVM at pre-qual (locked through closing)", "Surface as preferred lender for speed-certainty purchases"],
    ],
    widths=[1.8, 2.2, 3.0])

add_page_break(doc)

# ============== SECTION 13: PPP STATE MATRIX ==============
add_h2(doc, "13. PPP State Matrix (Sprint 2 — 2026 verified)")

add_h3(doc, "13.1 Three-Branch Logic (BEFORE any state lookup)")
add_code_block(doc,
    "BRANCH 1: Business-purpose + Entity vesting (LLC, Corp, Trust)\n"
    "BRANCH 2: Business-purpose + Individual vesting\n"
    "BRANCH 3: Consumer-purpose (disqualified from DSCR by definition)\n"
    "# DSCR loans are always business-purpose (investment property)\n"
    "# Entity vesting frequently unlocks exemption from consumer PPP statutes")

add_h3(doc, "13.2 Critical States (2026 thresholds — CORRECTED FROM v1.0)")
add_table(doc,
    ["State", "PPP Status (2026)", "Threshold", "Statute"],
    [
        ["MN", "ALLOWED (Aug 1, 2026)", "N/A", "MN HF 3437 enacted Apr 23, 2026; §58.137 applies only to personal/family/household"],
        ["PA", "THRESHOLD-RESTRICTED", "$329,411 (2026, was $319,777 in 2025)", "PA Act 6, 10 Pa. Code §7.2 — CPI indexed"],
        ["OH", "THRESHOLD-RESTRICTED", "$112,957 (2025; 2026 needs Jan pull)", "OH ORC §1343.011 — CPI indexed"],
        ["WA", "ALLOWED w/ ARM restriction", "60 days before ARM reset", "RCW 19.144.040"],
        ["NJ", "ENTITY-DEPENDENT (3-branch)", "C-Corp ALLOWED, LLC CONTESTED, LP/Trust/Ind PROHIBITED", "N.J.S.A. 46:10B-2; Arc Home Jul 2025; NPLA Oct 2025"],
        ["CA", "ALLOWED (business-purpose)", "N/A", "CA Civil Code §2954.10"],
        ["TX, FL, GA, NC, TN, SC, VA, AL, IN, KY, MI, MO, WI, LA, AZ, CO", "ALLOWED", "N/A", "Business-purpose exemption"],
        ["NY", "ALLOWED (business-purpose LLC)", "Banking Law §6-l", "Verify usury not violated"],
        ["IL", "ALLOWED (LLC)", "Individual prohibited if rate >8%", "IL Residential Real Property Disclosure Act"],
        ["MS", "DECLINING STRUCTURES only", "Flat banned >1yr", "§75-17-31"],
        ["AR", "ALLOWED first 3yr", "Penalty base = REMAINING balance", "State PPP matrix"],
    ],
    widths=[0.6, 1.7, 2.0, 2.5])

add_callout(doc, 
    "OH and PA thresholds are CPI-indexed annually. Engine must fire Celery cron on Jan 1 each year to "
    "pull from OH Dept. of Commerce + PA Bulletin.",
    color=WARN, label="ANNUAL RE-INDEXING REQUIRED")

add_h3(doc, "13.3 PPP Penalty Mechanics (Canonical)")
add_code_block(doc,
    "# PENALTY BASE (per state/contract):\n"
    "# DEFAULT = REMAINING balance x step rate (market-standard DSCR)\n"
    "# STATUTORY OVERRIDE = ORIGINAL principal (OH; some others)\n"
    "STORE penalty_base as per-state, per-loan field. Apply binding rule.\n"
    "\n"
    "# PARTIAL PREPAY EXCEPTION: Most DSCR lenders allow 20% of original principal/year\n"
    "# without triggering PPP. Engine must surface.\n"
    "\n"
    "# STRUCTURES: 5/4/3/2/1, 3/2/1, flat 5/5/5, floored, six-months-interest, ~20%/yr partial-prepay")

add_h3(doc, "13.4 NJ Three-Branch Resolution (Engine Implementation)")
add_code_block(doc,
    "NJ_PPP_GATE = {\n"
    "    'C_Corp':   {'status': 'ALLOWED',    'confidence': 95},\n"
    "    'LLC':      {'status': 'CONTESTED',  'confidence': 30,\n"
    "                'action': 'FLAG_FOR_ATTORNEY_REVIEW',\n"
    "                'note': 'No published NJ case law; NPLA won partial DOBI clarification Oct 2025; Arc Home banned Jul 2025. Advise C-Corp vesting or no-PPP structure.'},\n"
    "    'LP':       {'status': 'PROHIBITED', 'confidence': 85},\n"
    "    'Trust':    {'status': 'PROHIBITED', 'confidence': 85},\n"
    "    'Individual':{'status': 'PROHIBITED','confidence': 95}\n"
    "}\n"
    "# SAFE HARBOR: Use C-Corp in NJ when PPP is required; offer LLC deals as no-PPP or buy-down")

add_page_break(doc)

# ============== SECTION 14: ECOA ==============
add_h2(doc, "14. ECOA / FCRA / SR 26-02 Compliance Layer")

add_h3(doc, "14.1 SR 26-02 Architectural Status (v2.0 — CORRECTED FROM v1.0)")
add_table(doc,
    ["Component", "SR 26-02 Status", "Governance"],
    [
        ["DSCR Calculator (QuantLib/pyxirr)", "NOT a model (explicitly excluded)", "Unit tests + CI/CD"],
        ["Legal Rules Engine (state PPP, usury)", "NOT a model", "Quarterly counsel review"],
        ["After-Tax Engine (OBBBA, §1250, NIIT, PAL)", "NOT a model", "Unit tests + IRS source verification"],
        ["Lender Qualification Engine", "NOT a model (rule-based)", "Config + version control"],
        ["Monte Carlo Risk Engine", "HIGH-materiality MODEL", "Full model card + champion/challenger"],
        ["TimesFM 2.5 / TFT Forecasters", "MEDIUM-HIGH model", "Model card + backtesting"],
        ["Approval Predictor (XGBoost)", "HIGH-materiality MODEL", "Full card + outcomes analysis + disparate impact monitor"],
        ["LLM narrative generation", "Outside SR 26-02 but needs internal governance", "Risk policy + audit trail"],
    ],
    widths=[2.5, 2.5, 2.0])

add_callout(doc, 
    "SR 26-02 (OCC Bulletin 2026-13, effective April 17, 2026) supersedes SR 11-7 (2011). It narrows the "
    "definition of model to complex quantitative methods. Simple arithmetic + deterministic rule-based "
    "processes are EXPLICITLY EXCLUDED. This is a deliberate moat — deterministic layers ship faster.",
    color=GREEN, label="SR 26-02 MOAT")

add_h3(doc, "14.2 ECOA Reason Codes (Form C-1 verbatim — v1.0 CORRECTION)")
add_callout(doc, 
    "v1.0 said re-number Slice 1 ECOA codes. v2.0 says: DO NOT. Slice 1 codes 19/21/26/27/28 ARE Form C-1 "
    "verbatim. Per FCRA PDF (p.6) the mapping is:",
    color=GREEN, label="v1.0 CORRECTION")

add_table(doc,
    ["Code", "Verbatim Text", "Trigger (FCRA PDF p.6)"],
    [
        ["19", "Your income is not sufficient to meet your expenses and debt payments.", "DSCR low (rent-driven) / FICO < 620 / Reserves < 3mo"],
        ["21", "Your debt payments or other obligations are too high.", "DSCR low (ADS-driven)"],
        ["26", "You requested an amount that exceeds the maximum loan amount permitted by our regulations.", "LTV 80-90%"],
        ["27", "The collateral value is insufficient.", "LTV > 90%"],
        ["28", "The type of property you selected is not acceptable to us.", "STR prohibited / property type rejection"],
    ],
    widths=[0.6, 3.2, 2.7])

add_h3(doc, "14.3 Nuanced Mapping (Dynamic — not static lookup)")
add_para(doc, "Per FCRA PDF: a static one-to-one mapping is INSUFFICIENT. Engine must evaluate kill event data + lender-specific config to select the most accurate reason code:")
add_code_block(doc,
    "# LTV-driven reason:\n"
    "if ltv > 0.90: code = 27  # collateral insufficient\n"
    "elif ltv > 0.80: code = 26  # exceeds max loan amount\n"
    "# DSCR-driven reason (nuanced by sub-cause):\n"
    "if dscr_e < 1.0:\n"
    "    if rent_below_market: code = 19  # income insufficient\n"
    "    elif ads_too_high: code = 21  # debt payments too high\n"
    "# Property type:\n"
    "if property_type_rejected: code = 28")

add_page_break(doc)

# ============== SECTION 15: ADVERSE-ACTION ==============
add_h2(doc, "15. Adverse-Action Notice Payload (FCRA PDF Spec)")

add_para(doc, "Per FCRA PDF p.8-9, the payload schema:")
add_code_block(doc,
    "{\n"
    "  'version': '1.0',\n"
    "  'as_of': 'YYYY-MM-DDTHH:MM:SSZ',\n"
    "  'lender_id': 'CLIENT_XXXX',\n"
    "  'application_id': 'APP_YYYYYY',\n"
    "  'is_compliant': true,\n"
    "  'regulatory_notices': {\n"
    "    'ecoa_notice': {\n"
    "      'header': 'ADVERSE ACTION NOTICE',\n"
    "      'prohibition_statement': '<verbatim Reg B text>',\n"
    "      'reasons': [\n"
    "        {'code': '27', 'text': 'The collateral value is insufficient.'},\n"
    "        {'code': '21', 'text': 'Your debt payments or other obligations are too high.'}\n"
    "      ]\n"
    "    },\n"
    "    'fcra_disclosure': {\n"
    "      'header': 'DISCLOSURE REQUIRED BY THE FAIR CREDIT REPORTING ACT',\n"
    "      'statement': '<verbatim FCRA text>',\n"
    "      'data_source': 'CoreLogic, Inc.',\n"
    "      'source_address': '123 Main St, Anytown, ST 12345'\n"
    "    }\n"
    "  },\n"
    "  'state_specific_notices': [\n"
    "    // {'state': 'CA', 'text': 'Additional CA-specific language...'}\n"
    "  ],\n"
    "  'meta': {\n"
    "    'generation_timestamp': 'YYYY-MM-DDTHH:MM:SSZ',\n"
    "    'engine_version': 'DSCR_Engine_v15.0.0',\n"
    "    'explanation_layer_version': '1.0.0'\n"
    "  }\n"
    "}")

add_para(doc, "Configuration via YAML/JSON (per FCRA PDF — Slice 1 compliance.py already supports this):")
add_bullet(doc, "Default federally-compliant config shipped with engine")
add_bullet(doc, "Lender clients customize via YAML overlay")
add_bullet(doc, "California-specific state_notices appended for CA lenders")
add_bullet(doc, "FCRA disclosure required when CRA data (CoreLogic, credit reports) used in decision")

add_h3(doc, "15.1 Explainability_Layer Architecture")
add_para(doc, "Three-stage process: interception → enrichment → assembly")
add_bullet(doc, "Intercept: any KILL tagged event routes to Explainability_Layer")
add_bullet(doc, "Enrich: append raw inputs + intermediate values + threshold breached")
add_bullet(doc, "Assemble: dynamic config lookup → ECOA codes + FCRA disclosures + state notices")

add_page_break(doc)

# ============== PART V — LIVE DATA & ARCHITECTURE ==============
add_h1(doc, "Part V — Live Data & Architecture")

add_h2(doc, "16. Live Rate Triplet + SOFR Anchors (June 17-18, 2026)")

add_para(doc, "Refreshed on session open + every 4 hours via Celery. Source: FRED + NY Fed (free, no CME license).")

add_table(doc,
    ["Series", "Value (Jun 16-18, 2026)", "Source"],
    [
        ["DGS10 (10Y Treasury)", "4.43%", "FRED"],
        ["DGS30 (30Y Treasury)", "—", "FRED"],
        ["DGS5 (5Y Treasury — ARM benchmark)", "—", "FRED"],
        ["SOFR Overnight", "3.63%", "FRED"],
        ["30-Day Avg SOFR", "3.609%", "NY Fed"],
        ["90-Day Avg SOFR", "3.636%", "NY Fed"],
        ["Fed Funds Effective", "3.50-3.75%", "FRED (held 4th consecutive FOMC)"],
        ["Conventional 30yr IP (Freddie Mac)", "6.53%", "Freddie Mac Jun 8"],
        ["DSCR Premium vs Conventional", "+50-125 bps", "Multiple sources"],
    ],
    widths=[2.5, 2.0, 1.5])

add_h3(doc, "16.1 Pricing Calibration (June 17, 2026 — verified)")
add_table(doc,
    ["Tier", "Rate Range", "Profile"],
    [
        ["Competitive", "6.125-6.49% (par 6.125%, 0 pts)", "740+ FICO, ≤70-75% LTV, 1.0+ DSCR — Griffin Jun 2026"],
        ["ARM (from 5.125%)", "5.125%+", "Same profile, ARM structure"],
        ["Typical", "6.50-7.50%", "Standard files"],
        ["Full-market", "up to 10.75%", "Thin/non-prime/low DSCR/STR/FN"],
    ],
    widths=[1.5, 2.0, 3.0])

add_h3(doc, "16.2 LLPA Pricing Levers (verified off 740/par anchor)")
add_table(doc,
    ["Lever", "Adjustment"],
    [
        ["FICO 760+", "−0.05 to −0.125"],
        ["FICO 720-739", "+0.125"],
        ["FICO 700-719", "+0.125 to +0.25"],
        ["FICO 680-699", "+0.50 (cliff)"],
        ["FICO 660-679", "+0.875 (cliff)"],
        ["FICO 640-659", "+1.50 to +2.50"],
        ["LTV per 5% increment", "+0.125 to +0.25"],
        ["DSCR per 0.10 below 1.25", "+0.125"],
        ["85% LTV (select lenders)", "740+/SFR purchase/DSCR ≥1.0 only"],
        ["IO", "+0.25"],
        ["ARM", "−0.125 to −0.375 vs 30yr fixed"],
        ["1 discount point", "≈ −0.25% rate"],
        ["Cash-out", "+0.25 to +0.50"],
        ["Loan <$150K", "DSCR floor often 1.25"],
        ["Foreign national", "+0.50 to +1.50"],
        ["No-PPP", "+0.50 to +0.80"],
        ["6+ mo reserves", "−0.10 to −0.25"],
        ["Rate lock 60d", "+0.125"],
    ],
    widths=[2.5, 2.0])

add_page_break(doc)

# ============== SECTION 17: RENTCAST + AIRDNA ==============
add_h2(doc, "17. RentCast + AirDNA + Optimal Blue Integration")

add_h3(doc, "17.1 RentCast (LTR / Form 1007 Equivalent)")
add_table(doc,
    ["Aspect", "Detail"],
    [
        ["Coverage", "140M+ property records, all 50 states"],
        ["Free tier", "50 API calls/month (Developer plan) — CORRECTED FROM V1.0"],
        ["Auth", "X-Api-Key header"],
        ["Base URL", "https://api.rentcast.io/v1/"],
        ["Critical endpoints", "/avm/rent/long-term, /avm/value, /markets"],
        ["STR support", "NONE — STR requires AirDNA"],
    ],
    widths=[1.2, 5.0])

add_callout(doc, 
    "v1.0 cited consumer platform pricing ($29/$99/$199). Those are LANDLORD PORTFOLIO TRACKING plans, "
    "NOT API pricing. API pricing is volume-based. v2.0 correction.",
    color=GREEN, label="v1.0 CORRECTION")

add_h3(doc, "17.2 RentCast Confidence Gates")
add_code_block(doc,
    "RENTCAST_CONFIDENCE_GATES = {\n"
    "    'HIGH':    {'min': 80, 'max': 100, 'action': 'USE',\n"
    "               'note': 'High confidence — use as primary rent input'},\n"
    "    'MEDIUM':  {'min': 60, 'max': 79,  'action': 'USE_WITH_WARN'},\n"
    "    'LOW':     {'min': 0,  'max': 59,  'action': 'REQUIRE_MANUAL'},\n"
    "    'MISSING': {'min': None,'max': None,'action': 'REQUIRE_MANUAL'}\n"
    "}\n"
    "# LOW/MISSING: Order Form 1007 from licensed appraiser")

add_h3(doc, "17.3 AirDNA (STR)")
add_table(doc,
    ["Aspect", "Detail"],
    [
        ["Coverage", "10M+ STR properties, 120K+ markets; Airbnb/VRBO/Booking.com"],
        ["Data accuracy", "97% (daily scrape of 100% listings)"],
        ["Pricing", "$15-40/month per market (Professional); Enterprise API custom ($50K+/yr)"],
        ["Free tier", "Limited market data, no Rentalizer"],
        ["Critical for", "DSCR deals where STR income is primary qualifying metric"],
    ],
    widths=[1.2, 5.0])

add_h3(doc, "17.4 Optimal Blue PPE / Loansifter")
add_table(doc,
    ["Aspect", "Detail"],
    [
        ["Access", "Commercial lender/broker API entitlements required"],
        ["Pricing", "$15K-$50K+/yr"],
        ["2026 features", "Virtual Economist AI/ML forecasting, Profitability Center, Competitive Data License Plus"],
        ["Lock time", "Cuts from 15 minutes to seconds via API"],
    ],
    widths=[1.2, 5.0])

add_h3(doc, "17.5 Hybrid AI-OCR Intake (Sprint 0)")
add_para(doc, "Multi-engine OCR pipeline:")
add_bullet(doc, "Docling — digital PDFs and table reconstruction")
add_bullet(doc, "Mistral OCR 2505 — scanned and handwritten addenda")
add_bullet(doc, "GPT-4o Vision — structured JSON via Pydantic schemas")
add_para(doc, "Every extracted field tagged with bounding-box ID + confidence score. Fields with < 0.85 confidence → Human-in-the-Loop queue. Market Rent Guardrail: any lease deviating > ±30% from live RentCast AVM is flagged (instant fraud/stale-lease detection).", italic=True)

add_page_break(doc)

# ============== SECTION 18: THREE-METRIC CREDIT ==============
add_h2(doc, "18. Three-Metric Credit Standard + Dual-Audience Architecture")

add_h3(doc, "18.1 The Three-Metric Credit Standard (Replaces Single DSCR for Credit)")
add_table(doc,
    ["Metric", "Question", "Target"],
    [
        ["DSCR (Cash Control)", "Can the borrower make the payment?", "Per lender matrix (typically ≥ 1.00-1.25)"],
        ["Debt Yield (Workout Metric)", "What is the lender's cap rate if they foreclose?", "≥ 9% institutional standard"],
        ["LTV (Loss-Given-Default)", "How much asset deflation can the lender absorb?", "Per matrix + LGD model"],
    ],
    widths=[1.8, 2.5, 1.7])

add_h3(doc, "18.2 All-In Effective Yield (AEY) — True Cost of Capital")
add_code_block(doc,
    "True_Cost(hold) = Interest_During_Hold + Points$ + Lender/Broker/UW Fees +\n"
    "                   Lock_Cost + Prepay(exit_year) + Refi_Costs(if planned)\n"
    "Render at 12/24/36/60-mo + APR-equivalent.\n"
    "\n"
    "All-In Effective Yield (AEY) = XIRR of actual borrower cash flows:\n"
    "  [Net_Proceeds_0, -P_1, -P_2, ..., -(P_n + Balance_n + PPP_n)]\n"
    "Net_Proceeds_0 = Loan_Amount - (Points$ + Lender_Fees)\n"
    "\n"
    "Algorithm: SciPy brentq — guaranteed convergence on non-monotonic mortgage flows\n"
    "# Lender with LOWEST AEY over expected hold is cheapest, regardless of stated rate\n"
    "# TWO-QUOTE RULE: always one flex/fit + one rate-competitive, with AEY delta in dollars\n"
    "# NEVER a single quote")

add_h3(doc, "18.3 10-Minute Committee-Grade Verdict + 150-Word Investment Thesis Block")
add_para(doc, "Structured output for the Loan Officer:")
add_bullet(doc, "Property metrics (DSCR 4-track + Debt Yield + LTV + Cap Rate + Cash-on-Cash)")
add_bullet(doc, "Qualification status (PASS / CONDITIONAL / NO-GO with reason)")
add_bullet(doc, "Returns stack (Pre-Tax IRR + After-Tax IRR + CoC Y1/Y3/Y5 + Equity Multiple)")
add_bullet(doc, "Binding risks (Kill-Switch Conditions: \"If 1007 comes back below $2,100/mo, deal flips to NO-GO\")")

add_h3(doc, "18.4 Adverse-Case Recourse Table")
add_para(doc, "When file hits NO-GO, engine generates operator-action mapping:")
add_table(doc,
    ["Failure", "Suggested Fix (ranked)"],
    [
        ["Track A DSCR = 0.94", "Reduce loan by $14K (→1.00x) OR route to IO product (1 day)"],
        ["Debt Yield < 9%", "Lower price by X OR seek higher-rent lease comp"],
        ["LTV > matrix", "Increase down payment OR switch to lower-LTV matrix"],
        ["FICO below floor", "No automatic fix — refer to manual UW"],
    ],
    widths=[2.0, 4.0])

add_h3(doc, "18.5 Kill-Switch Monitor (Continuous)")
add_bullet(doc, "Polls RentCast API every 30 days")
add_bullet(doc, "Monitors lender guideline diffs")
add_bullet(doc, "Tracks 10-Year Treasury (current 4.43%)")
add_bullet(doc, "Alerts LO within 1 hour on breach")

add_page_break(doc)

# ============== PART VI — IMPLEMENTATION ==============
add_h1(doc, "Part VI — Implementation Roadmap")

add_h2(doc, "19. v1.0 → v2.0 Corrections Summary")

add_callout(doc, 
    "v1.0 synthesis had 7 critical errors. v2.0 corrects all. Do NOT act on v1.0 numbers.",
    color=CRIT, label="MANDATORY CORRECTIONS")

add_table(doc,
    ["#", "v1.0 Wrong", "v2.0 Correct", "Source Authority"],
    [
        ["1", "Renumber Slice 1 ECOA codes", "DO NOT renumber — codes ARE Form C-1 verbatim", "FCRA PDF p.6"],
        ["2", "PA PPP threshold = $319,777", "$329,411 (2026, was $319,777 in 2025)", "Sprint 2 + Definitive Blueprint v3"],
        ["3", "LLC non-bank financing triggers FinCEN BOI", "NOT triggered (FinCEN Mar 2025 interim final rule)", "Definitive Blueprint v3 §FinCEN"],
        ["4", "Deterministic core = model under SR 11-7", "NOT a model under SR 26-02 (Apr 17, 2026)", "OCC Bulletin 2026-13"],
        ["5", "TimesFM 2.0 with 500M params", "TimesFM 2.5: 200M, 15,360 ctx, quantile head, XReg", "Google BigQuery docs Jun 12, 2026"],
        ["6", "Two DSCR tracks (Track 1/2)", "Four tracks (Lender/Investor/Stabilized/All-In)", "v16 Master Consolidated"],
        ["7", "PAL MFJ phase-out = $200K", "PAL phase-out completes at $150K for ALL individual filers", "IRC §469(i); Sprint 4 statutory analysis"],
        ["8", "RentCast $29/$99/$199 tiers", "API free tier = 50 calls/mo; paid = volume-based", "rentcast.io/api"],
        ["9", "Gaussian copula acceptable", "t-Copula ν=5-7 mandatory; Gaussian BANNED", "Sprint 6 + multiple sources"],
        ["10", "Three-track stress (Base/Conservative/Severe)", "Four archetype scenarios + macro stacking", "Sprint 6"],
    ],
    widths=[0.3, 2.5, 2.5, 1.5])

add_h2(doc, "20. Slice-by-Slice Build Plan (Consolidated)")

add_h3(doc, "20.1 Slice 1 — Already Shipped (132 tests, 94.37% coverage)")
add_bullet(doc, "payment.py — payment_factor formula (Decimal prec=28, MAX_TERM=600)")
add_bullet(doc, "dscr.py — Track 1 (Lender) + Track 2 (Stressed) [BUG-01/05/06 fixed]")
add_bullet(doc, "leverage.py — brentq deal_break_rate + bisection max_purchase")
add_bullet(doc, "ltv.py — BUG-01 (min price/appraisal), BUG-05 (OpEx in breakeven), BUG-06 (Decimal rate)")
add_bullet(doc, "compliance.py — ECOA codes 19/21/26/27/28 (CORRECT — DO NOT RENUMBER)")

add_h3(doc, "20.2 Slice 2 Plan (4-6 weeks, ~200 hr) — Highest Priority")
add_para(doc, "Objective: Make the engine QbD-capable with Sequential Drawdown + Macro Archetypes + 6-class Recommendation.")

add_table(doc,
    ["#", "Module", "Effort", "Dependencies"],
    [
        ["P0-1", "Sequential Drawdown Array (month-by-month cash simulation)", "30 hr", "Slice 1 payment.py"],
        ["P0-2", "Stress Scenario Engine (Base/Conservative/Severe + 4 Macro Archetypes)", "40 hr", "Sprint 6 KBRA-calibrated distributions"],
        ["P0-3", "DFS + ISS (min-gate) + QbD (7-trigger)", "40 hr", "Sprint 6 + Doc 17 pseudocode"],
        ["P0-4", "Counterfactual Generator (binary search)", "20 hr", "P0-3"],
        ["P0-5", "6-class Recommendation State Machine", "20 hr", "P0-3 + P0-4"],
        ["P0-6", "BUG-02 NOI growth off-by-one + BUG-03 vacancy tornado labels", "10 hr", "Slice 1 dscr.py"],
        ["P0-7", "Stabilized Economic NOI (Ledger 2 foundation)", "20 hr", "P0-1 + Slice 1 dscr.py"],
        ["P0-8", "MCID detector (STR-specific)", "15 hr", "P0-1"],
        ["P0-9", "Triangulated Rent Validator (4-source weighted + CV)", "15 hr", "Slice 1 dscr.py"],
        ["TOTAL", "", "210 hr (~5-6 weeks)", ""],
    ],
    widths=[0.7, 3.5, 0.8, 1.8])

add_h3(doc, "20.3 Slice 3 Plan (8-12 weeks, ~400 hr)")
add_table(doc,
    ["#", "Module", "Effort"],
    [
        ["P1-1", "Track 3 Stabilized DSCR + All-In DSCR (4-track complete)", "40 hr"],
        ["P1-2", "Multi-Year DSCR Trajectory (10-year roll-up w/ IO+ARM reset)", "60 hr"],
        ["P1-3", "Cap-Rate Linked Refi Solver (dual gate DSCR+LTV)", "60 hr"],
        ["P1-4", "Matrix Grid Solver + Multi-Variable Constrained Opt", "80 hr"],
        ["P1-5", "After-Tax Engine (OBBBA 100% bonus + §1250 + NIIT + PAL + REP)", "80 hr"],
        ["P1-6", "t-Copula Monte Carlo (Sprint 6 implementation)", "60 hr"],
        ["P1-7", "SHA-256 + Merkle Audit Trail + SR 26-02 model cards", "40 hr"],
        ["TOTAL", "All modules", "420 hr (~10 weeks)"],
    ],
    widths=[0.7, 4.0, 0.8])

add_h3(doc, "20.4 Slice 4 Plan (12-16 weeks, ~600 hr)")
add_bullet(doc, "Live Data APIs (FRED + RentCast + AirDNA — full Sprint 5 integration)")
add_bullet(doc, "AEY True Cost of Capital (brentq XIRR)")
add_bullet(doc, "Three-Metric Credit Standard output (DSCR + Debt Yield + LTV)")
add_bullet(doc, "10-Minute Committee-Grade Verdict + 150-word Investment Thesis Block")
add_bullet(doc, "Kill-Switch Monitor (continuous polling)")
add_bullet(doc, "TimesFM 2.5 forecasters for DSCR trajectory + OBBBA-aware scenarios")
add_bullet(doc, "Approval Predictor (XGBoost ensemble) — SR 26-02 model card required")
add_bullet(doc, "Backtest + PSI drift detection + champion/challenger")

add_h3(doc, "20.5 Slice 5 Plan (16+ weeks) — Multi-Tenant + IC Memo")
add_bullet(doc, "IC Memo Command (institutional credit memo generation)")
add_bullet(doc, "1031 Exchange deferral modeling")
add_bullet(doc, "Multi-tenant (5+ unit) + Commercial underwriting")
add_bullet(doc, "Full 50-state compliance variations")
add_bullet(doc, "PPP annual re-index cron (Jan 1)")

add_page_break(doc)

# ============== SECTION 21: ACTION ITEMS ==============
add_h2(doc, "21. Specific Code-Level Action Items")

add_h3(doc, "21.1 Immediate (this week)")
add_para(doc, "AI-1: Add BUG-02 fix (NOI growth off-by-one) to Slice 1 dscr.py.", bold=True)
add_para(doc, "Convention: Year1_NOI is the first full operating year. Year3_NOI = Year1_NOI × (1+g)^2 (not ^3). Affects Track 3 Stabilized DSCR, Levered IRR, Monte Carlo exit NOI. Effort: 2 hr.")

add_para(doc, "AI-2: Add BUG-03 fix (vacancy tornado labels) to Slice 1 dscr.py.", bold=True)
add_para(doc, "vac_best = max(0, vac - 0.05); vac_worst = min(1, vac + 0.05). Swing must be non-negative. Right side = High DSCR / Low Vacancy / Best Case. Effort: 1 hr.")

add_para(doc, "AI-3: Document SR 26-02 status inline in Slice 1 code.", bold=True)
add_para(doc, "Add docstring to each module: \"SR 26-02 status: NOT a model (simple arithmetic). Unit tests + CI/CD sufficient.\" Effort: 1 hr.")

add_h3(doc, "21.2 Slice 2 Entry (next 2 weeks)")
add_para(doc, "AI-4: Build drawdown.py (Sequential Drawdown Array).", bold=True)
add_para(doc, "Highest-leverage Slice 2 addition. Use Sprint 6 pseudocode. Effort: 8 hr skeleton + 16 hr tests.")

add_para(doc, "AI-5: Adopt macro archetype library as config JSON.", bold=True)
add_para(doc, "4 archetypes (Stagflation/Recession/Climate/Local Distress) per Sprint 6. Effort: 4 hr.")

add_para(doc, "AI-6: Switch ISS to minimum-gate.", bold=True)
add_para(doc, "ISS = min(S_DSCR, S_LSC, S_Refi, S_CapEx), not weighted average. Effort: 4 hr.")

add_h3(doc, "21.3 Slice 2 Build (4-6 weeks)")
add_para(doc, "AI-7: Implement 6-class Recommendation State Machine.", bold=True)
add_para(doc, "STRONG / MONITOR / FRAGILE / QbD-MINOR / QbD-MODERATE / QbD-CRITICAL / HALT / REJECT. Maps directly to remediation actions. Effort: 8 hr.")

add_para(doc, "AI-8: Implement Extended QbD (7 triggers per Doc 17 #21).", bold=True)
add_para(doc, "QbD = TRUE if lender_pass AND any of: DSCR_E<1.0 OR Seasonal_DSCR<1.0 OR RWDSCR<1.0 OR drawdown_fail OR refi_fail OR ACS<0.6 OR ISS<0.6. Effort: 8 hr.")

add_page_break(doc)

# ============== SECTION 22: SR 26-02 ==============
add_h2(doc, "22. SR 26-02 Compliance Status (The Biggest Moat)")

add_para(doc, "SR 26-02 (OCC Bulletin 2026-13, effective April 17, 2026) supersedes SR 11-7. Key implications for DSCR Sovereign OS:")

add_table(doc,
    ["Component", "SR 26-02 Class", "Required Governance", "Engine Impact"],
    [
        ["DSCR Calculator", "NOT a model", "Unit tests + CI/CD", "Ship 5x faster than competitors"],
        ["Legal Rules Engine", "NOT a model", "Quarterly counsel review", "Ship 5x faster"],
        ["After-Tax Engine", "NOT a model", "IRS source verification + unit tests", "Ship 5x faster"],
        ["Lender Qualification", "NOT a model (rule-based)", "Config + version control", "Ship 5x faster"],
        ["Monte Carlo", "HIGH-materiality MODEL", "Full model card + champion/challenger + outcomes", "Required for production"],
        ["TimesFM 2.5 / TFT", "MEDIUM-HIGH model", "Model card + backtesting", "Required for production"],
        ["Approval Predictor", "HIGH-materiality MODEL", "Full card + outcomes + disparate impact monitor", "Required for production"],
    ],
    widths=[1.8, 2.0, 2.5, 1.8])

add_callout(doc, 
    "Architectural moat: deterministic layer ships WITHOUT model governance overhead. Monte Carlo + ML "
    "require full SR 26-02 cards. Competitors operating under blanket SR 11-7 definition must build full "
    "model governance for the deterministic layer too — we don't. This is 60-70% governance overhead reduction "
    "on the most-used components.",
    color=GREEN, label="SR 26-02 MOAT")

add_h3(doc, "22.1 Model Card Template (for Monte Carlo + ML components)")
add_code_block(doc,
    "Model Card: [Component Name]\n"
    "  - SR 26-02 Classification: [HIGH-MATERIALITY MODEL / etc.]\n"
    "  - Purpose: [Description]\n"
    "  - Inputs: [List with sources]\n"
    "  - Outputs: [List]\n"
    "  - Methodology: [Copula type, distributions, etc.]\n"
    "  - Validation: [Backtest results]\n"
    "  - Champion Model: [Production baseline]\n"
    "  - Challenger Models: [Alternatives]\n"
    "  - Performance Metrics: [AUC, calibration, etc.]\n"
    "  - Fairness Audit: [Disparate impact test results]\n"
    "  - Monitoring: [PSI thresholds, retraining cadence]\n"
    "  - Limitations: [Known edge cases]\n"
    "  - Approver: [Name + Date]")

add_page_break(doc)

# ============== PART VII — APPENDICES ==============
add_h1(doc, "Part VII — Appendices")

add_h2(doc, "Appendix A: Document Inventory (60+ sources)")

add_h3(doc, "A.1 Master Specifications (MDs — 8)")
add_bullet(doc, "six-function-doctrine.md")
add_bullet(doc, "Advisor_Grade_DSCR_Decision_Engine_Usable_Master_Spec.md")
add_bullet(doc, "Advisor_Grade_DSCR_Decision_Engine_Organized_Research.md")
add_bullet(doc, "AEGIS_DSCR_Algorithm_Gap_Upgrade_Pack.md")
add_bullet(doc, "AEGIS_DSCR_Advisor_Grade_Operating_Model_Upgrade_Pack.md")
add_bullet(doc, "AEGIS_DSCR_Deterministic_Core_Keeps_Detailed.md")
add_bullet(doc, "AEGIS_DSCR_Complete_Usable_Master_Doc_v3.md")
add_bullet(doc, "DSCR_Engine_Master_Specification.md")

add_h3(doc, "A.2 Architectural Blueprint PDFs (14)")
add_bullet(doc, "From Black Box to Glass Box — adversarial hardening")
add_bullet(doc, "From Calculator to Counselor — adversarial validation blueprint")
add_bullet(doc, "From Calculation to Counsel — quantitative innovation")
add_bullet(doc, "Architecting the Advisor-Grade DSCR Engine — 11-module + lender adapter")
add_bullet(doc, "Beyond the DSCR — dual-ledger + stabilization")
add_bullet(doc, "From Static Snapshot to Dynamic Trajectory — temporal/path")
add_bullet(doc, "From Calculator to Containment — adversarial hardening MVP")
add_bullet(doc, "AI Algorithm Improvement Prompt — 15 research loops")
add_bullet(doc, "AI Algorithm Improvement Prompt 2 — per-formula hardening")
add_bullet(doc, "FCRA Adverse Action Engine for Institutional Compliance — reason codes + payload")
add_bullet(doc, "Beyond the Rulebook (x2) — dynamic data + probabilistic underwriting")
add_bullet(doc, "From Blueprint to Sovereign Engine — TimesFM hardening")
add_bullet(doc, "From Policy to Profit — Cake Mortgage 2026 strategy")
add_bullet(doc, "From Restriction to Dominance — Cake Mortgage arbitrage")

add_h3(doc, "A.3 Sprint Research Execution (6)")
add_bullet(doc, "Sprint 0 & 1 — Live Research Execution Findings")
add_bullet(doc, "Sprint 2 — PPP State Matrix + STR Legality + 40-Year Amortization")
add_bullet(doc, "Sprint 3 — Lender Footprints + Securitization Pool + Competitive Threat")
add_bullet(doc, "Sprint 4 — After-Tax IRR + OBBBA + Insurance Kill + Flood Gate + Compliance")
add_bullet(doc, "Sprint 5 — Live Data APIs + Rate Anchors + Property Tax Matrix + Architecture")
add_bullet(doc, "Sprint 6 — t-Copula MC + QuantLib ARM + After-Tax IRR + IC Memo + 1031 + XGBoost")

add_h3(doc, "A.4 Master Specs (6)")
add_bullet(doc, "DSCR_Underwriting_Engine_v14_Complete_Master_Document.md")
add_bullet(doc, "DSCR_Underwriting_Engine_Master_Consolidated_v16.md (BUG-02/03 + FLAW-01/02)")
add_bullet(doc, "DSCR_Sovereign_OS_Definitive_Master_Blueprint_v3.md (SR 26-02 + 7 critical corrections)")
add_bullet(doc, "DSCR SOVEREIGN OS_ THE DEFINITIVE PRODUCT SPECIFICATION.md (Dual-Audience v12)")
add_bullet(doc, "DSCR Sovereign OS THE MASTER BLUEPRINT.md")
add_bullet(doc, "THE COMPLETE SOVEREIGN MASTER DOCUMENT.md (2 versions — full + shorter)")

add_h3(doc, "A.5 Knowledge Documents (5)")
add_bullet(doc, "Master DSCR Knowledge Document.md")
add_bullet(doc, "The 2026 DSCR Master Knowledge Paper.md (Manus AI synthesis)")
add_bullet(doc, "DSCR Sovereign OS MASTER RESEARCH SYNTHESIS.md")
add_bullet(doc, "DSCR Sovereign OS & Non-QM Wholesale Lender Definitive Master Research Report.md")
add_bullet(doc, "DSCR DUAL TRUTH ENGINE CHATGPT RESEARCH.md")

add_h3(doc, "A.6 Analysis & Strategy Documents (10)")
add_bullet(doc, "DSCR Sovereign OS Upgrade Intelligence Report (3 versions)")
add_bullet(doc, "dscr_sovereign_os_architectural_debt_and_math.md")
add_bullet(doc, "dscr_sovereign_os_deep_debt_analysis.md")
add_bullet(doc, "DSCR_Sovereign_OS_Feature_Engineering_Blueprint.md")
add_bullet(doc, "DSCR_Appendix_B_Research_Resolution_Report.md")
add_bullet(doc, "DSCR_Blueprint_Verification_Corrections_Log.md")
add_bullet(doc, "DSCR Forumals.md")
add_bullet(doc, "Actionable Next Steps for the 20X DSCR Deal Engine.md")
add_bullet(doc, "NEW_DSCR Deal Desk Build-Ready Research Report.md")
add_bullet(doc, "TimesFM_LoRA_Complete_Engineering_Spec.md")

add_h3(doc, "A.7 Slice 1 Codebase (read)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/payment.py (4,709 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/dscr.py (13,069 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/leverage.py (12,653 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/ltv.py (8,472 B, BUG-01/05/06 fixed)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/compliance.py (13,012 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/golden_vectors.json (3,211 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/audit doc + 132 tests / 94.37% coverage")

add_page_break(doc)

# ============== APPENDIX B: PSEUDOCODE ==============
add_h2(doc, "Appendix B: Pseudocode Library (Canonical, Expanded)")

add_h3(doc, "B.1 Sequential Drawdown Array (P0-1)")
add_code_block(doc,
    "def sequential_drawdown(\n"
    "    monthly_rent: list[float],       # 36 months\n"
    "    monthly_opex: list[float],      # 36 months\n"
    "    monthly_debt_service: list[float],\n"
    "    capex_events: dict[int, float],\n"
    "    starting_reserves: float,\n"
    ") -> dict:\n"
    "    balance = starting_reserves\n"
    "    ruin_month = None\n"
    "    for t in range(36):\n"
    "        ncf = monthly_rent[t] - monthly_opex[t] - monthly_debt_service[t]\n"
    "        if t in capex_events: ncf -= capex_events[t]\n"
    "        balance += ncf\n"
    "        if balance < 0 and ruin_month is None: ruin_month = t + 1\n"
    "    return {'liquidity_failure': ruin_month is not None,\n"
    "            'ruin_month': ruin_month}")

add_h3(doc, "B.2 t-Copula Monte Carlo (Sprint 6 implementation)")
add_code_block(doc,
    "from scipy import stats\n"
    "import numpy as np\n"
    "\n"
    "CORRELATION_MATRIX = np.array([\n"
    "    [ 1.00, -0.55,  0.25,  0.35, -0.10],\n"
    "    [-0.55,  1.00,  0.15, -0.30,  0.05],\n"
    "    [ 0.25,  0.15,  1.00,  0.10, -0.05],\n"
    "    [ 0.35, -0.30,  0.10,  1.00,  0.20],\n"
    "    [-0.10,  0.05, -0.05,  0.20,  1.00],\n"
    "])\n"
    "\n"
    "def run_monte_carlo(noi, monthly_pitia, hold_years, n_trials=10_000, nu=6):\n"
    "    rng = np.random.default_rng(42)\n"
    "    L = np.linalg.cholesky(CORRELATION_MATRIX)\n"
    "    Z = rng.standard_normal((n_trials, 5))\n"
    "    chi2 = rng.chisquare(nu, size=(n_trials, 1))\n"
    "    X = (Z @ L.T) / np.sqrt(chi2 / nu)\n"
    "    U = stats.t.cdf(X, df=nu)\n"
    "    rent_growth = stats.norm.ppf(U[:, 0], loc=0.02, scale=0.05)\n"
    "    vacancy = np.clip(stats.beta.ppf(U[:, 1], a=2, b=36), 0, 0.35)\n"
    "    # ... map remaining marginals\n"
    "    # Returns: dscr_p5, dscr_p50, prob_dscr_sub_1_0, stress_verdict")

add_h3(doc, "B.3 OBBBA After-Tax IRR (Sprint 4)")
add_code_block(doc,
    "def compute_after_tax_irr(deal, tax_profile, hold_years):\n"
    "    # OBBBA 100% bonus (acquired after Jan 19, 2025)\n"
    "    # Cost-seg reclassification: 60% structure / 40% bonus-eligible\n"
    "    building_basis = deal.purchase_price - deal.land_value\n"
    "    if deal.cost_seg_elected and deal.purchase_price >= 500000:\n"
    "        bonus = building_basis * 0.40\n"
    "        structure = building_basis * 0.60\n"
    "        yr1_dep = bonus + structure / 27.5\n"
    "        ongoing_dep = structure / 27.5\n"
    "    # §1250 recapture at exit (residential: Bucket 1=0, Bucket 2=25%, Bucket 3=LTCG)\n"
    "    # NIIT 3.8% if MAGI > threshold (FIXED, NOT CPI)\n"
    "    # PAL phase-out: $150K complete (not $200K)\n"
    "    # REP status: >750 hrs + >50% real property -> no NIIT on rental\n"
    "    # XIRR via pyxirr")

add_h3(doc, "B.4 Extended QbD (Doc 17 — 7 triggers)")
add_code_block(doc,
    "def qualifies_but_dangerous(lender_pass, dscr_e, seasonal_dscr_min,\n"
    "                             rwdscr, drawdown_pass, refi_pass, acs, iss):\n"
    "    if not lender_pass: return False, 'LENDER_FAIL'\n"
    "    triggers = []\n"
    "    if dscr_e < 1.0: triggers.append('DSCR_E_BREACH')\n"
    "    if seasonal_dscr_min < 1.0: triggers.append('SEASONAL_BREACH')\n"
    "    if rwdscr < 1.0: triggers.append('RWDSCR_BREACH')\n"
    "    if not drawdown_pass: triggers.append('LIQUIDITY_FAILURE')\n"
    "    if not refi_pass: triggers.append('REFI_FAILURE')\n"
    "    if acs < 0.6: triggers.append('LOW_DATA_CONFIDENCE')\n"
    "    if iss < 0.6: triggers.append('LOW_SURVIVAL')\n"
    "    if not triggers: return False, 'NONE'\n"
    "    severity = ('CRITICAL' if len(triggers) >= 3 else\n"
    "                'MODERATE' if len(triggers) == 2 else 'MINOR')\n"
    "    return True, severity")

add_h3(doc, "B.5 ARM Reset + Cap-Rate Linked Refi (Sprint 5/3)")
add_code_block(doc,
    "def arm_reset_with_sofr_curve(loan_balance, initial_rate, margin,\n"
    "                               periodic_cap, lifetime_cap, sofr_rates):\n"
    "    \"\"\"Sprint 5 QuantLib ARM Reset\"\"\"\n"
    "    new_rate = sofr_rates['6m'] + margin\n"
    "    new_rate = min(new_rate, initial_rate + periodic_cap)\n"
    "    new_rate = min(new_rate, initial_rate + lifetime_cap)\n"
    "    # PMT(loan, new_rate, remaining_term)\n"
    "\n"
    "def break_even_refi_cap_linked(noi, loan_balance, current_rate,\n"
    "                                target_dscr, current_cap_rate,\n"
    "                                cap_rate_beta, max_matrix_ltv,\n"
    "                                remaining_term_months):\n"
    "    # Dual gate: DSCR AND LTV\n"
    "    def check(rate):\n"
    "        pmt = compute_pmt(loan_balance, rate, remaining_term_months)\n"
    "        dscr = noi / (pmt * 12)\n"
    "        proj_cap = current_cap_rate + cap_rate_beta * (rate - current_rate)\n"
    "        proj_value = noi / proj_cap\n"
    "        ltv = loan_balance / proj_value\n"
    "        return (dscr >= target_dscr) and (ltv <= max_matrix_ltv)\n"
    "    # Binary search for break_even_rate")

add_h3(doc, "B.6 6-Class Recommendation (Doc 17 + State Machine)")
add_code_block(doc,
    "def recommend(acs, lender_pass, qbd_severity, iss, dfs, drawdown_pass):\n"
    "    if acs < 0.6:\n"
    "        return 'HALT - INSUFFICIENT DATA', ['ACS_BELOW_THRESHOLD']\n"
    "    if not lender_pass:\n"
    "        return 'REJECT - DOES NOT QUALIFY', ['LENDER_FAIL']\n"
    "    if qbd_severity == 'CRITICAL':\n"
    "        return 'QUALIFIES BUT CRITICAL RISK', ['QBD_CRITICAL']\n"
    "    if qbd_severity == 'MODERATE':\n"
    "        return 'QUALIFIES BUT DANGEROUS', ['QBD_MODERATE']\n"
    "    if qbd_severity == 'MINOR' or not drawdown_pass or iss < 50 or dfs < 30:\n"
    "        return 'FRAGILE - MONITOR CLOSELY', ['FRAGILITY_FLAG']\n"
    "    if iss >= 80 and dfs >= 60:\n"
    "        return 'STRONG DEAL', ['STRONG_PASS']\n"
    "    return 'ACCEPTABLE - MONITOR KEY RISKS', ['MONITOR_ONLY']")

add_h3(doc, "B.7 Adverse_Action_Notice_Payload (FCRA PDF spec)")
add_code_block(doc,
    "def emit_adverse_action(lender_id, application_id, kill_event,\n"
    "                       fcra_data_sources, state):\n"
    "    \"\"\"FCRA PDF p.8-9 compliant payload\"\"\"\n"
    "    reasons = dynamic_ecoa_mapping(kill_event)\n"
    "    payload = {\n"
    "        'version': '1.0',\n"
    "        'as_of': utcnow().isoformat(),\n"
    "        'lender_id': lender_id,\n"
    "        'application_id': application_id,\n"
    "        'is_compliant': True,\n"
    "        'regulatory_notices': {\n"
    "            'ecoa_notice': {\n"
    "                'header': 'ADVERSE ACTION NOTICE',\n"
    "                'prohibition_statement': REG_B_PROHIBITION_TEXT,\n"
    "                'reasons': reasons\n"
    "            },\n"
    "            'fcra_disclosure': {\n"
    "                'header': 'DISCLOSURE REQUIRED BY THE FAIR CREDIT REPORTING ACT',\n"
    "                'statement': FCRA_STANDARD_TEXT,\n"
    "                'data_source': fcra_data_sources[0]['name'],\n"
    "                'source_address': fcra_data_sources[0]['address']\n"
    "            }\n"
    "        },\n"
    "        'state_specific_notices': get_state_specific_notices(state),\n"
    "        'meta': {\n"
    "            'generation_timestamp': utcnow().isoformat(),\n"
    "            'engine_version': ENGINE_VERSION,\n"
    "            'explanation_layer_version': '1.0.0'\n"
    "        }\n"
    "    }\n"
    "    return payload")

add_page_break(doc)

# ============== APPENDIX C: LIVE DATA ==============
add_h2(doc, "Appendix C: Live Rate Triplet + Market Data (June 17-18, 2026)")

add_table(doc,
    ["Data Point", "Value (Jun 17-18, 2026)", "Source", "Refresh"],
    [
        ["10-Year Treasury (DGS10)", "4.43%", "FRED", "Every 4 hours"],
        ["SOFR Overnight", "3.63%", "FRED", "Every 4 hours"],
        ["30-Day Avg SOFR", "3.609%", "NY Fed (free)", "Daily"],
        ["90-Day Avg SOFR", "3.636%", "NY Fed", "Daily"],
        ["1M Term SOFR", "3.637%", "CME (paid license)", "Continuous"],
        ["3M Term SOFR", "3.668%", "CME", "Continuous"],
        ["6M Term SOFR", "3.731%", "CME", "Continuous"],
        ["12M Term SOFR", "3.869%", "CME", "Continuous"],
        ["Fed Funds Effective", "3.50-3.75%", "FRED", "After FOMC"],
        ["Conventional 30yr IP", "6.53%", "Freddie Mac Jun 8", "Weekly"],
        ["DSCR Premium over Conv", "+50-125 bps", "Multiple", "Per rate sheet"],
        ["Visio DSCR Q1 2026 range", "6.75-9.50%", "HonestCasa aggregator", "Per rate sheet"],
        ["KBRA CMBS 30+ DQ Feb 2026", "7.5%", "KBRA", "Monthly"],
        ["KBRA Office DQ Feb 2026", "12.8%", "KBRA", "Monthly"],
        ["KBRA CMBS Distress Feb 2026", "10.3%", "KBRA", "Monthly"],
        ["DSCR Origination Growth Jan 2025 YoY", "+123%", "Market data", "Monthly"],
        ["DSCR Avg Rate Q1 2026", "Below 7% (first time since Jun 2022)", "Market data", "Quarterly"],
    ],
    widths=[2.5, 2.2, 1.5, 1.3])

add_page_break(doc)

# ============== APPENDIX D: LENDER FOOTPRINT ==============
add_h2(doc, "Appendix D: Lender Footprint Matrix (June 2026)")

add_table(doc,
    ["Lender", "NMLS", "States", "Min FICO", "DSCR Floor", "Max LTV (P/CO)", "Key Feature"],
    [
        ["Visio Lending", "1935590", "41+DC", "680", "1.00 (sub-1.0 case-by-case)", "80%/75%", "#1 DSCR; 11 S&P deals ~$2B"],
        ["Kiavi", "—", "49+DC", "—", "—", "—", "AVM-heavy, rapid closings"],
        ["Angel Oak", "1160240", "47+DC", "—", "—", "—", "Rental AVM (Clear Capital) Nov 2025"],
        ["Griffin Funding", "—", "46+DC", "—", "Sub-1.0", "—", "No-ratio, jumbo DSCR"],
        ["LendingOne", "—", "All+exempt", "—", "—", "—", "—"],
        ["Lima One Capital", "—", "National", "—", "—", "—", "$10B+ lifetime; DSCR/bridge/construction"],
        ["Deephaven", "—", "National", "680", "1.00 (combined)", "80% CLTV", "DSCR Second (no reserves, no income docs)"],
        ["MortgageDepot", "—", "Verified", "—", "—", "—", "40-yr amort + 40-yr IO up to $3M"],
        ["Sistar Mortgage", "—", "Verified", "—", "—", "—", "40-yr IO confirmed 2026"],
        ["Cake Mortgage (Non-QM Wholesale)", "—", "National", "—", "—", "—", "Bank statement + DSCR; primary research target"],
    ],
    widths=[2.0, 0.8, 0.9, 0.7, 1.0, 1.0, 1.5])

add_h3(doc, "D.1 Angel Oak Rental AVM (Industry-First)")
add_para(doc, "Launched Nov 4, 2025. Powered by Clear Capital. Locks at pre-qualification, held through closing (unless property materially changes).")
add_callout(doc, 
    "Engine implication: surface Angel Oak as preferred lender when speed-of-certainty matters. "
    "Eliminates #1 DSCR deal-kill risk (appraisal coming back with lower rent than assumed).",
    color=ACCENT, label="COMPETITIVE EDGE")

add_page_break(doc)

# ============== APPENDIX E: REFERENCES ==============
add_h2(doc, "Appendix E: References & Source Anchors")

add_h3(doc, "E.1 Primary Regulatory Sources")
add_bullet(doc, "OCC Bulletin 2026-13 — SR 26-02 Model Risk Management (effective April 17, 2026)")
add_bullet(doc, "Fannie Mae Selling Guide (B2-2-03, B2-3-03, B7-3-02) — DSCR, NCF, replacement reserves")
add_bullet(doc, "Freddie Mac Single-Family Seller/Servicer Guide")
add_bullet(doc, "CFPB Circular 2022-03 — Adverse action reasons for complex algorithms")
add_bullet(doc, "Regulation B (ECOA) — Equal Credit Opportunity Act")
add_bullet(doc, "Regulation Z (TRID) — Truth in Lending disclosures")
add_bullet(doc, "IRC §1411 — Net Investment Income Tax (NIIT)")
add_bullet(doc, "IRC §469 — Passive Activity Loss (PAL) rules")
add_bullet(doc, "IRC §1250 — Unrecaptured gain")
add_bullet(doc, "Treasury/IRS Notice 2026-11 — OBBBA Bonus Depreciation")
add_bullet(doc, "FinCEN Interim Final Rule (Mar 21-26, 2025) — Corporate Transparency Act BOI")
add_bullet(doc, "HOEPA High-Cost Mortgage Thresholds (2026: $27,592 loan / $1,380 P&F)")
add_bullet(doc, "FHFA AB-2022-03 — Fair Lending AI/ML")
add_bullet(doc, "Basel III Finalising Post-Crisis Reforms (BCBS d424)")
add_bullet(doc, "EBA 2025 EU-wide Stress Test")
add_bullet(doc, "PA Act 6, 10 Pa. Code §7.2; OH ORC §1343.011")
add_bullet(doc, "MN HF 3437 (enacted April 23, 2026; effective Aug 1, 2026)")

add_h3(doc, "E.2 Market Data Sources")
add_bullet(doc, "KBRA Single-Borrower CMBS Default and Loss Study")
add_bullet(doc, "KBRA CMBS Loan Performance Trends (Feb/Apr 2026)")
add_bullet(doc, "KBRA CRE CLO Loan Default and Loss Study (Jun 2026)")
add_bullet(doc, "Trepp CMBS distress rate")
add_bullet(doc, "CRED iQ CMBS Conduit Underwriting Trends (Feb 2026)")
add_bullet(doc, "S&P Global Ratings DSCR adjustment factor methodology (NRMLT 2026-NQM1)")
add_bullet(doc, "Fitch Ratings Multifamily/Office CMBS Delinquency")
add_bullet(doc, "Matthews 2026 Cap Rate Analysis")
add_bullet(doc, "NCREIF Property Index")
add_bullet(doc, "MSCI Real Capital Analytics")
add_bullet(doc, "MMCG CRE Insights")
add_bullet(doc, "John Burns Real Estate Consulting")
add_bullet(doc, "AirDNA Best Places to Invest 2026")
add_bullet(doc, "RentCast API (rentcast.io/api)")
add_bullet(doc, "Optimal Blue PPE 2026")

add_h3(doc, "E.3 Academic / Methodology")
add_bullet(doc, "Demarta & McNeil (2005) — t-copula dependency modeling, DOI")
add_bullet(doc, "Artzner et al (1999) — Coherent risk measures")
add_bullet(doc, "BIS Supervisory Stress Testing (BCBS d427)")
add_bullet(doc, "IMF Macro-Financial Stress Test Framework")
add_bullet(doc, "OECD Commercial Real Estate Markets 2024")
add_bullet(doc, "ESRB CRE Lending Risks (Occasional Paper No. 29)")
add_bullet(doc, "UNC Cap Rate Determinants (Tsui-Morgan 2025)")
add_bullet(doc, "MIT Sloan Statistical Learning for Finance")
add_bullet(doc, "TimesFM 2.5 paper (Google Research)")

add_h3(doc, "E.4 Reference Preservation")
add_para(doc, "All source documents preserved in workspace. PDF text extractions in RESEARCH/pdf_extractions/. Sprint files in RESEARCH/sprint_short/. Full bibliography of 170+ entries preserved in the original PDFs (Calculator to Counselor, Calculation to Counsel, Beyond the DSCR, From Static Snapshot, From Calculator to Containment).")

# ============== END ==============
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_run = end_p.add_run("--- END OF v2 SYNTHESIS ---\n"
                        "Prepared for DSCR Sovereign OS / 20X DSCR Deal Engine\n"
                        "Workspace: C:\\Users\\serge\\OneDrive\\Documents\\DSCR_LOAN OFFICE\n"
                        "Output: output/doc/DSCR_Advisor_Engine_Cross_Doc_Synthesis_v2_20260619.docx\n"
                        "v2.0 corrects 7 errors in v1.0; adds Sprint 0-6, FCRA, Definitive Blueprint,\n"
                        "Definitive Product Spec, v16 fixes, Master Knowledge, 2026 Master Knowledge Paper\n"
                        "60+ source documents analyzed end-to-end")
end_run.font.size = Pt(10)
end_run.font.italic = True
end_run.font.color.rgb = DARK_GRAY

# Save
doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
print(f"Size: {OUT_FILE.stat().st_size / 1024:.1f} KB")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Sections: {len(doc.sections)}")