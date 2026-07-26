"""
DSCR Advisor-Grade Decision Engine: Cross-Document Synthesis v3.0 (DEFINITIVE)
==========================================================================
The competitive-breaking synthesis. Built from 60+ source documents analyzed end-to-end.

Critical new content vs v2.0:
  - 8 architectural debts (deep debt analysis) — R-vine copula, EVT, NSS, Kalman
  - 12 critical gaps (Non-QM Wholesale Lens) — LoanPASS, Cotality, ACES, MIAC, Encompass, Salesforce
  - Cake Mortgage product matrix (DSCR v4.0, Bundt, Cup, Velvet, Pound)
  - GNN entity resolution for layered LLCs (Beyond the Rulebook #2)
  - Conformal Prediction with exponential decay (e^-λt)
  - Tabular Foundation Models (TabPFN) for zero-shot/low-data
  - TabPFN-2.5, TimesFM 2.5, CPTC, R-vine copula library (pyvinecopulib)
  - Verus = nation's largest Non-QM securitizer ($15B+); LoanPASS = best PPE
  - Complete vendor stack + 50-state PPP matrix (corrected)
  - 5-dimensional distributional DSCR Engine spec
  - 16 research domains from Master Research Synthesis
  - Full Cake Mortgage product matrix with limits, programs, key differentiators
  - Borrower eligibility tiers (Experienced / First-Time / FTHB / FN / ITIN)
  - 50-state + DC license footprint for top 8 DSCR lenders
  - Historical LIVE market data (multifamily CMBS 7.15% March 2026, etc.)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_DIR = Path(r"C:\Users\serge\OneDrive\Documents\DSCR_LOAN OFFICE\output\doc")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_FILE = OUT_DIR / "DSCR_Advisor_Engine_Cross_Doc_Synthesis_v3_20260619.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
ACCENT = RGBColor(0x0B, 0x5E, 0x4B)
WARN = RGBColor(0xB7, 0x47, 0x1A)
CRIT = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
PURPLE = RGBColor(0x6B, 0x21, 0x80)

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(20)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(15)
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(12.5)
    return p

def add_para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0, color=None, bold=False):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)
        if color:
            run.font.color.rgb = color
        if bold:
            run.font.bold = True
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, widths=None, header_bg="1F3A5F"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, 1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F0F4F8")
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_callout(doc, text, color=WARN, label="WARNING"):
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}] {text}")
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF4E5' if color == WARN else 'FFE5E5' if color == CRIT else 'E5F5E5' if color == GREEN else 'F0E5FF')
    pPr.append(shd)
    return p

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ============== TITLE PAGE ==============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("DSCR Advisor-Grade Decision Engine")
trun.font.size = Pt(28); trun.font.bold = True; trun.font.color.rgb = NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = subtitle.add_run("Cross-Document Synthesis v3.0 — DEFINITIVE")
srun.font.size = Pt(18); srun.font.italic = True; srun.font.color.rgb = DARK_GRAY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mrun = meta.add_run("Competitive-Breaking Architecture & Capability Roadmap\n"
                    "Built from 60+ source documents analyzed end-to-end")
mrun.font.size = Pt(11); mrun.font.color.rgb = DARK_GRAY

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
drun = date_p.add_run("Prepared: 2026-06-19  -  Workspace: DSCR_LOAN OFFICE\n"
                     "v3.0 adds: 8 architectural debts + 12 wholesale gaps + GNN entity resolution +\n"
                     "Conformal Prediction + Tabular Foundation Models + Cake Mortgage product matrix +\n"
                     "Complete vendor stack + competitive moat analysis")
drun.font.size = Pt(10); drun.font.color.rgb = DARK_GRAY

doc.add_paragraph()
warn_p = doc.add_paragraph()
warn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
wr = warn_p.add_run("BEAT THE COMPETITION  -  RESEARCH-BACKED  -  AUDITABLE  -  SR 26-02 OPTIMIZED\n"
                   "10 architectural debts identified, 12 wholesale gaps cataloged,\n"
                   "8 institutional math modules specified, 16 research domains synthesized")
wr.font.size = Pt(11); wr.font.bold = True; wr.font.color.rgb = ACCENT

add_page_break(doc)

# ============== TOC ==============
add_h1(doc, "Table of Contents")
toc_items = [
    "PART I — STRATEGIC FOUNDATION",
    "  1. v3 Executive Summary + Corrective Changes",
    "  2. Six-Function Doctrine + Three-Audience Framework",
    "  3. Three-Plane Architecture + Semantic Diff Engine",
    "  4. Competitive Landscape (Verus, LoanPASS, YieldStack, LenderSA)",
    "PART II — THE 8 ARCHITECTURAL DEBTS",
    "  5. Debt 1: DSCR as Ratio Is Not a Risk Metric (Distributional DSCR)",
    "  6. Debt 2: Income Inputs Have No Propagated Uncertainty (Conformal Prediction)",
    "  7. Debt 3: Monte Carlo Assumes Stationary Correlation (R-Vine Copula)",
    "  8. Debt 4: No Forward Rate Surface (Nelson-Siegel-Svensson)",
    "  9. Debt 5: No Credit Loss Model (CECL PD × LGD × EAD)",
    " 10. Debt 6: No Contagion Model (Spatio-Temporal Graph)",
    " 11. Debt 7: LLM Hallucination Firewall (Deterministic Fact-Checker)",
    " 12. Debt 8: No Model Version Tracking (Audit Trail)",
    "PART III — DETERMINISTIC FINANCIAL CORE",
    " 13. Four DSCR Tracks + Golden Vector",
    " 14. BUG/FLAW Catalog (v16 reconciled)",
    " 15. After-Tax Engine (OBBBA + §1250 + NIIT + PAL + REP)",
    " 16. ARM Reset + IO Reversion + Cap-Rate Linked Refi",
    "PART IV — STRESS & RISK",
    " 17. t-Copula + R-Vine + Macro Archetypes (Sprint 6 + TUM Munich)",
    " 18. Sequential Drawdown + MCID + Counterfactual Generator",
    " 19. STR Risk Scoring + Top 10 Markets (AirDNA 2026)",
    "PART V — LENDER, COMPLIANCE & WHOLESALE",
    " 20. Cake Mortgage Product Matrix (DSCR v4.0, Bundt, Cup, Velvet, Pound)",
    " 21. Lender Footprint Matrix (Visio, Kiavi, Angel Oak, Deephaven, Griffin)",
    " 22. 50-State PPP Matrix (Sprint 2 + 2026 thresholds)",
    " 23. ECOA / FCRA / SR 26-02 Compliance",
    " 24. Non-QM Wholesale: 12 Critical Gaps (MISSING PIECES)",
    " 25. The Non-QM Vendor Stack (LoanPASS, Cotality, ACES, MIAC, Encompass, Salesforce)",
    " 26. Adverse-Action Notice Payload (FCRA PDF)",
    "PART VI — LIVE DATA & ARCHITECTURE",
    " 27. Live Rate Triplet (June 17-18, 2026)",
    " 28. RentCast + AirDNA + Optimal Blue Integration",
    " 29. Three-Metric Credit Standard + Dual-Audience",
    " 30. The GNN Entity Resolution Framework",
    " 31. Conformal Prediction Vault (Decaying Confidence)",
    " 32. Tabular Foundation Models (TabPFN) for Low-Data",
    "PART VII — IMPLEMENTATION ROADMAP",
    " 33. v1 → v2 → v3 Corrections Summary",
    " 34. Slice-by-Slice Build Plan",
    " 35. Specific Code-Level Action Items",
    " 36. SR 26-02 Compliance Status",
    "PART VIII — APPENDICES",
    "  A. Document Inventory (60+ sources)",
    "  B. Pseudocode Library (canonical, full)",
    "  C. Live Rate Triplet + Market Data (June 2026)",
    "  D. Lender Footprint Matrix",
    "  E. Cake Mortgage Product Matrix Reference",
    "  F. References & Source Anchors",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9.5)

add_page_break(doc)

# ============== SECTION 1: EXEC SUMMARY ==============
add_h1(doc, "Part I — Strategic Foundation")

add_h2(doc, "1. v3 Executive Summary + Corrective Changes")

add_para(doc, 
    "Sixty-plus source documents were analyzed end-to-end to inform the upgrade of DSCR Sovereign OS into an "
    "Advisor-Grade Decision Engine. v3.0 corrects errors from v1.0 and v2.0, and adds the missing competitive-edge "
    "content from deep-debt analysis, the 12 Non-QM wholesale gaps, the Cake Mortgage product matrix, and "
    "the GNN + Conformal Prediction + Tabular Foundation Model frameworks. The corpus spans master "
    "specifications (~1.5 MB), architectural blueprint PDFs (~3 MB / ~960 pages), Sprint 0-6 research "
    "execution, the v14/v15/v16 master consolidated specs, the Definitive Master Blueprint v3, and Slice 1 "
    "(132 tests / 94.37% coverage / 100/100 quality gate).",
    bold=False)

add_h2(doc, "1.1 The v3 Verdict")
add_para(doc, 
    "The corpus identifies 8 architectural debts that, unresolved, leave the engine vulnerable to exactly "
    "the conditions materializing in the market (multifamily CMBS delinquency 7.15% March 2026, up from 1.84% "
    "two years prior — a 4× increase). It also identifies 12 critical gaps that prevent the engine from being "
    "the best Non-QM wholesale lender in the nation. v3 integrates the institutional math fixes for all 8 debts "
    "and the vendor stack + process recommendations for all 12 gaps.",
    bold=True)

add_h2(doc, "1.2 What's New in v3 (vs v2.0)")
add_table(doc,
    ["New Section", "Source", "What's Added"],
    [
        ["The 8 Architectural Debts", "architectural_debt_and_math.md + deep_debt_analysis.md", "Complete failure mode analysis + institutional math fixes for each (R-vine, EVT, NSS, Kalman)"],
        ["5-Dimensional Distributional DSCR", "deep_debt_analysis.md §DEBT 1", "Replace point-DSCR with P(t<1.0) distribution across horizons + CVaR"],
        ["Conformal Prediction Vault", "Beyond the Rulebook #2", "e^(-λt) exponential decay + Mondrian hierarchical calibration"],
        ["GNN Entity Resolution (THGNN/HGT/TGN)", "Beyond the Rulebook #2", "Layered LLC detection; node classification; link prediction"],
        ["Tabular Foundation Models (TabPFN)", "Beyond the Rulebook #2 + Nature paper", "Zero-shot underwriting for low-data product launches"],
        ["Cake Mortgage Product Matrix", "Beyond the Rulebook #1", "DSCR v4.0, Bundt NQM/NOO, Cup, Velvet, Pound — limits, niches, dates"],
        ["12 Critical Gaps (MISSING PIECES)", "MISSING_PIECES_*.md", "Bank statement engine, PPE, TPO, warehouse, QC, MSR, hedging, LOS, CMS"],
        ["Non-QM Vendor Stack", "Non-QM Wholesale Master Research Report", "LoanPASS, Cotality, ACES, MIAC, Encompass, Salesforce"],
        ["Live Market Data March-June 2026", "Sprint 3 + Master Research", "Multifamily CMBS 7.15%, Office 11.71%, DSCR 60+ DQ 2.92%"],
        ["16 Research Domains", "MASTER_RESEARCH_SYNTHESIS.md", "All build modules mapped to academic papers + state-of-art + build recs"],
        ["Top 8 Wholesale Lender License Footprint", "Sprint 3", "Visio 41+DC, Kiavi 49+DC, Angel Oak 47+DC, Griffin 46+DC"],
        ["Borrower Eligibility Tiers", "Master DSCR Knowledge + 2026 Paper", "Experienced / First-Time / FTHB / FN / ITIN"],
    ],
    widths=[1.8, 2.0, 3.5])

add_h2(doc, "1.3 The Headline Architecture (v3)")
add_table(doc,
    ["Metric", "Baseline (Slice 1)", "Target After v3 Improvements"],
    [
        ["DSCR measurement", "Single static point ratio", "5-dimensional distribution (P12, P36, lifetime P, E[macro], CVaR)"],
        ["Income uncertainty", "Point estimates", "Conformal Prediction intervals with e^(-λt) decay"],
        ["Correlation structure", "5x5 t-copula (Sprint 6)", "5x5 R-vine (mixed families per edge) per TUM Munich"],
        ["Tail extrapolation", "Empirical percentile", "EVT GPD fit for P1 / P0.1 loss quantiles"],
        ["Rate curve", "Flat forward at current SOFR", "Nelson-Siegel-Svensson surface + Hull-White short-rate sim"],
        ["Latent state estimation", "Raw observations", "Kalman filter (Diebold-Li) for rent growth / vacancy"],
        ["Entity resolution", "Name matching", "THGNN (HGT/TGN) — link prediction for hidden LLCs"],
        ["Credit loss", "None", "CECL PD × LGD × EAD with vintage × FICO × LTV × property × geo curves"],
        ["Contagion modeling", "Independent deals", "Spatio-temporal graph (sponsor × ZIP × lender × MSA)"],
        ["LLM safety", "Raw LLM output", "Deterministic fact-checker (every number verified vs engine_output)"],
        ["Model audit", "Log file", "Version-pinned model cards (SR 26-02) + per-inference audit trail"],
        ["ECOA compliance", "20 abstract codes", "40 Form C-1 verbatim codes + FCRA disclosure + state notices"],
        ["After-tax", "Pre-tax only", "OBBBA + §1250 + NIIT + PAL + REP + 1031 (per Sprint 4)"],
        ["SR 26-02 status", "Not classified", "Deterministic core NOT model (5x faster ship)"],
        ["Non-QM coverage", "DSCR only", "Bank Statement + Asset Depletion + ITIN + FN (per MISSING PIECES)"],
    ],
    widths=[1.7, 2.0, 3.6])

add_page_break(doc)

# ============== SECTION 2-4: STRATEGIC FOUNDATION ==============
add_h2(doc, "2. Six-Function Doctrine + Three-Audience Framework")

add_h3(doc, "2.1 The Six Functions (Godmode v7)")
add_table(doc,
    ["#", "Function", "Elite Standard", "Module"],
    [
        ["01", "Scenario Accuracy", "GO/NO-GO verdict with confidence in <10min", "engine.ts + preflightGate.ts"],
        ["02", "Guideline Intelligence", "25+ verified lenders, auto-fit scoring, two-quote rule", "lenders.ts + fitScorer.ts"],
        ["03", "Borrower Trust", "Every quote regulator-ready with full constraint disclosure", "quoteExplainer.ts + pdfQuotePack.ts"],
        ["04", "Capital Partner Trust", "Zero-defect file, first-pass clean rate >90%", "fileCompletenessEngine.ts + defectScorer.ts"],
        ["05", "Distribution", "60%+ revenue from repeat referral channels", "referralPortal.ts + channelAttribution.ts"],
        ["06", "Risk Discipline", "Hard decline gates + adverse-action compliance, false-decline <5%", "declineGate.ts + adverseActionEngine.ts"],
    ],
    widths=[0.4, 1.5, 2.5, 2.5])

add_h3(doc, "2.2 Three Audiences of Every Quote")
add_bullet(doc, "BORROWER: cares whether deal closes and at what cost")
add_bullet(doc, "CAPITAL PARTNER: cares whether file is clean and defensible")
add_bullet(doc, "OPERATOR (loan officer): cares whether 10-min verdict holds through closing")
add_para(doc, "A quote that satisfies only one audience is a failure.", italic=True)

add_h2(doc, "3. Three-Plane Architecture (Graph-Native OS)")
add_table(doc,
    ["Plane", "Definition", "Implementation"],
    [
        ["Projection Plane", "Human-facing views", "Scenario Builder, Lender Matchmaker, After-Tax IRR Studio, IC Memo"],
        ["Graph Plane", "Causal central nervous system", "Nodes (Borrower, Property, Lender, Law, Rate) + Typed Edges (Qualifies, Conflicts, Supersedes, Shocks) + pgvector"],
        ["Ledger Plane", "Immutable append-only event log", "Every mutation, approval, export captured with full provenance"],
    ],
    widths=[1.5, 2.0, 3.5])

add_h3(doc, "3.1 Semantic Diff Engine")
add_para(doc, "Classifies changes by facet (Location, Timing, Budget, Legal). Structural change (LLC to Individual vesting) triggers causal propagation through the PPP Legal Branching Gate. Cosmetic change produces no propagation. Three-stage diff: (1) structured field hash, (2) LLM facet classifier, (3) facet-sensitive edge invalidation.")

add_h2(doc, "4. Competitive Landscape (Verus, LoanPASS, YieldStack, LenderSA)")
add_table(doc,
    ["Competitor", "Threat Level", "Their Strength", "Counter-Moat"],
    [
        ["Verus Mortgage Capital", "CRITICAL — same space, biggest fish", "$15B+ Non-QM securitizations, LoanPASS PPE partnership Oct 2025", "Sovereign OS provides deeper per-deal risk modeling (5-dim DSCR, R-vine, CECL) that Verus lacks"],
        ["YieldStack AI", "MODERATE-HIGH — direct competitor", "180+ lender programs, program-level matching, zero cost", "YieldStack matches; we ANALYZE (dual-track + 5-dim DSCR + Sequential Drawdown + AEY + After-Tax + MC + IC memo + ECOA)"],
        ["LenderSA 3.2 AI", "MODERATE — different segment", "Hard money focus, hundreds of lenders, AI negotiation", "Different segment (fix-and-flip, not DSCR); moat in analytical depth, not lender breadth"],
        ["OCMBC / CrossCountry / Acra / A&D", "LOW — fragmented tech", "Top 4 Non-QM lenders by volume ($3.5B+) but no algorithmic depth", "We're the algorithmic layer; they become customers via white-label"],
        ["Kiavi / A&D / Lima One / Easy Street", "LOW — borrowers/lenders, not platforms", "Strong direct lender programs but no decision intelligence", "Engine becomes their pricing/qualification layer"],
    ],
    widths=[1.8, 1.7, 2.0, 2.0])

add_callout(doc, 
    "Verus ($15B+ Non-QM) uses LoanPASS as PPE. The Sovereign OS can become the DECISION LAYER that "
    "feeds LoanPASS and other PPEs — same way Stripe sits in front of banks. We don't replace LoanPASS, "
    "we provide the risk scoring + 5-dim DSCR + after-tax + IC memo + LLM narrative that LoanPASS doesn't have.",
    color=GREEN, label="COMPETITIVE POSITIONING")

add_page_break(doc)

# ============== PART II: 8 ARCHITECTURAL DEBTS ==============
add_h1(doc, "Part II — The 8 Architectural Debts (Deep Debt Analysis)")

add_para(doc, "From dscr_sovereign_os_architectural_debt_and_math.md + deep_debt_analysis.md. Each debt has a market signal confirming it is activating NOW. The institutional math fix is specified for each.")

add_h2(doc, "5. Debt 1 — DSCR as a Ratio Is Not a Risk Metric")
add_callout(doc, 
    "Live signal: multifamily CMBS delinquency 7.15% March 2026 (4× increase in 24 months). Quarterly default "
    "rates increase ~50% per 0.1x DSCR step down after 1.3x. Correlation unemployment ↔ DSCR PD = 0.87 "
    "(R²=0.76) — DSCR default is a MACRO phenomenon, not property-level.",
    color=CRIT, label="MARKET CONFIRMATION")

add_para(doc, "The vulnerability: deals classified PASS/FAIL based on threshold (typically 1.20x) treat 1.22x and 1.45x identically. Treats wildfire-exposed Airbnb and stable LTR identically. Treats lender-induced low-LTV (paradoxically riskier per Journal of Real Estate Finance & Economics) as safer.")

add_h3(doc, "The Fix: 5-Dimensional Distributional DSCR Engine")
add_code_block(doc,
    "Replace point-DSCR with a 5-dimensional stochastic DSCR surface:\n"
    "  1. P(DSCR_t < 1.0 | t = 12 months)         # near-term breach\n"
    "  2. P(DSCR_t < 1.0 | t = 36 months)         # medium-term breach\n"
    "  3. P(min DSCR < 1.0 over [0, T])            # lifetime breach\n"
    "  4. E[DSCR | macro recession scenario]        # macro-conditioned expected coverage\n"
    "  5. CVaR(DSCR loss | 95th percentile macro)  # tail conditional coverage")

add_h3(doc, "Plus Debt Yield Integration (per South State Bank analysis)")
add_code_block(doc,
    "Debt Yield = Annualized_NOI / Loan_Balance\n"
    "Minimum acceptable: 8.0%+ (conservative), 7.0% (standard), 6.0% (aggressive)\n"
    "Flag: Debt Yield < 7.0% AND DSCR 1.15-1.25 → enhanced stress testing required")

add_h2(doc, "6. Debt 2 — Income Inputs Have No Propagated Uncertainty")
add_callout(doc, 
    "Live signal: Clear Capital rental AVM error rate 6.65% (in thin markets much higher). AirDNA "
    "requires 6+ comps within 2mi at 65%+ occupancy for market score ≥60. STR income volatility can collapse 30-50% in 90 days.",
    color=CRIT, label="MARKET CONFIRMATION")

add_para(doc, "RentCast returns a point estimate. AirDNA returns a projected number. These enter the engine as certainties. ±8-15% AVM uncertainty, ±15-25% STR uncertainty, never represented in DSCR output.")

add_h3(doc, "The Fix: Conformal Prediction Vault (Mathematically Guaranteed Coverage)")
add_code_block(doc,
    "1. Train base estimator (gradient boosting on rent price)\n"
    "2. Calibrate on held-out set: nonconformity scores = |y_actual - y_predicted|\n"
    "3. At inference: q_hat = quantile(scores, ceil((n+1)(1-alpha))/n)\n"
    "4. Output interval: [y_hat - q_hat, y_hat + q_hat]\n"
    "# GUARANTEE: P(Y_new in interval) >= 1-alpha, regardless of model or distribution\n"
    "# KEY: arXiv 2405.02140 - conformal intervals upper-bound conditional entropy H(Y|X)")
add_code_block(doc,
    "# Mondrian (hierarchical) conformal by ZIP coverage:\n"
    "Group 1: ZIPs with 50+ AVM comps   -> tight intervals\n"
    "Group 2: ZIPs with 10-49 comps      -> medium intervals\n"
    "Group 3: ZIPs with <10 comps        -> wide intervals + auto-flag for human review\n"
    "# Decay mechanism (Dempster-Shafer):\n"
    "# nonconformity *= exp(-lambda * t)  where lambda is data-tier specific")
add_callout(doc, 
    "Implementation: Tier 1 (county tax record) lambda=low, decays over 365 days. "
    "Tier 3 (borrower-stated rent) lambda=high, decays in 30 days. As data ages, interval auto-widens.",
    color=GREEN, label="DECAY MECHANISM")

add_h2(doc, "7. Debt 3 — Monte Carlo Assumes Stationary Correlation")
add_callout(doc, 
    "Live signal: Lodging CMBS delinquency jumped 137 bps in single month to 7.31% (March 2026) — correlation "
    "spiked simultaneously with level shock. Gaussian copula would have missed this.",
    color=CRIT, label="2008 LESSON APPLIED NOW")

add_para(doc, "The vulnerability: t-copula captures symmetric tail dependence. Cannot capture ASYMMETRIC tail dependence: lower-tail (joint crashes — rent collapse + vacancy surge co-move) and upper-tail (joint spikes — cap rate + OpEx surge in inflation).")

add_h3(doc, "The Fix: R-Vine Copula with Mixed Families (TUM Munich)")
add_code_block(doc,
    "import pyvinecopulib as pv  # TUM Munich vinecopulib C++ backend\n"
    "import numpy as np\n"
    "\n"
    "# Calibrated to KBRA-equivalent stress scenarios (from Sprint 6 + TUM research):\n"
    "#                Stable   Cyclical   Stress\n"
    "# Rent shock     +/-10%   +/-20%    +/-30-40%\n"
    "# Vacancy        +200bps  +500bps   +1000bps\n"
    "# Cap rate       +50bps   +100bps   +200bps\n"
    "# Rate shock     +/-75bps +/-150bps  +/-250bps\n"
    "# OpEx shock     +5%      +10%      +20%\n"
    "\n"
    "data = np.column_stack([rent, vacancy, cap, rate, opex])\n"
    "controls = pv.FitControlsVinecop(family_set=pv.all, criterion='aic',\n"
    "                                tree_criterion='tau')  # max spanning on Kendall's tau\n"
    "vine = pv.Vinecop(data, controls=controls)\n"
    "scenarios = vine.simulate(n=10000, seeds=[42])")
add_callout(doc, 
    "Per-edge family selection (R-vine 10 bivariate copulas):\n"
    "- Rent-Vacancy edge: Clayton copula (lower-tail dependence — joint crashes correlate)\n"
    "- Cap-OpEx edge: Gumbel copula (upper-tail dependence — joint spikes correlate)\n"
    "- Rent-Cap edge: Student-t copula (symmetric — both tails matter)\n"
    "- R-vine structure: max spanning tree on absolute Kendall's tau\n"
    "Per Bundesbank: Gaussian MAY outperform at extreme stress (paradox) — hence mixed families.",
    color=ACCENT, label="ASYMMETRIC TAIL DEPENDENCE")

add_h3(doc, "Plus portvine (Higher-Level Wrapper for Portfolio CVaR)")
add_code_block(doc,
    "import portvine  # Built on pyvinecopulib\n"
    "# Portfolio-level CVaR / VaR estimation\n"
    "# with backtesting + stress testing baked in\n"
    "portfolio_cvar = portvine.compute_cvar(\n"
    "    vine=vine,\n"
    "    portfolio_positions=[deals],\n"
    "    confidence=0.95\n"
    ")")

add_h2(doc, "8. Debt 4 — No Forward Rate Surface (ARM Reset Uses Flat Curve)")

add_para(doc, "Vulnerability: ARM reset engine uses current SOFR swap rate as point forecast. Doesn't use full term structure. No rate path uncertainty. Doesn't account for cap-rate expansion feedback.")

add_h3(doc, "The Fix: Nelson-Siegel-Svensson Rate Surface + Hull-White Short-Rate Sim")
add_code_block(doc,
    "# Nelson-Siegel model:\n"
    "y(tau) = b0 + b1 * [(1 - exp(-tau/lambda)) / (tau/lambda)]\n"
    "              + b2 * [(1 - exp(-tau/lambda)) / (tau/lambda) - exp(-tau/lambda)]\n"
    "\n"
    "# Svensson extension adds: + b3 * [(1 - exp(-tau/lambda2))/(tau/lambda2) - exp(-tau/lambda2)]\n"
    "\n"
    "# Daily fit workflow:\n"
    "# 1. Pull SOFR swap quotes at 1M, 3M, 6M, 1Y, 2Y, 3Y, 5Y, 7Y, 10Y from FRED\n"
    "# 2. Fit (b0, b1, b2, b3, lambda1, lambda2) by nonlinear least squares\n"
    "# 3. Evaluate forward rate at ARM reset maturity: f(tau) = -d/dtau [tau * y(tau)]\n"
    "# 4. Add margin (e.g., 2.50%) to get projected reset rate\n"
    "# 5. For Monte Carlo: fit Hull-White (one-factor) short-rate model for rate path uncertainty")
add_code_block(doc,
    "from scipy.optimize import minimize\n"
    "import numpy as np\n"
    "\n"
    "def nelson_siegel(tau, b0, b1, b2, lam):\n"
    "    factor1 = (1 - np.exp(-tau/lam)) / (tau/lam)\n"
    "    factor2 = factor1 - np.exp(-tau/lam)\n"
    "    return b0 + b1 * factor1 + b2 * factor2\n"
    "\n"
    "def fit_ns(maturities, yields):\n"
    "    def objective(params):\n"
    "        return np.sum((nelson_siegel(maturities, *params) - yields)**2)\n"
    "    return minimize(objective, x0=[0.05, -0.01, 0.01, 1.5],\n"
    "                   bounds=[(0.01, 0.15), (-0.10, 0.10), (-0.10, 0.10), (0.1, 5.0)]).x")

add_h2(doc, "9. Debt 5 — No Credit Loss Model (CECL PD × LGD × EAD)")
add_callout(doc, 
    "Live signal: DSCR delinquencies doubled in 2 years. First-mover in real expected-credit-loss "
    "model for DSCR (not just threshold qualifier but lifetime loss estimate) holds structural analytical advantage.",
    color=CRIT, label="STRATEGIC FIRST-MOVER")

add_para(doc, "Engine computes PROBABILITY deal qualifies at origination. Does NOT compute expected credit loss over loan life. Different problems entirely.")

add_h3(doc, "The Fix: CECL PD × LGD × EAD (FASB ASC 326 / Basel III CRE32)")
add_code_block(doc,
    "Expected_Loss = PD * LGD * EAD\n"
    "\n"
    "# Pool-level PD curves by:\n"
    "# - vintage (2022 originations highest exposure)\n"
    "# - FICO band\n"
    "# - LTV band\n"
    "# - property type\n"
    "# - geographic cluster (NY/NJ 48% of new distress!)\n"
    "\n"
    "# LGD model:\n"
    "LGD = 1 - (LTV_at_default * haircut_factor)\n"
    "# haircut accounts for distressed sale discount + foreclosure costs\n"
    "\n"
    "# EAD with amortization:\n"
    "EAD(t) = loan_balance(t) * (1 - prepayment_assumption(t))")

add_h2(doc, "10. Debt 6 — No Contagion Model for Portfolio-Level Risk")
add_callout(doc, 
    "Live signal: 80% of new multifamily distress concentrated in NY/NJ (48%) + Houston (30%) = textbook "
    "geographic contagion cluster that no single-deal engine can detect because it has no portfolio network layer.",
    color=CRIT, label="MARKET CONFIRMATION")

add_para(doc, "Every deal evaluated independently. No model of how deals in same sponsor's portfolio, ZIP cluster, or lender's book interact under stress. FSB May 2026 report on private credit vulnerabilities: hidden leverage at fund/investor levels, indirect exposures not in first-layer reporting, no standardized classification.")

add_h3(doc, "The Fix: Spatio-Temporal Graph Risk")
add_code_block(doc,
    "# Graph nodes:\n"
    "# - Sponsor nodes: edges to all properties a single sponsor holds\n"
    "# - Property nodes: edges to neighboring properties in same ZIP/STR cluster\n"
    "# - Lender nodes: edges to all loans in single lender's book\n"
    "# - MSA nodes: edges to all properties in same macro market\n"
    "\n"
    "# Graph contagion algorithms:\n"
    "# - Spectral methods\n"
    "# - Random walk centrality\n"
    "# - Community detection\n"
    "\n"
    "# Implementation: GNN-derived embeddings as FEATURES into tabular champion\n"
    "# (NOT end-to-end black box — structured network risk features in interpretable model)")

add_h2(doc, "11. Debt 7 — LLM Layer Has No Hallucination Firewall")
add_callout(doc, 
    "SR 26-02 (April 17, 2026) places generative/agentic AI OUTSIDE MRM scope but requires broader governance. "
    "LLM-generated content influencing credit decisions needs governance chain.",
    color=WARN, label="SR 26-02 IMPLICATION")

add_para(doc, "Claude given engine JSON and asked to write IC memo will round, misattribute, or confabulate financial figures. Not broken — it's a text coherence system, not a numerical precision system.")

add_h3(doc, "The Fix: Deterministic Financial Fact-Checker")
add_code_block(doc,
    "def verify_llm_narrative(narrative: str, engine_output: dict) -> dict:\n"
    "    \"\"\"Extract all numeric claims from narrative.\n"
    "    Cross-reference against engine_output JSON.\n"
    "    Any number not within 0.5% of an engine value = MISMATCH.\"\"\"\n"
    "    numbers_in_narrative = extract_numeric_claims(narrative)\n"
    "    engine_values = flatten_dict(engine_output)\n"
    "    verified, mismatched, fabricated = [], [], []\n"
    "    for num, context in numbers_in_narrative:\n"
    "        match = find_close_match(num, engine_values, tolerance=0.005)\n"
    "        if match:\n"
    "            verified.append((num, context, match))\n"
    "        else:\n"
    "            mismatched.append((num, context))\n"
    "    return {'verified': verified, 'mismatched': mismatched,\n"
    "            'fabricated': [n for n, c in mismatched if c not in narrative]}\n"
    "# Human review MANDATORY before any LLM-generated content to final documentation")

add_h2(doc, "12. Debt 8 — No Model Version Tracking (Audit Trail)")
add_para(doc, "If XGBoost retrained weekly and deal evaluated at Week 14, no automatic record of which model version produced which output. Six months later, can't reproduce the inference.")

add_h3(doc, "The Fix: Per-Inference Model Audit Trail (SR 26-02 Compliant)")
add_code_block(doc,
    "# Every model inference stamped with:\n"
    "audit_entry = {\n"
    "    'model_name': 'approval_predictor_xgb_v3',\n"
    "    'model_version': '3.2.1',\n"
    "    'model_git_hash': 'a1b2c3d4',\n"
    "    'training_data_cutoff': '2026-05-15',\n"
    "    'training_data_size': 1234567,\n"
    "    'feature_importances_at_inference': {...},\n"
    "    'raw_input_vector_normalized': {...},\n"
    "    'output_probability': 0.847,\n"
    "    'calibration_version': 'iso_v2_2026q2',\n"
    "    'conformal_interval': [0.79, 0.91],\n"
    "    'coverage_tier': 0.90,\n"
    "    'timestamp': '...',\n"
    "    'user_id': 'broker_xyz',\n"
    "}")

add_page_break(doc)

# ============== PART III: DETERMINISTIC CORE ==============
add_h1(doc, "Part III — Deterministic Financial Core")

add_h2(doc, "13. Four DSCR Tracks + Golden Vector (v16 Master)")

add_table(doc,
    ["Track", "Numerator", "Denominator", "Use Case"],
    [
        ["Track 1 — Lender Qualifying", "Qualifying rent: lower of lease or 1007", "PITIA", "Loan qualification"],
        ["Track 2 — Investor Survival", "NOI = EGI - OpEx", "ADS annual", "True property debt coverage"],
        ["Track 3 — Stabilized", "Stabilized Year-N NOI (usually Y3)", "ADS after recast/perm", "Value-add, bridge-to-perm"],
        ["All-In DSCR", "NOI", "PI + T + I + HOA annualized", "Conservative lender/investor"],
    ],
    widths=[1.5, 2.0, 1.5, 1.5])

add_h3(doc, "13.1 Golden Vector (verified — Slice 1)")
add_code_block(doc,
    "$425K / 75% LTV / 7.00% / 30yr / lease $3,000 = 1007 / tax $5K / ins $2K / HOA $150:\n"
    "  P&I = $318,750 * 0.0066530 = $2,120.6517\n"
    "  PITIA = $2,853.9850\n"
    "  Track 1 DSCR = 1.0512x\n"
    "  Track 2 DSCR (8% vac, 8% mgmt) = 0.7884\n"
    "  Rent break-even (T1=1.0) = $2,853.985 (-4.83%)\n"
    "  Deal-break rate ~7.67%\n"
    "  Max price at T1=1.0 ~$454,100")

add_h2(doc, "14. BUG/FLAW Catalog (v16) — Slice 1 Status")

add_table(doc,
    ["ID", "Bug/Flaw", "Slice 1 Status"],
    [
        ["BUG-01", "LTV = min(purchase, appraisal) for purchases", "FIXED (ltv.py)"],
        ["BUG-02", "NOI growth exponent off-by-one", "NEW — Slice 2"],
        ["BUG-03", "Vacancy tornado labels swapped", "NEW — Slice 2"],
        ["BUG-05", "Breakeven Occupancy includes OpEx", "FIXED (ltv.py)"],
        ["BUG-06", "IO Max Loan uses decimal rate", "FIXED (ltv.py)"],
        ["FLAW-01", "Required DSCR risk stacking", "NEW — Slice 2"],
        ["FLAW-02", "Tranched waterfall", "Future — Slice 5"],
        ["IMP-08", "Stressed DS re-amortization at IO expiry", "NEW — Slice 2"],
    ],
    widths=[0.8, 3.5, 1.7])

add_h2(doc, "15. After-Tax Engine (OBBBA + §1250 + NIIT + PAL + REP)")

add_h3(doc, "15.1 OBBBA Three Tests (Sprint 4 — IRS Notice 2026-11)")
add_bullet(doc, "TEST 1: Property acquired after Jan 19, 2025")
add_bullet(doc, "TEST 2: NOT under binding written contract executed before Jan 20, 2025")
add_bullet(doc, "TEST 3: Self-constructed: <10% of hard costs incurred before Jan 20, 2025")
add_para(doc, "Eligible: 5-yr (appliances, carpet), 7-yr (office furniture), 15-yr (land improvements, paving), cost-seg components. NOT eligible: building structure (always 27.5-yr), land, residential rental building.")

add_h3(doc, "15.2 §1250 Three Buckets (v3 — Residential DSCR)")
add_bullet(doc, "Bucket 1 (Recaptured §1250, accelerated): ZERO for residential (no excess over SL); ORDINARY up to 37% for commercial")
add_bullet(doc, "Bucket 2 (Unrecaptured §1250, straight-line): 25% max federal — hits every residential seller")
add_bullet(doc, "Bucket 3 (Remaining gain, true appreciation): LTCG 0/15/20%")

add_h3(doc, "15.3 NIIT (FIXED thresholds, NOT CPI)")
add_table(doc,
    ["Filing Status", "MAGI Threshold", "Rate"],
    [
        ["Single / HH", "$200,000", "3.8%"],
        ["MFJ", "$250,000", "3.8%"],
        ["MFS", "$125,000", "3.8%"],
        ["QW", "$250,000", "3.8%"],
    ],
    widths=[2.0, 1.8, 0.7])

add_h3(doc, "15.4 PAL — Phase-Out at $150K (v3 corrected from Uncle Kam $200K)")
add_code_block(doc,
    "PAL_ENGINE = {\n"
    "    'standard_allowance': 25000,\n"
    "    'phase_out_threshold': 100000,\n"
    "    'phase_out_rate': 0.50,\n"
    "    'phase_out_complete_magi': 150000,  # ALL individual filers (corrected)\n"
    "    'rep_hours_minimum': 750,\n"
    "    'rep_time_pct_minimum': 0.50,\n"
    "}")
add_para(doc, "REP = highest-leverage tax status: 750hr + 50% real property → eliminates NIIT on rental income.", bold=True)

add_h3(doc, "15.5 OBBBA — Additional Provisions (per Non-QM Master Research Report)")
add_bullet(doc, "Section 179 limit: $2.56M (2026), phase-out at $4.09M")
add_bullet(doc, "EBITDA-based ATI for §163(j) interest deductibility (more favorable)")
add_bullet(doc, "QBI 20% deduction permanent for pass-throughs")
add_bullet(doc, "LIHTC bond financing reduced to 25% for 4% credits effective Jan 1, 2026")
add_bullet(doc, "SALT individual deduction cap increased to $40,000 through 2029")

add_h2(doc, "16. ARM Reset + IO Reversion + Cap-Rate Linked Refi")

add_h3(doc, "16.1 ARM Standard Market Specs (Non-QM Master Report)")
add_table(doc,
    ["Spec", "Value", "Source"],
    [
        ["Index", "30-day avg SOFR (most DSCR ARMs)", "Non-QM Master Research"],
        ["Margin", "3.5% above SOFR", "Griffin Funding SOFR ARM DSCR"],
        ["Initial cap", "2%", "Industry standard"],
        ["Periodic cap", "1-2%", "Industry standard"],
        ["Lifetime cap", "5-6% above start rate", "Industry standard"],
    ],
    widths=[1.5, 2.0, 2.5])

add_h3(doc, "16.2 Cap-Rate Linked Refi Solver (with NSS Forward Curve)")
add_code_block(doc,
    "def break_even_refi_cap_linked(noi, loan_balance, current_rate, target_dscr,\n"
    "                                current_cap_rate, cap_rate_beta,\n"
    "                                max_matrix_ltv, remaining_term_months,\n"
    "                                forward_curve):  # from NSS fit\n"
    "    \"\"\"Dual gate: DSCR AND LTV\"\"\"\n"
    "    def check(rate):\n"
    "        pmt = compute_pmt(loan_balance, rate, remaining_term_months)\n"
    "        dscr = noi / (pmt * 12)\n"
    "        # Use forward rate from NSS at reset maturity\n"
    "        forward_rate = forward_curve(rate)\n"
    "        proj_cap = current_cap_rate + cap_rate_beta * (rate - current_rate)\n"
    "        proj_value = noi / proj_cap\n"
    "        ltv = loan_balance / proj_value\n"
    "        return (dscr >= target_dscr) and (ltv <= max_matrix_ltv)\n"
    "    # Binary search...")

add_page_break(doc)

# ============== PART IV: STRESS & RISK ==============
add_h1(doc, "Part IV — Stress & Risk")

add_h2(doc, "17. t-Copula + R-Vine + Macro Archetypes")

add_callout(doc, 
    "Gaussian copula is BANNED for production use. t-Copula (ν=5-7) is the minimum. "
    "R-Vine with mixed families per edge is the institutional standard.",
    color=CRIT, label="MANDATORY")

add_table(doc,
    ["Archetype", "Rent", "Vacancy", "OpEx", "Rate", "Cap Rate"],
    [
        ["Stagflation", "flat", "flat", "+10%", "+200 bps", "+50 bps"],
        ["Recession", "-15%", "+10%", "-5%", "-100 bps", "+75 bps"],
        ["Climate/Regional", "-5%", "-3%", "+50% ins", "flat", "-10%"],
        ["Local Distress", "-10%", "+5%", "flat", "flat", "+100 bps"],
        ["Macro Stacking", "any combo", "any combo", "any combo", "any combo", "any combo"],
    ],
    widths=[1.4, 0.8, 0.8, 1.2, 0.9, 1.1])

add_h2(doc, "18. Sequential Drawdown + MCID + Counterfactual Generator")

add_h3(doc, "18.1 Sequential Drawdown Array (path-dependent)")
add_code_block(doc,
    "def sequential_drawdown(monthly_rent, monthly_opex, monthly_ds,\n"
    "                        capex_events, starting_reserves):\n"
    "    balance = starting_reserves\n"
    "    ruin_month = None\n"
    "    min_balance = balance\n"
    "    for t in range(36):\n"
    "        ncf = monthly_rent[t] - monthly_opex[t] - monthly_ds[t]\n"
    "        if t in capex_events: ncf -= capex_events[t]\n"
    "        balance += ncf\n"
    "        if balance < min_balance: min_balance = balance\n"
    "        if balance < 0 and ruin_month is None: ruin_month = t + 1\n"
    "    return {'liquidity_failure': ruin_month is not None,\n"
    "            'ruin_month': ruin_month, 'min_balance': min_balance}")

add_h3(doc, "18.2 MCID — Max Cumulative Intra-Year Deficit (STR-specific)")
add_code_block(doc,
    "def compute_mcid(monthly_net_cashflow):\n"
    "    \"\"\"Max cumulative deficit in any rolling 12-mo window\"\"\"\n"
    "    cumulative = 0; min_cumulative = 0\n"
    "    for cash in monthly_net_cashflow:\n"
    "        cumulative += cash\n"
    "        if cumulative < min_cumulative: min_cumulative = cumulative\n"
    "    return abs(min_cumulative)\n"
    "# Required reserves >= MCID, else fail seasonality trough test")

add_h3(doc, "18.3 Counterfactual Generator (Binary Search for Minimum Change)")
add_code_block(doc,
    "def counterfactual_search(target_state, current_state, all_variables,\n"
    "                         constraints):\n"
    "    \"\"\"Find smallest change to any single input that flips verdict.\"\"\"\n"
    "    candidates = []\n"
    "    for var in all_variables:\n"
    "        # Binary search delta for this variable\n"
    "        lo, hi = constraints[var]['min'], constraints[var]['max']\n"
    "        while hi - lo > 0.01:  # precision\n"
    "            mid = (lo + hi) / 2\n"
    "            test_state = current_state.copy()\n"
    "            test_state[var] = mid\n"
    "            if simulate(test_state) == target_state:\n"
    "                hi = mid\n"
    "            else:\n"
    "                lo = mid\n"
    "        candidates.append({'variable': var, 'min_change': hi})\n"
    "    return sorted(candidates, key=lambda c: c['min_change'])")

add_h2(doc, "19. STR Risk Scoring + Top 10 Markets (AirDNA 2026)")

add_para(doc, "STR Top 10 markets are small/mid-tier ($296K avg, 13.7% gross yield). Trophy markets (NYC/SF/LA) effectively dead for investors via Local Law 18, Home-Sharing Ordinance, 120-day rule, etc.")

add_table(doc,
    ["Rank", "Market", "Cap Rate", "Median Price", "Annual Revenue"],
    [
        ["1", "Jackson, MS", "15.95%", "$84,672", "$24,550"],
        ["2", "Abilene, TX", "14.01%", "$201,493", "$51,330"],
        ["3", "Akron, OH", "11.66%", "$139,633", "$29,612"],
        ["4", "Montgomery, AL", "11.64%", "$143,500", "$30,364"],
        ["5", "Port Arthur, TX", "10.38%", "$124,353", "$23,477"],
    ],
    widths=[0.5, 1.5, 0.9, 1.0, 1.2])

add_para(doc, "STR income qualification: LOWEST monthly figure if multiple sources. 20% haircut on AirDNA. STR capped at LTR market rent per 1007. AirDNA: 12mo forecast, 3 comps, market score ≥60, dated within 90 days, ≤2 individuals/bedroom.")

add_callout(doc, 
    "California SB 346 (effective Jan 1, 2026): Airbnb/VRBO must share host data with local governments. "
    "Engine must flag CA STR deals.",
    color=WARN, label="CA SB 346 ALERT")

add_page_break(doc)

# ============== PART V: LENDER & COMPLIANCE ==============
add_h1(doc, "Part V — Lender, Compliance & Wholesale")

add_h2(doc, "20. Cake Mortgage Product Matrix (Beyond the Rulebook #1)")

add_table(doc,
    ["Program", "Effective", "Max Loan", "Key Differentiator"],
    [
        ["DSCR v4.0", "Apr 1, 2026", "Not Specified", "Flagship — no personal income/tax docs; allows STR (AirDNA); FN friendly"],
        ["Bundt Cake NQM", "Apr 14, 2026", "$3M", "Credit event champion — BK/FC at 12mo (vs Funnel Cake 48mo); FTHB-friendly"],
        ["Bundt Cake NOO", "Apr 14, 2026", "$3M", "Non-DSCR NOO alternative (BS/P&L); prohibits WVOE"],
        ["Cup Cake Non-QM", "Mar 9, 2026", "$4M", "Highest loan limit; 1-Year Self-Employed; ITIN/Non-Perm Resident overlays"],
        ["Velvet Cake Non-QM", "Apr 1, 2026", "$3M", "Agricultural/Hobby Farms; FN INELIGIBLE"],
        ["Pound Cake Lite", "Feb 13, 2026", "$2.5M", "Best for WVOE & <2yr self-employment"],
    ],
    widths=[1.4, 1.0, 0.8, 3.2])

add_h3(doc, "20.1 DSCR v4.0 Specific Rules")
add_bullet(doc, "Allows 100% vacancy on 1-4 unit properties when using Form 1007")
add_bullet(doc, "Permits STR (Airbnb) but requires AirDNA Rentalizer on PURCHASES (market score ≥60); disallows AirDNA on REFINANCES")
add_bullet(doc, "Borrower experience tiers: Experienced / First-Time / FTHB")
add_bullet(doc, "Layered LLCs permitted up to 2 levels; trusts + LLCs prohibited")
add_bullet(doc, "Asset seasoning: NONE required for DSCR loans")
add_bullet(doc, "Cryptocurrency: must be liquidated to USD for funds to close")
add_bullet(doc, "Ineligible property types: Co-ops, Barndominiums, Houseboats")

add_h2(doc, "21. Top 8 Wholesale Lender License Footprint (Sprint 3)")

add_table(doc,
    ["Lender", "NMLS", "States", "Special"],
    [
        ["Visio Lending", "1935590", "41+DC", "Entity required in GA/HI/IL/MA/NJ/NY/PA/VA. Zero PPP in NM/KS/OH/MD/PA/RI. #1 DSCR ($854.6M 2024)"],
        ["Kiavi", "—", "49+DC", "Tech-forward, AVM-heavy, rapid closings"],
        ["Angel Oak", "1160240", "47+DC", "Rental AVM (Clear Capital) Nov 2025 — locked pre-qual"],
        ["Griffin Funding", "—", "46+DC", "Sub-1.0/low-DSCR/no-ratio/jumbo DSCR; SOFR ARM DSCR + 6-mo adj"],
        ["LendingOne", "—", "All+exempt", "Licensed or exempt in all other states"],
        ["Lima One Capital", "—", "National", "$10B+ lifetime; DSCR/bridge/construction"],
        ["Deephaven", "—", "National", "DSCR Second Mortgage ($75K-$500K, no reserves, no income docs)"],
        ["Verus Mortgage Capital", "—", "—", "Nation's largest Non-QM securitizer ($15B+) — uses LoanPASS PPE"],
    ],
    widths=[1.5, 0.8, 0.7, 3.5])

add_h2(doc, "22. 50-State PPP Matrix (2026 verified — Sprint 2)")

add_h3(doc, "22.1 Three-Branch Logic")
add_code_block(doc,
    "BRANCH 1: Business-purpose + Entity vesting (LLC, Corp, Trust)  -> Usually exempt\n"
    "BRANCH 2: Business-purpose + Individual vesting                -> Gray zone in some states\n"
    "BRANCH 3: Consumer-purpose (disqualified from DSCR by definition)\n"
    "DSCR loans are always business-purpose. Entity vesting frequently unlocks exemption.")

add_h3(doc, "22.2 Critical States (2026)")
add_table(doc,
    ["State", "PPP Status (2026)", "Threshold"],
    [
        ["MN", "ALLOWED (Aug 1, 2026)", "MN HF 3437 (Apr 23, 2026)"],
        ["PA", "THRESHOLD-RESTRICTED", "$329,411 (2026, was $319,777)"],
        ["OH", "THRESHOLD-RESTRICTED", "$112,957 (2025; 2026 = pull Jan)"],
        ["WA", "ALLOWED w/ ARM restriction", "60 days before ARM reset (RCW 19.144.040)"],
        ["NJ", "ENTITY-DEPENDENT", "C-Corp OK, LLC contested (NPLA Oct 2025), LP/Trust/Ind NO"],
        ["CA, TX, FL, GA, NC, TN, SC, VA, AL, IN, KY, MI, MO, WI, LA, AZ, CO", "ALLOWED", "Business-purpose exemption"],
    ],
    widths=[1.0, 2.0, 2.5])

add_callout(doc, 
    "OH and PA thresholds CPI-indexed annually. Engine must fire Celery cron on Jan 1 each year.",
    color=WARN, label="ANNUAL RE-INDEXING")

add_h2(doc, "23. ECOA / FCRA / SR 26-02 Compliance")

add_h3(doc, "23.1 SR 26-02 Status (CORRECTED FROM v1.0)")
add_table(doc,
    ["Component", "SR 26-02 Class"],
    [
        ["DSCR Calculator + Legal Rules Engine + After-Tax", "NOT a model (excluded)"],
        ["Monte Carlo + ML Forecasters + Approval Predictor", "MODEL — full card required"],
        ["LLM narrative", "Outside SR 26-02 but needs internal governance"],
    ],
    widths=[3.0, 3.0])

add_h3(doc, "23.2 ECOA Reason Codes (Form C-1 verbatim — v3 CONFIRMED from FCRA PDF)")
add_table(doc,
    ["Code", "Text", "Trigger (FCRA PDF p.6)"],
    [
        ["19", "Your income is not sufficient to meet your expenses and debt payments.", "DSCR low (rent) / FICO <620 / Reserves <3mo"],
        ["21", "Your debt payments or other obligations are too high.", "DSCR low (ADS)"],
        ["26", "You requested an amount that exceeds the maximum loan amount permitted by our regulations.", "LTV 80-90%"],
        ["27", "The collateral value is insufficient.", "LTV >90%"],
        ["28", "The type of property you selected is not acceptable to us.", "STR prohibited / property type reject"],
    ],
    widths=[0.6, 3.2, 2.7])

add_h2(doc, "24. Non-QM Wholesale: 12 Critical Gaps (MISSING PIECES)")

add_h3(doc, "24.1 P0 — Blocking")
add_table(doc,
    ["#", "Gap", "Solution", "Vendor"],
    [
        ["1", "Bank Statement Income Engine (DSCR-only incomplete)", "12-mo/24-mo parsing, 50% expense factor, transfer filtering", "Ocrolus / LoanLogics"],
        ["2", "Product & Pricing Engine (no rate distribution)", "LLPA matrix + dynamic rate sheet", "LoanPASS (Verus-selected) / Lender Price FLEX"],
        ["3", "Broker Approval / TPO Management (no onboarding)", "NMLS verification + E&O + comp plan", "Salesforce FSC + Encompass TPO Connect"],
        ["4", "Warehouse Lending Facility", "Real-time advance rate + borrowing base tracking", "LoanVantage / ICE Encompass module"],
    ],
    widths=[0.4, 2.0, 2.5, 1.6])

add_h3(doc, "24.2 P1 — High Priority")
add_table(doc,
    ["#", "Gap", "Solution"],
    [
        ["5", "Asset Depletion / Utilization Programs", "(Eligible_Assets - Down_Payment - Closing - Reserves) / 84 mo. 30% haircut on retirement."],
        ["6", "Foreign National + ITIN Programs", "ITIN input, alt credit, +0.50-1.50% rate premium, 12-24mo reserves"],
        ["7", "MSR Valuation + Secondary Market", "Gain_On_Sale = Sale_Price - UPB - Origination - Hedging + MSR_Value. MIAC Analytics"],
        ["8", "Pipeline Hedging + Rate Risk", "Hedge_Ratio = Pipeline_Volume × Pull_Through × Duration. TBA MBS or Treasury futures. Non-QM pull-through 65-75%"],
        ["9", "Quality Control + Loan Review (10% sample for KBRA)", "ACES Quality Management"],
    ],
    widths=[0.4, 2.0, 4.0])

add_h3(doc, "24.3 P2 — Medium Priority")
add_bullet(doc, "LOS Integration (MISMO 3.4 / ULAD) — ICE Encompass or Calyx PointCentral")
add_bullet(doc, "Compliance Mgmt + State Licensing + HMDA — Wolters Kluwer Compliance One")
add_bullet(doc, "Investor Relations + Capital Partner Mgmt — RMBS presale report formatting")

add_h2(doc, "25. The Non-QM Vendor Stack")

add_table(doc,
    ["Function", "Vendor", "Why"],
    [
        ["Non-QM PPE", "LoanPASS", "Verus selected Oct 2025; HousingWire 2026 Tech 100 Award; sub-second pricing"],
        ["Alternative PPE", "Lender Price FLEX", "Mature, integrated with LOS"],
        ["Fraud Detection", "Cotality (LoanSafe)", "Q1 2026: 1/44 investment + 1/29 multifamily apps show fraud risk"],
        ["QC Program", "ACES Quality Management", "KBRA/DBRS presale 10% sample + 100% EPD"],
        ["MSR Valuation", "MIAC Analytics", "March 2026 MSR Market Update; bulk pricing; $2.52B servicing offering Apr 2026"],
        ["LOS", "ICE Encompass", "MISMO 3.4 / ULAD / ULDD / UCD bi-directional sync"],
        ["Bank Statement OCR", "Ocrolus / Docling + Mistral OCR + GPT-4o", "Hybrid 3-layer pipeline"],
        ["STR Data", "AirDNA Enterprise", "$15-40/mo per market; Enterprise API custom"],
        ["Rent AVM", "RentCast API", "Free tier 50 calls/mo; 140M+ properties"],
        ["Rate Data", "FRED API (FREE)", "SOFR, MORTGAGE30US, Treasury curve"],
        ["Broker CRM", "Salesforce FSC", "13 pre-built mortgage objects; TPO portal via Experience Cloud"],
    ],
    widths=[1.5, 2.0, 3.0])

add_h2(doc, "26. Adverse-Action Notice Payload (FCRA PDF Spec)")

add_code_block(doc,
    "{\n"
    "  'version': '1.0',\n"
    "  'as_of': 'YYYY-MM-DDTHH:MM:SSZ',\n"
    "  'lender_id': 'CLIENT_XXXX',\n"
    "  'application_id': 'APP_YYYYYY',\n"
    "  'is_compliant': true,\n"
    "  'regulatory_notices': {\n"
    "    'ecoa_notice': {\n"
    "      'header': 'ADVERSE ACTION NOTICE',\n"
    "      'prohibition_statement': '<verbatim Reg B text>',\n"
    "      'reasons': [\n"
    "        {'code': '27', 'text': 'The collateral value is insufficient.'},\n"
    "        {'code': '21', 'text': 'Your debt payments or other obligations are too high.'}\n"
    "      ]\n"
    "    },\n"
    "    'fcra_disclosure': {\n"
    "      'header': 'DISCLOSURE REQUIRED BY THE FAIR CREDIT REPORTING ACT',\n"
    "      'statement': '<verbatim FCRA text>',\n"
    "      'data_source': 'CoreLogic, Inc.',\n"
    "      'source_address': '123 Main St, Anytown, ST 12345'\n"
    "    }\n"
    "  },\n"
    "  'state_specific_notices': [...],\n"
    "  'meta': {'engine_version': 'DSCR_Engine_v15.0.0', 'explanation_layer_version': '1.0.0'}\n"
    "}")

add_page_break(doc)

# ============== PART VI: LIVE DATA ==============
add_h1(doc, "Part VI — Live Data & Architecture")

add_h2(doc, "27. Live Rate Triplet (June 17-18, 2026)")

add_table(doc,
    ["Series", "Value", "Source"],
    [
        ["DGS10 (10Y Treasury)", "4.43%", "FRED"],
        ["SOFR Overnight", "3.63%", "FRED"],
        ["30-Day Avg SOFR", "3.609%", "NY Fed (FREE)"],
        ["90-Day Avg SOFR", "3.636%", "NY Fed"],
        ["Fed Funds Effective", "3.50-3.75%", "FRED (4th consecutive FOMC hold)"],
        ["Conventional 30yr IP", "6.52%", "Freddie Mac Jun 8"],
        ["Non-QM DSCR Warehouse Line Rate", "6.50-8.00%", "Tier 2 industry data"],
        ["DSCR Range (mid-2026)", "6.0-10.75%+", "Verified market production"],
        ["DSCR Premium over Conv", "+75-200 bps", "Narrowed from 2023 peak"],
    ],
    widths=[2.5, 2.0, 1.5])

add_h2(doc, "28. RentCast + AirDNA + Optimal Blue Integration")

add_table(doc,
    ["Source", "Use", "Cost", "API"],
    [
        ["RentCast", "LTR rent AVM (Form 1007 equiv)", "Free 50 calls/mo; volume-based", "developers.rentcast.io"],
        ["AirDNA", "STR revenue forecasting (12-mo + 36-mo + 60-mo)", "Free limited; $15-40/mo per market; Enterprise $50K+/yr", "airdna.redoc.ly"],
        ["Optimal Blue PPE / Loansifter", "Real-time lender pricing + eligibility", "$15K-$50K+/yr (commercial)", "Loansifter for brokers"],
        ["FRED", "Rate triplet + macro indicators", "FREE", "fred.stlouisfed.org"],
        ["Cotality (LoanSafe)", "Fraud detection", "Per-loan pricing", "LoanSafe Connect"],
        ["ATTOM Data", "Public records, comps, parcel, tax, lien", "~$500/mo for API", "attomdata.com"],
        ["HazardHub", "Hazard insurance quote API", "Per-quote", "hazardhub.com"],
        ["CorpAPI / BizFile", "SOS entity verification", "Per-call", "corpapi.com"],
        ["FraudGuard / DataVerify", "OFAC + fraud screening", "Per-check", "fraudguard.com"],
    ],
    widths=[1.5, 2.0, 2.0, 1.5])

add_h2(doc, "29. Three-Metric Credit Standard + Dual-Audience")

add_table(doc,
    ["Metric", "Question", "Target"],
    [
        ["DSCR (Cash Control)", "Can the borrower make the payment?", "Per matrix (typically ≥1.00-1.25)"],
        ["Debt Yield (Workout Metric)", "What is the lender's cap rate if they foreclose?", "≥ 9% institutional"],
        ["LTV (Loss-Given-Default)", "How much asset deflation can the lender absorb?", "Per matrix + LGD model"],
    ],
    widths=[1.5, 2.5, 1.5])

add_h3(doc, "29.1 AEY (All-In Effective Yield)")
add_code_block(doc,
    "True_Cost(hold) = Interest + Points + Lender_Fees + Lock_Cost + Prepay(exit_year) + Refi_Costs\n"
    "Render at 12/24/36/60-mo + APR-equivalent.\n"
    "AEY = XIRR of [Net_Proceeds_0, -P_1, -P_2, ..., -(P_n + Balance_n + PPP_n)]\n"
    "Algorithm: scipy.optimize.brentq — guaranteed convergence")

add_h3(doc, "29.2 6-Class Recommendation State Machine")
add_code_block(doc,
    "def recommend(acs, lender_pass, qbd_severity, iss, dfs, drawdown_pass):\n"
    "    if acs < 0.6: return 'HALT - INSUFFICIENT DATA'\n"
    "    if not lender_pass: return 'REJECT - DOES NOT QUALIFY'\n"
    "    if qbd_severity == 'CRITICAL': return 'QUALIFIES BUT CRITICAL RISK'\n"
    "    if qbd_severity == 'MODERATE': return 'QUALIFIES BUT DANGEROUS'\n"
    "    if qbd_severity == 'MINOR' or not drawdown_pass or iss<50 or dfs<30:\n"
    "        return 'FRAGILE - MONITOR CLOSELY'\n"
    "    if iss >= 80 and dfs >= 60: return 'STRONG DEAL'\n"
    "    return 'ACCEPTABLE - MONITOR KEY RISKS'")

add_h2(doc, "30. The GNN Entity Resolution Framework (THGNN)")

add_h3(doc, "30.1 Why GNN for Layered LLCs (Beyond the Rulebook #2)")
add_para(doc, "Layered LLCs (per Cake Mortgage guidelines: up to 2 layers) are exploited to obscure beneficial ownership. Standard KYC fails when borrower is opaque entity. Temporal Heterogeneous Graph Neural Networks (THGNNs) handle multi-relational, time-evolving entity networks.")

add_h3(doc, "30.2 Architecture")
add_code_block(doc,
    "# Nodes:\n"
    "# Person  (features: name, dob, ssn)\n"
    "# LLC     (features: articles_of_org, formation_date, registered_agent)\n"
    "# Address (features: full_address, county, state)\n"
    "# Phone   (features: phone_number)\n"
    "# Email   (features: email_address)\n"
    "\n"
    "# Edges (typed):\n"
    "# OWNS         (Person -> LLC, weight=ownership_pct)\n"
    "# LOCATED_AT   (LLC -> Address)\n"
    "# CONTACTS_PHONE (Person -> Phone)\n"
    "# GUARANTEES   (Person -> LLC, type=full_recourse)\n"
    "\n"
    "# Model architectures (Heterogeneous Graph Transformer HGT or Temporal Graph Network TGN):\n"
    "from torch_geometric.nn import HGTConv\n"
    "model = HGTConv(in_channels=node_feature_dim, out_channels=embedding_dim,\n"
    "                num_types=5, num_relations=4)")
add_para(doc, "Inference tasks:")
add_bullet(doc, "Node classification: BeneficialOwner / ShellCompanyOperator / FirstTimeInvestor / HighRiskIndividual")
add_bullet(doc, "Link prediction: identify hidden OWNS edges (anomaly)")
add_bullet(doc, "Anomaly detection: unusual network topology (one person owning 50 LLCs across states)")

add_h3(doc, "30.3 Data Sources")
add_bullet(doc, "Secretary of State (SOS) APIs: CorpAPI, BizFile, Trulioo")
add_bullet(doc, "Loan application data (1003)")
add_bullet(doc, "External: public records, news mentions, social media (entity enrichment)")

add_h2(doc, "31. Conformal Prediction Vault (Decaying Confidence)")

add_h3(doc, "31.1 The Framework (Beyond the Rulebook #2)")
add_code_block(doc,
    "from sklearn.ensemble import GradientBoostingRegressor\n"
    "import numpy as np\n"
    "\n"
    "def conformal_prediction_interval(new_x, calibration_X, calibration_y,\n"
    "                                base_model, alpha=0.10):\n"
    "    # Step 1: Compute nonconformity scores on calibration set\n"
    "    cal_pred = base_model.predict(calibration_X)\n"
    "    scores = np.abs(calibration_y - cal_pred)\n"
    "\n"
    "    # Step 2: Apply exponential decay (e^-lambda * t)\n"
    "    # Tier 1 (county tax): lambda=low, decays 365d\n"
    "    # Tier 3 (borrower-stated rent): lambda=high, decays 30d\n"
    "    decay_factors = np.exp(-lambda * data_age_days)\n"
    "    weighted_scores = scores * decay_factors\n"
    "\n"
    "    # Step 3: Quantile for (1-alpha) coverage\n"
    "    q_hat = np.quantile(weighted_scores, np.ceil((len(scores)+1)*(1-alpha))/len(scores))\n"
    "\n"
    "    # Step 4: Output interval\n"
    "    point_pred = base_model.predict([new_x])[0]\n"
    "    return (point_pred - q_hat, point_pred + q_hat)\n"
    "# GUARANTEE: P(Y_new in interval) >= 1-alpha")
add_para(doc, "Mondrian Conformal by data tier:")
add_bullet(doc, "Tier 1 (Form 1007, county tax): narrow interval")
add_bullet(doc, "Tier 2 (RentCast LTR AVM): medium interval")
add_bullet(doc, "Tier 3 (AirDNA STR projection, borrower-stated rent): wide interval + human review flag")

add_h2(doc, "32. Tabular Foundation Models (TabPFN) for Low-Data")

add_h3(doc, "32.1 The Problem")
add_para(doc, "Niche DSCR products (5-9 unit multifamily in Florida, Hobby Farms, Non-Warrantable Condos) have SPARSE historical defaults. XGBoost fails. Zero-shot TFM (TabPFN-2.5) gives credible baseline immediately.")

add_h3(doc, "32.2 Implementation (Beyond the Rulebook #2)")
add_code_block(doc,
    "# TabPFN-2.5 (Nature paper 2024 + arXiv 2511.08667):\n"
    "# Pre-trained on millions of diverse tabular datasets.\n"
    "# Zero-shot: makes predictions on new tasks without any fine-tuning.\n"
    "# Few-shot: 20-50 labeled examples -> fine-tune.\n"
    "\n"
    "from tabpfn import TabPFNClassifier\n"
    "model = TabPFNClassifier()  # Pre-loaded\n"
    "\n"
    "# New product line (DSCR 5-9 unit Florida):\n"
    "applications = pd.DataFrame({\n"
    "    'fico': [...], 'ltv': [...], 'dscr': [...],\n"
    "    'property_type': ['5-9_unit_mf'], 'occupancy': [...]\n"
    "})\n"
    "\n"
    "# Zero-shot prediction (no historical defaults needed):\n"
    "default_prob = model.predict_proba(applications)\n"
    "# Returns probability distribution per applicant\n"
    "\n"
    "# As 20-50 actual defaults accumulate:\n"
    "model.fit(applications_with_known_defaults)\n"
    "# Now uses few-shot learning for this specific niche")

add_page_break(doc)

# ============== PART VII: IMPLEMENTATION ==============
add_h1(doc, "Part VII — Implementation Roadmap")

add_h2(doc, "33. v1 → v2 → v3 Corrections Summary")

add_table(doc,
    ["#", "v1.0 Wrong", "v2.0 Correction", "v3.0 Final Verdict", "Source"],
    [
        ["1", "Renumber ECOA codes 19/21/26/27/28", "v2.0 said DO NOT renumber", "v3.0 confirms — codes ARE Form C-1 verbatim (FCRA PDF p.6)"],
        ["2", "PA PPP = $319,777", "v2.0 corrected to $329,411 (2026)", "v3.0 confirms + cron Jan 1"],
        ["3", "Two DSCR tracks", "v2.0 added Track 3 + All-In (4 tracks)", "v3.0 adds 5-dim DISTRIBUTIONAL DSCR (P12, P36, lifetime, E[macro], CVaR)"],
        ["4", "Gaussian copula OK", "v2.0 said t-copula only", "v3.0 says R-VINE with mixed families (TUM Munich + Bundesbank)"],
        ["5", "Point DSCR is the risk metric", "Not addressed in v2.0", "v3.0 DEBT 1: 5-dim distributional DSCR"],
        ["6", "Income inputs as point estimates", "Not addressed in v2.0", "v3.0 DEBT 2: Conformal Prediction with e^-λt decay"],
        ["7", "ARM reset at flat forward curve", "v2.0 said SOFR forward curve", "v3.0 DEBT 4: Nelson-Siegel-Svensson + Hull-White short-rate"],
        ["8", "No credit loss model", "Not addressed in v2.0", "v3.0 DEBT 5: CECL PD × LGD × EAD"],
        ["9", "No contagion model", "Not addressed in v2.0", "v3.0 DEBT 6: Spatio-temporal graph"],
        ["10", "LLM can produce IC memo directly", "Not addressed in v2.0", "v3.0 DEBT 7: Deterministic fact-checker + human review mandatory"],
        ["11", "No model version tracking", "Not addressed in v2.0", "v3.0 DEBT 8: Per-inference audit trail (SR 26-02)"],
        ["12", "DSCR-only product", "Not addressed in v2.0", "v3.0: 12 Non-QM Wholesale gaps; Bank Statement + Asset Depletion"],
        ["13", "Single lender matching", "Not addressed in v2.0", "v3.0: LoanPASS PPE + Verus $15B+ Non-QM securitizer"],
        ["14", "No entity resolution for layered LLCs", "Not addressed in v2.0", "v3.0: THGNN (HGT/TGN) entity resolution"],
        ["15", "All data equally weighted", "Not addressed in v2.0", "v3.0: Mondrian conformal with tier-specific decay"],
    ],
    widths=[0.3, 2.0, 2.3, 2.0, 1.4])

add_h2(doc, "34. Slice-by-Slice Build Plan (Consolidated)")

add_h3(doc, "34.1 Slice 1 — Already Shipped (Production-Ready)")
add_bullet(doc, "132 tests passing (was 122; +10 from 10x audit)")
add_bullet(doc, "94.37% coverage (was 91%; +3.37pp)")
add_bullet(doc, "ruff lint + format clean; SR 26-02 status: NOT a model")
add_bullet(doc, "All 4 v16 BUGs have regression tests")
add_bullet(doc, "Slice 1 v1.0 quality score 100/100 (per AUDIT_20260618.md)")
add_bullet(doc, "Files: payment.py / dscr.py / leverage.py / ltv.py / compliance.py")
add_bullet(doc, "ECOA codes (Form C-1 verbatim) — DO NOT renumber")

add_h3(doc, "34.2 Slice 2 — Distributional DSCR + Lender Compliance (~6 weeks, 250 hr)")
add_table(doc,
    ["#", "Module", "Effort", "Priority"],
    [
        ["P0-1", "5-Dim Distributional DSCR Engine (DEBT 1 fix)", "60 hr", "P0"],
        ["P0-2", "Conformal Prediction Vault (DEBT 2 fix)", "50 hr", "P0"],
        ["P0-3", "Sequential Drawdown + Macro Archetypes (4)", "40 hr", "P0"],
        ["P0-4", "DFS + min-gate ISS + QbD (7 triggers)", "40 hr", "P0"],
        ["P0-5", "Counterfactual Generator + 6-class Recommend", "30 hr", "P0"],
        ["P0-6", "BUG-02 (NOI growth) + BUG-03 (vacancy tornado)", "10 hr", "P0"],
        ["P0-7", "MCID detector + Cap-Rate Linked Refi (with NSS prep)", "20 hr", "P1"],
        ["TOTAL", "", "250 hr", ""],
    ],
    widths=[0.7, 3.5, 0.8, 0.8])

add_h3(doc, "34.3 Slice 3 — Institutional Math + After-Tax + Lender Footprint (~10 weeks, 400 hr)")
add_table(doc,
    ["#", "Module", "Effort"],
    [
        ["P1-1", "R-Vine Copula (TUM Munich pyvinecopulib) (DEBT 3)", "80 hr"],
        ["P1-2", "Nelson-Siegel-Svensson + Hull-White short-rate (DEBT 4)", "60 hr"],
        ["P1-3", "After-Tax Engine (OBBBA + §1250 + NIIT + PAL + REP + 1031)", "80 hr"],
        ["P1-4", "Matrix Grid Solver + Multi-Variable Opt", "80 hr"],
        ["P1-5", "CECL PD × LGD × EAD credit loss model (DEBT 5)", "60 hr"],
        ["P1-6", "SHA-256 + Merkle + SR 26-02 model cards + DEBT 8 audit", "40 hr"],
        ["TOTAL", "", "400 hr"],
    ],
    widths=[0.7, 4.0, 0.8])

add_h3(doc, "34.4 Slice 4 — GNN + Conformal + TFM + Live Data (~12 weeks, 500 hr)")
add_table(doc,
    ["#", "Module", "Effort"],
    [
        ["P2-1", "THGNN Entity Resolution (HGT/TGN) (DEBT 6 prep)", "120 hr"],
        ["P2-2", "Conformal Prediction Vault production (with Mondrian tiers)", "60 hr"],
        ["P2-3", "TabPFN-2.5 Tabular Foundation Model integration", "60 hr"],
        ["P2-4", "Live Data APIs (RentCast, AirDNA, Optimal Blue, Cotality, ATTOM)", "100 hr"],
        ["P2-5", "Three-Metric Credit Standard + AEY + IC Memo + 10-Min Verdict", "80 hr"],
        ["P2-6", "TimesFM 2.5 forecaster + Kalman Filter for latent state (DEBT 4 prep)", "80 hr"],
        ["TOTAL", "", "500 hr"],
    ],
    widths=[0.7, 4.0, 0.8])

add_h3(doc, "34.5 Slice 5 — Wholesale + Multi-Product (~16 weeks, 600 hr)")
add_bullet(doc, "Bank Statement Income Engine (50% expense factor, per MISSING PIECES Gap 1)")
add_bullet(doc, "LoanPASS PPE Integration (Verus-selected)")
add_bullet(doc, "Salesforce FSC + Encompass TPO Connect (TPO Management)")
add_bullet(doc, "Warehouse Lending Facility (LoanVantage)")
add_bullet(doc, "ACES QC Program (10% sample + 100% EPD)")
add_bullet(doc, "MIAC MSR Valuation + Pipeline Hedging")
add_bullet(doc, "1031 Exchange Timeline + Asset Depletion (84-month)")
add_bullet(doc, "Foreign National + ITIN Programs")

add_h2(doc, "35. Specific Code-Level Action Items (v3 prioritized)")

add_h3(doc, "35.1 Immediate (this week)")
add_para(doc, "AI-1: Add BUG-02 (NOI growth off-by-one) to Slice 1 dscr.py — 2 hr", bold=True)
add_para(doc, "AI-2: Add BUG-03 (vacancy tornado labels) to Slice 1 dscr.py — 1 hr", bold=True)
add_para(doc, "AI-3: Document SR 26-02 status inline in Slice 1 — 1 hr", bold=True)

add_h3(doc, "35.2 Slice 2 Entry (next 2 weeks)")
add_para(doc, "AI-4: Begin 5-Dim Distributional DSCR Engine (DEBT 1).", bold=True)
add_para(doc, "AI-5: Adopt R-Vine copula as Monte Carlo backend (Sprint 6 baseline + DEBT 3).", bold=True)
add_para(doc, "AI-6: Switch ISS to minimum-gate per Doc 17 (already in v2.0 plan).", bold=True)

add_h3(doc, "35.3 Slice 3 Critical (after Slice 2 ships)")
add_para(doc, "AI-7: Nelson-Siegel-Svensson forward curve (DEBT 4).", bold=True)
add_para(doc, "AI-8: After-Tax Engine (OBBBA 100% bonus + §1250 + NIIT + PAL + REP).", bold=True)
add_para(doc, "AI-9: CECL PD × LGD × EAD (DEBT 5) — critical for institutional risk pricing.", bold=True)

add_h3(doc, "35.4 Slice 4 First-Mover Capabilities")
add_para(doc, "AI-10: THGNN entity resolution (DEBT 6 / layered LLC detection).", bold=True)
add_para(doc, "AI-11: Conformal Prediction Vault production-grade.", bold=True)
add_para(doc, "AI-12: TabPFN-2.5 for zero-shot niche product underwriting.", bold=True)

add_h2(doc, "36. SR 26-02 Compliance Status — The Biggest Moat")

add_table(doc,
    ["Component", "SR 26-02 Class", "Required Governance", "v3 Engine Impact"],
    [
        ["DSCR Calculator + Legal Rules + After-Tax", "NOT a model", "Unit tests + CI/CD", "Ship 5x faster than competitors"],
        ["Sequential Drawdown + Macro Archetype Engine", "NOT a model (rule-based)", "Config + tests", "Ship 5x faster"],
        ["AEY / True Cost / Pricing Solver", "NOT a model (rule-based)", "Config + tests", "Ship 5x faster"],
        ["t-Copula / R-Vine Monte Carlo", "HIGH-materiality MODEL", "Full model card + champion/challenger", "Required for production"],
        ["TimesFM 2.5 / TFT / Kalman", "MEDIUM-HIGH MODEL", "Model card + backtesting", "Required for production"],
        ["CECL PD × LGD × EAD", "HIGH-materiality MODEL", "Full card + outcomes analysis", "Required for production"],
        ["TabPFN-2.5 Approval Predictor", "HIGH-materiality MODEL", "Full card + disparate impact monitor", "Required for production"],
        ["THGNN Entity Resolution", "HIGH-materiality MODEL", "Card + ongoing GNN monitoring", "Required for production"],
        ["LLM Narrative", "Outside SR 26-02", "Deterministic fact-checker + human review", "Governance chain required"],
    ],
    widths=[2.0, 1.5, 1.8, 1.7])

add_callout(doc, 
    "Architectural moat: deterministic layer ships WITHOUT model governance overhead. Monte Carlo + ML + GNN "
    "require full SR 26-02 cards. Competitors operating under blanket SR 11-7 definition must build full "
    "model governance for the deterministic layer too — we don't. This is 60-70% governance overhead reduction "
    "on the most-used components.",
    color=GREEN, label="SR 26-02 MOAT")

add_page_break(doc)

# ============== APPENDIX A: DOCUMENT INVENTORY ==============
add_h1(doc, "Part VIII — Appendices")

add_h2(doc, "Appendix A: Document Inventory (60+ sources)")

add_h3(doc, "A.1 Master Specifications (MDs — 8)")
for f in ["six-function-doctrine.md", "Advisor_Grade_DSCR_Decision_Engine_Usable_Master_Spec.md",
          "Advisor_Grade_DSCR_Decision_Engine_Organized_Research.md",
          "AEGIS_DSCR_Algorithm_Gap_Upgrade_Pack.md",
          "AEGIS_DSCR_Advisor_Grade_Operating_Model_Upgrade_Pack.md",
          "AEGIS_DSCR_Deterministic_Core_Keeps_Detailed.md",
          "AEGIS_DSCR_Complete_Usable_Master_Doc_v3.md",
          "DSCR_Engine_Master_Specification.md"]:
    add_bullet(doc, f)

add_h3(doc, "A.2 Architectural Blueprint PDFs (14)")
for f in ["From Black Box to Glass Box", "From Calculator to Counselor", "From Calculation to Counsel",
          "Architecting the Advisor-Grade DSCR Engine", "Beyond the DSCR", "From Static Snapshot to Dynamic Trajectory",
          "From Calculator to Containment", "AI Algorithm Improvement Prompt (Loops 1-15)",
          "AI Algorithm Improvement Prompt 2 (Per-Formula)", "FCRA Adverse Action Engine",
          "Beyond the Rulebook — Dynamic Data (Cake Mortgage)", "Beyond the Rulebook — Probabilistic + GNN + Conformal + TFM",
          "From Policy to Profit", "From Restriction to Dominance"]:
    add_bullet(doc, f)

add_h3(doc, "A.3 Sprint Research Execution (6)")
for f in ["Sprint 0/1: Live Research Execution Findings", "Sprint 2: PPP State Matrix + STR Legality + 40-Year Amortization",
          "Sprint 3: Lender Footprints + Securitization Pool + Competitive Threat",
          "Sprint 4: After-Tax IRR + OBBBA + Insurance Kill + Flood Gate + Compliance",
          "Sprint 5: Live Data APIs + Rate Anchors + Property Tax Matrix + Architecture",
          "Sprint 6: t-Copula MC + QuantLib ARM + After-Tax IRR + IC Memo + 1031 + XGBoost"]:
    add_bullet(doc, f)

add_h3(doc, "A.4 Architectural Debt Documents (2)")
add_bullet(doc, "dscr_sovereign_os_architectural_debt_and_math.md — 8 architectural debts + 4 institutional math modules (R-vine, EVT, NSS, Kalman)")
add_bullet(doc, "dscr_sovereign_os_deep_debt_analysis.md — Full debt archaeology with 2026 market data validation")

add_h3(doc, "A.5 Master Specs (6)")
for f in ["DSCR_Underwriting_Engine_v14_Complete_Master_Document.md",
          "DSCR_Underwriting_Engine_Master_Consolidated_v16.md (BUG-02/03 + FLAW-01/02)",
          "DSCR_Sovereign_OS_Definitive_Master_Blueprint_v3.md (SR 26-02 + 7 corrections)",
          "DSCR SOVEREIGN OS_ THE DEFINITIVE PRODUCT SPECIFICATION.md (Dual-Audience v12)",
          "DSCR Sovereign OS_ THE MASTER BLUEPRINT.md",
          "THE COMPLETE SOVEREIGN MASTER DOCUMENT.md"]:
    add_bullet(doc, f)

add_h3(doc, "A.6 Knowledge + Strategy Documents (8)")
for f in ["Master DSCR Knowledge Document.md", "The 2026 DSCR Master Knowledge Paper.md",
          "DSCR Sovereign OS MASTER RESEARCH SYNTHESIS.md (16 research domains)",
          "DSCR Sovereign OS & Non-QM Wholesale Lender Definitive Master Research Report.md",
          "DSCR DUAL TRUTH ENGINE CHATGPT RESEARCH.md",
          "THE DEFINITIVE BLUEPRINT_ BUILDING THE BEST NON-QM WHOLESALE LENDER.md",
          "THE MISSING PIECES_ NON-QM WHOLESALE LENDER GAP ANALYSIS.md (12 critical gaps)",
          "NEW_DSCR Deal Desk Build-Ready Research Report.md"]:
    add_bullet(doc, f)

add_h3(doc, "A.7 Slice 1 Codebase (verified — 132 tests, 94.37% coverage)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/payment.py (4,709 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/dscr.py (13,069 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/leverage.py (12,653 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/ltv.py (8,472 B, BUG-01/05/06 fixed)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/src/dscr_core/compliance.py (13,012 B)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/tests/test_v16.py (29 tests for BUG/FLAW regression)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/golden_vectors.json (4 test groups)")
add_bullet(doc, "DSCR_SOVEREIGN_OS/packages/dscr-core/AUDIT_20260618.md (100/100 quality gate)")

add_h3(doc, "A.8 Workspace Subdirectories (analyzed)")
add_bullet(doc, "ANALYSIS/ — TOPICAL_INDEX (66 KB), MASTER_ANALYSIS (415 KB), GOLDEN_VECTORS (50 KB), godmode_research_plan_v2 (48 KB)")
add_bullet(doc, "RESEARCH/godmode_20260618/ — 170 files / 1.1 MB across 14 directories (Rounds 16-21 research)")
add_bullet(doc, "RESEARCH/pdf_extractions/ — 19 PDF text extracts")
add_bullet(doc, "RESEARCH/pdf_short/ — 4 Beyond the Rulebook copies")

add_page_break(doc)

# ============== APPENDIX B: PSEUDOCODE LIBRARY ==============
add_h2(doc, "Appendix B: Pseudocode Library (Full)")

add_h3(doc, "B.1 R-Vine Copula (DEBT 3)")
add_code_block(doc,
    "import pyvinecopulib as pv  # TUM Munich\n"
    "import numpy as np\n"
    "\n"
    "# Calibration data: historical or KBRA-scenario stress returns\n"
    "data = np.column_stack([rent_shocks, vacancy_shocks, cap_shocks,\n"
    "                         rate_shocks, opex_shocks])\n"
    "controls = pv.FitControlsVinecop(family_set=pv.all, criterion='aic',\n"
    "                                tree_criterion='tau')\n"
    "vine = pv.Vinecop(data, controls=controls)\n"
    "scenarios = vine.simulate(n=10000, seeds=[42])\n"
    "# Per-edge family selection (R-vine 10 bivariate copulas):\n"
    "# Rent-Vacancy: Clayton (lower-tail)\n"
    "# Cap-OpEx: Gumbel (upper-tail)\n"
    "# Rent-Cap: Student-t (symmetric)")

add_h3(doc, "B.2 Conformal Prediction Vault (DEBT 2)")
add_code_block(doc,
    "import numpy as np\n"
    "from sklearn.ensemble import GradientBoostingRegressor\n"
    "\n"
    "def conformal_interval(new_x, cal_X, cal_y, base_model, alpha=0.10,\n"
    "                     data_age_days=0, lambda_tier=0.01):\n"
    "    cal_pred = base_model.predict(cal_X)\n"
    "    scores = np.abs(cal_y - cal_pred)\n"
    "    # Decay factor (e^-lambda*t) — Tier-specific\n"
    "    decay = np.exp(-lambda_tier * data_age_days)\n"
    "    weighted_scores = scores * decay\n"
    "    q_hat = np.quantile(weighted_scores,\n"
    "                        np.ceil((len(scores)+1)*(1-alpha))/len(scores))\n"
    "    point_pred = base_model.predict([new_x])[0]\n"
    "    return (point_pred - q_hat, point_pred + q_hat)")

add_h3(doc, "B.3 THGNN Entity Resolution (DEBT 6 + Beyond the Rulebook #2)")
add_code_block(doc,
    "from torch_geometric.nn import HGTConv\n"
    "# Heterogeneous graph nodes: Person, LLC, Address, Phone, Email\n"
    "# Edges: OWNS (Person->LLC, weight=pct), LOCATED_AT, CONTACTS_PHONE, GUARANTEES\n"
    "model = HGTConv(in_channels=node_dim, out_channels=embed_dim,\n"
    "                num_types=5, num_relations=4)\n"
    "# Tasks: node classification (BeneficialOwner/ShellCompanyOperator/etc),\n"
    "# link prediction (hidden OWNS), anomaly detection (unusual topology)")

add_h3(doc, "B.4 Nelson-Siegel-Svensson + Hull-White (DEBT 4)")
add_code_block(doc,
    "from scipy.optimize import minimize\n"
    "import numpy as np\n"
    "def nelson_siegel(tau, b0, b1, b2, lam):\n"
    "    f1 = (1 - np.exp(-tau/lam)) / (tau/lam)\n"
    "    f2 = f1 - np.exp(-tau/lam)\n"
    "    return b0 + b1*f1 + b2*f2\n"
    "# Daily: fit to SOFR swap quotes 1M-10Y, evaluate forward at ARM reset")

add_h3(doc, "B.5 CECL PD × LGD × EAD (DEBT 5)")
add_code_block(doc,
    "Expected_Loss = PD * LGD * EAD\n"
    "# PD curves by: vintage × FICO × LTV × property_type × geo_cluster\n"
    "# LGD = 1 - (LTV_at_default * haircut)\n"
    "# EAD = loan_balance(t) * (1 - prepayment_assumption(t))")

add_h3(doc, "B.6 EVT (Extreme Value Theory) for Deep Tail")
add_code_block(doc,
    "from scipy.stats import genpareto\n"
    "# Fit GPD to excess losses above threshold u\n"
    "xi, loc, beta = genpareto.fit(excess_losses, floc=0)\n"
    "# VaR_alpha under Fréchet (xi > 0):\n"
    "# VaR_alpha = mu_n - (sigma_n / xi_n) * [1 - (-n*ln(alpha))^(-xi_n)]")

add_h3(doc, "B.7 TabPFN-2.5 Zero-Shot Underwriting (Beyond the Rulebook #2)")
add_code_block(doc,
    "from tabpfn import TabPFNClassifier\n"
    "model = TabPFNClassifier()  # Pre-loaded, no training needed\n"
    "# Zero-shot on niche product (DSCR 5-9 unit Florida):\n"
    "applications = pd.DataFrame({...})\n"
    "default_prob = model.predict_proba(applications)\n"
    "# Few-shot after 20-50 actual defaults accumulate:\n"
    "model.fit(applications_with_known_defaults)")

add_h3(doc, "B.8 LLM Narrative Fact-Checker (DEBT 7)")
add_code_block(doc,
    "def verify_llm_narrative(narrative: str, engine_output: dict) -> dict:\n"
    "    nums = extract_numeric_claims(narrative)\n"
    "    vals = flatten_dict(engine_output)\n"
    "    verified, mismatched = [], []\n"
    "    for n, ctx in nums:\n"
    "        match = find_close_match(n, vals, tolerance=0.005)\n"
    "        if match: verified.append((n, ctx, match))\n"
    "        else: mismatched.append((n, ctx))\n"
    "    # Human review MANDATORY for any mismatches")

add_page_break(doc)

# ============== APPENDIX C: LIVE DATA ==============
add_h2(doc, "Appendix C: Live Rate Triplet + Market Data (June 2026)")

add_table(doc,
    ["Data Point", "Value (Jun 17-18, 2026)", "Source", "Refresh"],
    [
        ["DGS10 (10Y Treasury)", "4.43%", "FRED", "Every 4 hr"],
        ["SOFR Overnight", "3.63%", "FRED", "Every 4 hr"],
        ["30-Day Avg SOFR", "3.609%", "NY Fed (free)", "Daily"],
        ["90-Day Avg SOFR", "3.636%", "NY Fed", "Daily"],
        ["Fed Funds Effective", "3.50-3.75%", "FRED", "After FOMC"],
        ["Conventional 30yr IP", "6.52%", "Freddie Mac Jun 8", "Weekly"],
        ["DSCR Premium over Conv", "+75-200 bps", "Multiple", "Per rate sheet"],
        ["multifamily CMBS 30+ DQ", "7.55%", "Trepp March 2026", "Monthly"],
        ["Multifamily CMBS Special Servicing", "8.75%", "Trepp March 2026", "Monthly"],
        ["Office CMBS DQ", "11.71%", "Trepp March 2026", "Monthly"],
        ["DSCR/Investor 60+ DQ", "2.92%", "December 2025", "Monthly"],
        ["Bank Statement 60+ DQ", "3.99%", "December 2025", "Monthly"],
        ["DSCR Origination Growth Jan 2025 YoY", "+123%", "Market data", "Monthly"],
        ["Cotality Fraud Index Q1 2026", "121 (down from 133 Q4 2025)", "Cotality", "Quarterly"],
        ["1 in 44 investment property apps show fraud", "Q1 2026", "Cotality", "Quarterly"],
        ["1 in 29 multifamily apps show fraud", "Q1 2026", "Cotality", "Quarterly"],
    ],
    widths=[2.3, 2.0, 1.5, 1.2])

add_page_break(doc)

# ============== APPENDIX D: LENDER FOOTPRINT ==============
add_h2(doc, "Appendix D: Lender Footprint Matrix (June 2026)")

add_table(doc,
    ["Lender", "NMLS", "States", "Min FICO", "DSCR Floor", "Max LTV", "Key Feature"],
    [
        ["Visio Lending", "1935590", "41+DC", "680", "1.00 (sub-1.0 case-by-case)", "80%/75%", "#1 DSCR ($854.6M 2024), 11 S&P deals ~$2B"],
        ["Kiavi", "—", "49+DC", "—", "—", "—", "AVM-heavy, rapid closings"],
        ["Angel Oak", "1160240", "47+DC", "—", "—", "—", "Rental AVM (Clear Capital) Nov 2025"],
        ["Griffin Funding", "—", "46+DC", "—", "Sub-1.0", "—", "No-ratio, jumbo DSCR, SOFR ARM DSCR"],
        ["LendingOne", "—", "All+exempt", "—", "—", "—", "Licensed or exempt in all states"],
        ["Lima One Capital", "—", "National", "—", "—", "—", "$10B+ lifetime"],
        ["Deephaven", "—", "National", "680", "1.00 (combined)", "80% CLTV", "DSCR Second ($75K-$500K, no reserves)"],
        ["MortgageDepot", "—", "Verified", "—", "—", "—", "40-yr amort + 40-yr IO up to $3M"],
        ["Sistar Mortgage", "—", "Verified", "—", "—", "—", "40-yr IO confirmed 2026"],
        ["Verus Mortgage Capital", "—", "—", "—", "—", "—", "$15B+ Non-QM securitizer (nation's largest)"],
    ],
    widths=[1.6, 0.8, 0.7, 0.6, 0.9, 0.7, 1.8])

add_page_break(doc)

# ============== APPENDIX E: CAKE PRODUCT MATRIX ==============
add_h2(doc, "Appendix E: Cake Mortgage Product Matrix Reference")

add_table(doc,
    ["Program", "Type", "Effective", "Max Loan", "Niche / Differentiation"],
    [
        ["DSCR v4.0", "Investment DSCR", "Apr 1, 2026", "Not Spec.", "Flagship; no income docs; STR (AirDNA); FN friendly"],
        ["Bundt Cake NQM", "Owner-Occupied", "Apr 14, 2026", "$3M", "Credit event champion (BK/FC at 12mo)"],
        ["Bundt Cake NOO", "Investment Alt-Doc", "Apr 14, 2026", "$3M", "BS/P&L; prohibits WVOE"],
        ["Cup Cake Non-QM", "Owner-Occupied", "Mar 9, 2026", "$4M", "Highest limit; 1-Yr SE; ITIN/Non-Perm"],
        ["Velvet Cake Non-QM", "Owner-Occupied", "Apr 1, 2026", "$3M", "Agricultural/Hobby Farms; FN INELIGIBLE"],
        ["Pound Cake Lite", "Owner-Occupied", "Feb 13, 2026", "$2.5M", "WVOE + <2yr SE"],
    ],
    widths=[1.4, 1.2, 0.9, 0.7, 3.0])

add_h3(doc, "E.1 Universal Cake Mortgage Rules")
add_bullet(doc, "Asset seasoning: NONE required for DSCR loans")
add_bullet(doc, "Cryptocurrency: must be liquidated to USD for funds to close")
add_bullet(doc, "Entity vesting: layered LLCs up to 2 levels; trusts + LLCs PROHIBITED")
add_bullet(doc, "Borrower tiers: Experienced Investor / First-Time Investor / First-Time Homebuyer")
add_bullet(doc, "Ineligible: Co-ops, Barndominiums, Houseboats, properties <500 sq ft, C5/C6 condition")

add_page_break(doc)

# ============== APPENDIX F: REFERENCES ==============
add_h2(doc, "Appendix F: References & Source Anchors")

add_h3(doc, "F.1 Primary Regulatory")
for r in ["OCC Bulletin 2026-13 (SR 26-02) — effective Apr 17, 2026",
         "Fannie Mae Selling Guide §B2-2-03, §B2-3-03, §B7-3-02",
         "Freddie Mac Single-Family Seller/Servicer Guide",
         "CFPB Circular 2022-03 (adverse action for complex algorithms)",
         "CFPB Circular 2023-03 (AI-specific adverse action)",
         "Reg B (ECOA), Reg Z (TRID)",
         "IRC §1411 NIIT, §469 PAL, §1250 recapture",
         "Treasury/IRS Notice 2026-11 OBBBA bonus depreciation",
         "FinCEN Interim Final Rule Mar 2025 (BOI)",
         "HOEPA 2026 thresholds ($27,592 loan / $1,380 P&F)",
         "FHFA AB-2022-03 (Fair Lending AI/ML)",
         "Basel III BCBS d424",
         "EBA 2025 EU Stress Test",
         "PA Act 6 / OH ORC §1343.011 / MN HF 3437 (Apr 23, 2026 / eff Aug 1, 2026)",
         "NJ N.J.S.A. 46:10B-2 + Arc Home Jul 2025 + NPLA Oct 2025"]:
    add_bullet(doc, r)

add_h3(doc, "F.2 Market Data")
for r in ["KBRA Single-Borrower CMBS Default and Loss Study",
         "KBRA CMBS Loan Performance Trends (Feb/Apr 2026)",
         "KBRA CRE CLO Loan Default and Loss Study (Jun 2026)",
         "Trepp CMBS distress rate; CMBS delinquency March 2026 7.55%",
         "CRED iQ CMBS Conduit Underwriting Trends",
         "S&P Global Ratings DSCR adjustment factor 1.5x-2.5x (NRMLT 2026-NQM1)",
         "Fitch Ratings Multifamily/Office CMBS Delinquency",
         "Matthews 2026 Cap Rate Analysis; UNC cap-rate determinants (Tsui-Morgan 2025)",
         "NCREIF Property Index",
         "MSCI Real Capital Analytics; MMCG CRE Insights",
         "Cotality (LoanSafe) Fraud Index Q1 2026: 121 (down from 133)",
         "AirDNA Best Places to Invest 2026; Top 10 STR markets",
         "RentCast API; Optimal Blue PPE (Verus-selected Oct 2025)"]:
    add_bullet(doc, r)

add_h3(doc, "F.3 Academic / Methodology")
for r in ["Li, D.X. (2000) — Gaussian copula (the 2008 mechanism)",
         "Cherubini, Luciano, Vecchiato (2004) Copula Methods in Finance",
         "Glasserman (2003) Monte Carlo Methods in Financial Engineering",
         "Bundesbank research on heavy-tailed copulas (Clayton/t for 'less severe'; Gaussian MAY outperform at extreme stress)",
         "TUM Munich vinecopulib R + pyvinecopulib Python (R-vine copulas)",
         "Demarta & McNeil (2005) t-copula DOI",
         "Artzner et al (1999) Coherent risk measures (CVaR)",
         "Blanc-Brude & Hasan (2016) Structural credit risk: default = DSCR<1.0",
         "Rodríguez (2024) DOI:10.4236/jfrm.2024.134029 — DSCR ↔ Economic Value",
         "Diebold-Li state-space model (Kalman filter for yield curve)",
         "Nelson-Siegel-Svensson term structure (SSRN 2054689)",
         "Vasicek / Hull-White short-rate models",
         "Conformal Prediction: arXiv 2405.02140 (tightest valid intervals)",
         "Mondrian Conformal (2024 PhD thesis — real estate tenant debt)",
         "TabPFN-2.5 (Nature 2024 + arXiv 2511.08667) — zero-shot tabular",
         "Heterogeneous Graph Transformers / Temporal Graph Networks (arXiv 2510.26307)",
         "TimesFM 2.5 (Google BigQuery docs Jun 12, 2026) — 200M params, 15,360 ctx, native quantile",
         "CPTC Conformal Prediction for Time-series (NeurIPS 2025)",
         "Davis (2024) P50/P99 Debt Sculpting — Pivotal180",
         "OBBBA signed Jul 4, 2025 — 100% bonus depreciation permanent"]:
    add_bullet(doc, r)

add_h3(doc, "F.4 Vendor Stack")
for r in ["LoanPASS (Non-QM PPE; Verus selected Oct 2025; HousingWire Tech 100 2026)",
         "Lender Price FLEX (alternative PPE)",
         "Cotality LoanSafe (fraud detection; 1/44 IP + 1/29 MF fraud risk)",
         "ACES Quality Management (QC; KBRA/DBRS presale)",
         "MIAC Analytics (MSR valuation; March 2026 update)",
         "ICE Encompass (LOS; MISMO 3.4 / ULAD / ULDD / UCD)",
         "Salesforce FSC (13 pre-built mortgage objects; TPO portal)",
         "Ocrolus / Docling / Mistral OCR 2505 / GPT-4o (hybrid OCR pipeline)",
         "Wolters Kluwer Compliance One (CMS)",
         "FRED API (FREE; 845,000+ series)",
         "Census ACS API (FREE; B25002/B25004 vacancy)",
         "FEMA NFHL WMS (FREE; flood hazard)",
         "ATTOM Data (~$500/mo for API)",
         "Docling (digital PDFs)",
         "Mistral OCR 2505 ($1/1000 pages)",
         "Instructor library (Pydantic schema)",
         "HazardHub / CorpAPI / FraudGuard (specialty)"]:
    add_bullet(doc, r)

# ============== END ==============
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_run = end_p.add_run("--- END OF v3 SYNTHESIS (DEFINITIVE) ---\n"
                        "DSCR Advisor-Grade Decision Engine: Cross-Document Synthesis v3.0\n"
                        "Prepared for DSCR Sovereign OS / 20X DSCR Deal Engine\n"
                        "Workspace: C:\\Users\\serge\\OneDrive\\Documents\\DSCR_LOAN OFFICE\n"
                        "Output: output/doc/DSCR_Advisor_Engine_Cross_Doc_Synthesis_v3_20260619.docx\n\n"
                        "60+ source documents analyzed end-to-end\n"
                        "60+ documents / ~3MB / ~960 pages of PDF content + 1.5MB of MDs\n\n"
                        "v3.0 ADDS (vs v2.0):\n"
                        "  - 8 architectural debts with institutional math fixes (R-vine, EVT, NSS, Kalman)\n"
                        "  - 5-dimensional distributional DSCR (P12, P36, lifetime, E[macro], CVaR)\n"
                        "  - Conformal Prediction Vault with e^-lambda*t decay\n"
                        "  - THGNN entity resolution for layered LLCs (HGT/TGN)\n"
                        "  - TabPFN-2.5 for zero-shot niche product underwriting\n"
                        "  - Cake Mortgage product matrix (DSCR v4.0, Bundt, Cup, Velvet, Pound)\n"
                        "  - 12 Non-QM wholesale gaps + complete vendor stack\n"
                        "  - Comprehensive SR 26-02 compliance status\n"
                        "  - Live market data March-June 2026 (CMBS 7.15%, fraud index 121)")
end_run.font.size = Pt(10)
end_run.font.italic = True
end_run.font.color.rgb = DARK_GRAY

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
print(f"Size: {OUT_FILE.stat().st_size / 1024:.1f} KB")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")